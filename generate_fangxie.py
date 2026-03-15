# -*- coding: utf-8 -*-
"""
百家号文案仿写生成脚本
"""

import anthropic
from docx import Document
from datetime import datetime
import os
import re

# 初始化API客户端
client = anthropic.Anthropic(
    base_url="https://yunyi.rdzhvip.com/claude",
    api_key="A68YGT53-41HW-8GT7-353U-MJVDNY8NE5KX"
)

# 带货素材
product_info = """
商品：黄铜貔貅摆件（一对）

核心卖点：
1. 公貔貅主招财，脚踩绣球，象征主动吸纳外界财气
2. 母貔貅主守财，脚踩如意，象征稳固家中财库
3. 成对摆放寓意"招财守财，财富双收"

材质工艺：
- 黄铜铸造，真材实料，分量十足
- 多层工序打磨，加厚外壁
- 纳米镀层，防水防污，细节精致

适用场景：
- 家中客厅财位：寓意家庭兴旺和睦
- 办公桌/前台：祝事业顺利
- 收银台：祝生意兴隆
"""

# 参考文案（用于分析结构和爆点）
reference = """你根本无法想象命运对你的宠爱有多么夸张，它像一条从天际倾泻下来的光河，把你整个笼罩其中..."""

prompt = f"""【重要指令】你是一个文案生成器，直接输出3篇完整的情感带货文案，不要询问任何问题，不要输出任何解释性文字。

【参考文案分析】
参考文案的核心爆点：
1. 用"命运宠爱"的概念让读者感到被认可和特别
2. 将"孤独"重新定义为"被净化的轨迹"，化负面为正面
3. 用灯塔、星辰等意象塑造读者的独特价值感
4. 最后自然过渡到貔貅产品，作为"祥瑞信物"

【你的任务】
基于以上爆点结构，创作3篇全新的情感带货文案，主题围绕"善待自己、肯定自我价值"，最后自然引导到黄铜貔貅产品。

【带货商品信息】
{product_info}

【写作模式】
这是第二人称对话式文案，用"你"贯穿全文，像朋友促膝长谈。

【3篇文案的主题方向（必须不同）】
第1篇：从"被忽视的付出"切入，肯定读者默默承受的一切，最后引导到貔貅作为对自己的犒赏
第2篇：从"中年觉醒"切入，讲述人到中年才懂得善待自己的重要性，貔貅作为新阶段的开始
第3篇：从"独处的智慧"切入，将独处重新定义为内心富足的表现，貔貅作为这份智慧的见证

【开头要求 - 3篇必须完全不同】
第1篇开头类型：直接点破钩子 - 直接说出读者心里的话
第2篇开头类型：场景代入钩子 - 描述一个具体场景让读者代入
第3篇开头类型：反常识钩子 - 说一个颠覆认知的观点

【爆点提炼硬规则】（必须执行，不能漏）
请先从参考文案中提炼并融入以下爆点，再开始写正文：
1) 开场必须轮换：3篇使用不同开场样式（直呼点名/好友对话/场景切入/反问切入/共情陈述），不能固定“孩子/有缘人/天选之人/道友”。
2) 情绪识别与共情：先承认“你很累、你在硬扛、你被消耗”，再给理解与安抚。
3) 国学意象背书：可用“节律、时机、守拙、厚德、积善”等表达做短论证，避免空泛堆砌。
4) 窗口期与行动感：强调“现在是收拢状态的窗口”，但禁止绝对承诺和恐吓式表述。
5) 物件定位：把产品写成“行动锚点/生活仪式感/状态提醒”，不是神奇道具。
6) 结尾引导：引流段80-140字；置顶写清“看第几条+能得到什么”；橱窗写清“优先看哪类+怎么选”，不能宽泛。

【国学气质增强】（必须体现）
- 每篇至少2句宏观意象句：大千世界、芸芸众生、天地节律、人间万象、四时流转。
- 每篇至少1组“长句铺陈 + 短句收束”的节奏表达。
- 语气方向：守拙、知止、顺势、积善、厚德；强调先稳心再起势。
- 可有轻玄学意境，但必须落地到现实行动，不能神化物件。
- 禁止神异断言和绝对结果承诺。

【低敏替换规则】（命中时必须替换）
- 天命/改命 -> 人生节律/转折阶段
- 贵人降临 -> 关键支持者/重要机会
- 法器/开运 -> 寓意物件/行动锚点
- 驱邪挡灾 -> 减少干扰/稳住状态
- 神明加持 -> 经验方法/心理支持
- 必发财/必翻盘 -> 提高把握/增加概率/更稳推进

【标题要求】
每篇5个标题，3篇共15个标题，必须各不相同：
- 标题要口语化，像朋友说话
- 字数控制在10-20字
- 禁止使用"震惊""竟然""99%的人不知道"等低质词汇

【结尾带货要求】
- 自然过渡到黄铜貔貅产品
- 强调"善待自己"的理念
- 3篇结尾的表达方式必须不同：
  - 第1篇：自我补偿型 - "亏欠自己太久了，是时候补上了"
  - 第2篇：重新开始型 - "新的开始，从善待自己这一刻起"
  - 第3篇：岁月沉淀型 - "到了这个年纪，更懂得什么值得拥有"
- 引导点击左下角小黄车或头像进主页橱窗

【可引用的经典（每篇引用不同）】
第1篇引用：《道德经》或王阳明
第2篇引用：《菜根谭》或曾国藩
第3篇引用：《庄子》或杨绛

【禁止内容】
- 不得出现：命运、运势、运气、福报、因果报应、天命、气场、能量场等玄学词汇
- 不得出现：神仙、菩萨、算命、占卜、风水、阴阳、五行等迷信词汇
- 不得出现：改运、转运、暴富、发横财等承诺类词汇
- 不得出现死亡、疾病、丧事等消极内容

【字数要求】
每篇1200-1500字

【输出格式】
【第一篇】

【标题1】xxx
【标题2】xxx
【标题3】xxx
【标题4】xxx
【标题5】xxx

---

（正文内容）

═══════════════════════════════════════════════════

【第二篇】

【标题1】xxx
【标题2】xxx
【标题3】xxx
【标题4】xxx
【标题5】xxx

---

（正文内容）

═══════════════════════════════════════════════════

【第三篇】

【标题1】xxx
【标题2】xxx
【标题3】xxx
【标题4】xxx
【标题5】xxx

---

（正文内容）

【再次强调】直接输出文案内容，从"【第一篇】"开始，不要有任何开场白或解释。"""

print("正在调用AI生成文案...")

response = client.messages.create(
    model="claude-sonnet-4-20250514",
    max_tokens=16000,
    messages=[{"role": "user", "content": prompt}]
)

generated_content = response.content[0].text
print(f"AI返回内容长度：{len(generated_content)}字")

# 解析文章
def parse_articles(content):
    articles = []
    pattern = r'【第[一二三1-3]篇】'
    parts = re.split(pattern, content)

    for idx, part in enumerate(parts):
        part = part.strip()
        if len(part) < 500:
            continue

        # 提取标题
        titles = []
        title_matches = re.findall(r'【标题[1-5]】\s*(.+?)(?=\n|$)', part)
        if title_matches:
            titles = [t.strip() for t in title_matches]

        # 提取正文
        body = part
        body = re.sub(r'【标题[1-5]】.+?\n', '', body)
        body = re.sub(r'^-{3,}$', '', body, flags=re.MULTILINE)
        body = re.sub(r'^═+$', '', body, flags=re.MULTILINE)
        body = body.strip()
        body = re.sub(r'\n{3,}', '\n\n', body)

        if len(body) >= 800:
            articles.append({
                'titles': titles if titles else [f"文案{len(articles)+1}"],
                'content': body
            })
            print(f"解析第{len(articles)}篇，{len(body)}字，{len(titles)}个标题")

    return articles

articles = parse_articles(generated_content)
print(f"共解析出 {len(articles)} 篇文章")

# 保存文档
if articles:
    timestamp = datetime.now().strftime('%Y%m%d%H%M%S')
    filename = f"{timestamp}_1.docx"
    output_path = r"D:\A百家号带货视频\带货文案"
    filepath = os.path.join(output_path, filename)

    doc = Document()
    doc.add_paragraph(f"引流类型：带货引流（黄铜貔貅）")
    doc.add_paragraph(f"生成时间：{datetime.now().strftime('%Y-%m-%d %H:%M:%S')}")
    doc.add_paragraph("═" * 50)
    doc.add_paragraph()

    for i, article in enumerate(articles, 1):
        doc.add_paragraph(f"【第{i}篇】", style='Heading 1')
        doc.add_paragraph()

        for j, title in enumerate(article['titles'][:5], 1):
            doc.add_paragraph(f"【标题{j}】{title}")

        doc.add_paragraph()
        doc.add_paragraph("---")
        doc.add_paragraph()
        doc.add_paragraph(article['content'])

        word_count = len(article['content'])
        doc.add_paragraph()
        doc.add_paragraph(f"（本篇字数：{word_count}字）")
        doc.add_paragraph()
        doc.add_paragraph("═" * 50)
        doc.add_paragraph()

    doc.save(filepath)
    print(f"\n文档已保存：{filepath}")
    print(f"共生成 {len(articles)} 篇文章")
else:
    print("未能解析出有效文章，原始内容：")
    print(generated_content[:2000])
