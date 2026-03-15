# -*- coding: utf-8 -*-
import anthropic
from docx import Document
from datetime import datetime
import os
import re

client = anthropic.Anthropic(
    base_url="https://yunyi.rdzhvip.com/claude",
    api_key="A68YGT53-41HW-8GT7-353U-MJVDNY8NE5KX"
)

prompt = """你是专业情感文案写手。现在直接生成3篇带货文案，每篇1200-1500字。

【商品】黄铜貔貅摆件（一对）
- 公貔貅主招财，脚踩绣球；母貔貅主守财，脚踩如意
- 成对摆放寓意"招财守财，财富双收"
- 黄铜铸造，真材实料，纳米镀层防水防污
- 适合客厅、办公桌、收银台摆放

【写作要求】
1. 第二人称"你"贯穿全文，像朋友聊天
2. 3篇开头必须完全不同
3. 每篇5个标题，共15个标题不能重复

4. 必须先提炼并融入以下爆点（不可省略）：
- 开场强代入但不固定：3篇必须使用不同开场样式（直呼点名/好友对话/场景切入/反问切入/共情陈述），禁止固定“孩子/有缘人/天选之人/道友”
- 共情安抚：先承认用户操劳、委屈、硬扛，再给情绪缓冲
- 国学感论证：使用“节律、时机、守拙、积善、厚德”等低敏表达
- 窗口期提醒：强调“现在要稳住节奏并行动”，不得绝对化承诺
- 物件定位：写成“行动锚点/仪式感提醒”，不要写成神奇改命
- 收束引导：结尾引流段控制80-140字；置顶必须写“看第几条+解决什么”；橱窗必须写“先看哪类+怎么选”，避免宽泛表述

5. 必须体现道家语感与宏观意象：
- 每篇至少2句“宏观意象句”（大千世界/芸芸众生/天地节律/人间万象/四时流转）
- 每篇至少1组“长句+短句断拍”节奏
- 用“守拙、知止、顺势、积善、厚德”这类低敏国学词，不要神化
- 可有轻玄学氛围，但必须落到现实行动与情绪安顿

5. 命中高风险词时，必须低敏替换：
- 天命/改命 -> 人生节律/转折阶段
- 贵人降临 -> 关键支持者/重要机会
- 法器/开运 -> 寓意物件/行动锚点
- 驱邪挡灾 -> 减少干扰/稳住状态
- 神明加持 -> 经验方法/心理支持
- 必发财/必翻盘 -> 提高把握/增加概率/更稳推进
4. 结尾自然引导到貔貅产品，点击小黄车购买
5. 每篇引用不同经典（道德经/菜根谭/庄子等）

【3篇主题】
第1篇：从"被忽视的付出"切入，肯定读者默默承受的一切
第2篇：从"中年觉醒"切入，人到中年才懂善待自己
第3篇：从"独处智慧"切入，独处是内心富足的表现

【禁止词汇】命运、运势、福报、因果、天命、气场、能量场、风水、阴阳、五行、改运、转运、暴富

【输出格式】严格按以下格式，直接开始：

【第一篇】

【标题1】xxx
【标题2】xxx
【标题3】xxx
【标题4】xxx
【标题5】xxx

---

正文内容（1200-1500字）

═══════════════════════════════════════════════════

【第二篇】

【标题1】xxx
【标题2】xxx
【标题3】xxx
【标题4】xxx
【标题5】xxx

---

正文内容（1200-1500字）

═══════════════════════════════════════════════════

【第三篇】

【标题1】xxx
【标题2】xxx
【标题3】xxx
【标题4】xxx
【标题5】xxx

---

正文内容（1200-1500字）

现在直接从【第一篇】开始输出："""

print("正在生成文案...")

response = client.messages.create(
    model="claude-opus-4-5-20251101",
    max_tokens=16000,
    messages=[{"role": "user", "content": prompt}]
)

content = response.content[0].text
print(f"生成内容：{len(content)}字")
print("="*50)
print(content[:2000])
print("="*50)

# 解析
articles = []
parts = re.split(r'【第[一二三]篇】', content)
for part in parts:
    part = part.strip()
    if len(part) < 500:
        continue
    titles = re.findall(r'【标题[1-5]】\s*(.+?)(?=\n|$)', part)
    body = re.sub(r'【标题[1-5]】.+?\n', '', part)
    body = re.sub(r'^-{3,}$', '', body, flags=re.MULTILINE)
    body = re.sub(r'^═+$', '', body, flags=re.MULTILINE)
    body = body.strip()
    if len(body) >= 500:
        articles.append({'titles': titles or [f"文案{len(articles)+1}"], 'content': body})
        print(f"第{len(articles)}篇：{len(body)}字，{len(titles)}个标题")

# 保存
if articles:
    os.makedirs(r"D:\A百家号带货视频\带货文案", exist_ok=True)
    ts = datetime.now().strftime('%Y%m%d%H%M%S')
    fp = rf"D:\A百家号带货视频\带货文案\{ts}_1.docx"
    doc = Document()
    doc.add_paragraph(f"生成时间：{datetime.now().strftime('%Y-%m-%d %H:%M:%S')}")
    doc.add_paragraph("═" * 40)
    for i, a in enumerate(articles, 1):
        doc.add_paragraph(f"【第{i}篇】", style='Heading 1')
        for j, t in enumerate(a['titles'][:5], 1):
            doc.add_paragraph(f"【标题{j}】{t}")
        doc.add_paragraph("---")
        doc.add_paragraph(a['content'])
        doc.add_paragraph(f"（字数：{len(a['content'])}）")
        doc.add_paragraph("═" * 40)
    doc.save(fp)
    print(f"\n已保存：{fp}")
else:
    print("解析失败，原文：")
    print(content[:3000])
