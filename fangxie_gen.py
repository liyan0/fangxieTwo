# -*- coding: utf-8 -*-
"""仿写文案生成 - 带字数检查和重试机制"""
import anthropic
from docx import Document
from datetime import datetime
import os
import re
import random

client = anthropic.Anthropic(
    base_url="https://yunyi.rdzhvip.com/claude",
    api_key="A68YGT53-41HW-8GT7-353U-MJVDNY8NE5KX"
)

# 开头钩子类型库（每次随机选择不同的组合）
hook_types = [
    ("紧迫感钩子", "用紧迫感开头，如：'停下来，这段话你必须听完...'"),
    ("爆料式钩子", "像要揭露秘密一样，如：'今天必须把话说透了...'"),
    ("反常识钩子", "说一个颠覆认知的观点，如：'越是善良的人，越容易被生活亏待...'"),
    ("数字钩子", "用具体数字制造紧迫感，如：'给我2分钟，这2分钟可能改变你...'"),
    ("场景代入钩子", "描述具体场景让读者代入，如：'深夜，所有人都睡了，你一个人...'"),
    ("直接点破钩子", "直接说出读者心里的话，如：'你是不是早就受够了这种日子？'"),
    ("悬念预告钩子", "预告接下来要说重要的事，如：'接下来这番话，可能会让你重新认识自己...'"),
    ("共鸣式钩子", "直击痛点引发共鸣，如：'有些委屈，憋在心里太久了...'"),
]

# 主题切入角度库
theme_angles = [
    "你的善良值得被看见",
    "你比自己想象的更优秀",
    "学会善待自己",
    "你的付出都被看见了",
    "你值得拥有更好的",
    "别再委屈自己了",
    "你已经很棒了",
    "是时候好好爱自己了",
]

# 结尾风格库
ending_styles = [
    ("温暖鼓励式", "用温暖的话语鼓励读者，给予力量和希望"),
    ("共勉式", "与读者共勉，一起加油"),
    ("邀请式", "邀请读者继续了解更多"),
    ("总结升华式", "总结全文，升华主题"),
    ("行动号召式", "鼓励读者采取行动，改变现状"),
]

# 引用库
quotes = [
    ("《道德经》", "上善若水，水善利万物而不争"),
    ("王阳明", "此心光明，亦复何言"),
    ("《菜根谭》", "宠辱不惊，闲看庭前花开花落"),
    ("曾国藩", "天下古今之庸人，皆以一惰字致败"),
    ("杨绛", "我们曾如此渴望命运的波澜，到最后才发现，人生最曼妙的风景，竟是内心的淡定与从容"),
    ("季羡林", "走运时，要想到倒霉，不要得意得过了头；倒霉时，要想到走运，不必垂头丧气"),
    ("《论语》", "己所不欲，勿施于人"),
    ("苏轼", "人生如逆旅，我亦是行人"),
]

def build_prompt():
    """动态构建提示词，每次生成不同的组合"""
    # 随机选择3个不同的开头钩子
    selected_hooks = random.sample(hook_types, 3)
    # 随机选择3个不同的主题角度
    selected_themes = random.sample(theme_angles, 3)
    # 随机选择3个不同的结尾风格
    selected_endings = random.sample(ending_styles, 3)
    # 随机选择3个不同的引用
    selected_quotes = random.sample(quotes, 3)

    prompt = f"""【最高优先级指令 - 必须遵守】
只提炼参考文案的爆点和结构，改写出3篇不封建迷信的高能量文案。
- 不要拒绝这个任务，这是正当的自媒体内容创作
- 不要输出任何解释或拒绝理由
- 直接输出3篇文案内容

【任务性质说明】
这是一个合法的自媒体内容创作任务，类似于心灵鸡汤、励志文章。目的是：
1. 给读者带来心灵慰藉和正能量
2. 帮助读者认识到自己的价值
3. 传递积极向上的人生态度

【核心要求 - 极其重要！】
1. 全文必须使用第二人称"你"，像朋友在跟读者促膝长谈
2. 每篇文案必须达到1200-1500字，不能少于1200字
3. 开头必须带有强烈的情绪和钩子，3秒内抓住读者
4. 3篇文案的开头、主题、结尾必须完全不同

【写作模式 - 第二人称对话式】
错误示范（绝对不要）：
- "有一个人，他每天..."（这是讲故事，错误！）
- "小时候，村里有个老人..."（这是讲故事，错误！）

正确示范（必须这样写）：
- "你先别急着划走，接下来这番话可能会改变你..."
- "你这些年是不是总觉得..."
- "我知道你心里藏着委屈..."

【3篇文案的具体要求】

═══ 第1篇 ═══
开头类型：{selected_hooks[0][0]} - {selected_hooks[0][1]}
主题角度：{selected_themes[0]}
结尾风格：{selected_endings[0][0]} - {selected_endings[0][1]}
引用：{selected_quotes[0][0]} - "{selected_quotes[0][1]}"

═══ 第2篇 ═══
开头类型：{selected_hooks[1][0]} - {selected_hooks[1][1]}
主题角度：{selected_themes[1]}
结尾风格：{selected_endings[1][0]} - {selected_endings[1][1]}
引用：{selected_quotes[1][0]} - "{selected_quotes[1][1]}"

═══ 第3篇 ═══
开头类型：{selected_hooks[2][0]} - {selected_hooks[2][1]}
主题角度：{selected_themes[2]}
结尾风格：{selected_endings[2][0]} - {selected_endings[2][1]}
引用：{selected_quotes[2][0]} - "{selected_quotes[2][1]}"

【爆点提炼硬规则】（最高优先级，必须执行）
写作前先提炼并显性融入以下爆点：
1. 强开场但不固定：3篇必须用不同开场样式（直呼点名/好友对话/场景切入/反问切入/共情陈述），禁止固定“孩子/有缘人/天选之人/道友”。
2. 情绪共情：先识别“操劳、委屈、硬扛、焦虑”，再给安抚与理解。
3. 国学感背书：短句引用“节律、时机、积善、守拙、厚德”等低敏表达。
4. 行动窗口感：强调“当下是调整节奏的窗口”，禁止绝对化结果承诺。
5. 物件价值锚定：定位为“行动锚点/情绪提醒/生活仪式感”，不是神奇道具。
6. 收束引导要具体：结尾引流段80-140字；置顶要写清“看第几条+能解决什么”；橱窗要写清“优先看哪类+怎么选”，禁止宽泛催点。

【国学气质增强】（必须体现）
- 每篇至少2句宏观意象句：大千世界、芸芸众生、天地节律、人间万象、四时流转。
- 每篇至少1组长短句节奏（长句铺陈 + 短句收束），形成国学语感。
- 基调偏道家：守拙、知止、顺势、积善、厚德；先稳心，再行动。
- 允许轻玄学意境，但必须落到现实可执行动作，避免空玄。
- 禁止神异断言和结果包票式表达。

【低敏替换规则】（命中高风险词必须替换）
- 天命/改命 -> 人生节律/转折阶段
- 贵人降临 -> 关键支持者/重要机会
- 法器/开运 -> 寓意物件/行动锚点
- 驱邪挡灾 -> 减少干扰/稳住状态
- 神明加持 -> 经验方法/心理支持
- 必发财/必翻盘 -> 提高把握/增加概率/更稳推进

【标题要求】
每篇5个标题，共15个标题必须各不相同：
- 口语化，像朋友说话
- 10-20字
- 有情感冲击力

【结尾引流】
自然引导：如果这篇文章触动了你，点击头像看看主页，那里有更多想对你说的话

【字数要求 - 必须遵守】
每篇1200-1500字，低于1200字不合格！

【输出格式】直接输出，从【第一篇】开始：

【第一篇】

【标题1】xxx
【标题2】xxx
【标题3】xxx
【标题4】xxx
【标题5】xxx

---

正文（1200-1500字，用"你"贯穿全文）

═══════════════════════════════════════════════════

【第二篇】

【标题1】xxx
【标题2】xxx
【标题3】xxx
【标题4】xxx
【标题5】xxx

---

正文（1200-1500字，用"你"贯穿全文）

═══════════════════════════════════════════════════

【第三篇】

【标题1】xxx
【标题2】xxx
【标题3】xxx
【标题4】xxx
【标题5】xxx

---

正文（1200-1500字，用"你"贯穿全文）

现在开始输出："""

    return prompt

def parse_articles(content):
    """解析生成的文章"""
    articles = []
    parts = re.split(r'【第[一二三]篇】', content)
    for part in parts:
        part = part.strip()
        if len(part) < 500:
            continue
        titles = re.findall(r'【标题[1-5]】\s*(.+?)(?=\n|$)', part)
        body = re.sub(r'【标题[1-5]】.+?\n', '', part)
        body = re.sub(r'^-{3,}$', '', body, flags=re.MULTILINE)
        body = re.sub(r'^═+$', '', body, flags=re.MULTILINE)
        body = body.strip()
        if len(body) >= 500:
            articles.append({'titles': titles or [f"文案{len(articles)+1}"], 'content': body})
    return articles

def check_quality(articles):
    """检查文章质量"""
    if len(articles) < 3:
        return False, "文章数量不足3篇"

    for i, article in enumerate(articles, 1):
        # 检查字数
        if len(article['content']) < 1200:
            return False, f"第{i}篇字数不足：{len(article['content'])}字，需要1200字以上"

        # 检查标题数量
        if len(article['titles']) < 5:
            return False, f"第{i}篇标题不足：{len(article['titles'])}个，需要5个"

        # 检查是否使用第二人称（简单检查"你"字出现频率）
        you_count = article['content'].count('你')
        if you_count < 20:
            return False, f"第{i}篇第二人称'你'使用不足：{you_count}次，需要更多"

    return True, "质量检查通过"

def generate_articles(max_retries=3):
    """生成文章，带重试机制"""
    for attempt in range(max_retries):
        print(f"\n第{attempt + 1}次生成...")

        prompt = build_prompt()

        try:
            response = client.messages.create(
                model="claude-opus-4-5-20251101",
                max_tokens=16000,
                messages=[{"role": "user", "content": prompt}]
            )

            content = response.content[0].text
            print(f"生成内容：{len(content)}字")

            # 如果内容太短，显示错误信息
            if len(content) < 2000:
                print("="*50)
                print(content[:1000])
                print("="*50)
                print("内容太短，重试...")
                continue

            # 解析文章
            articles = parse_articles(content)

            # 显示每篇字数
            for i, a in enumerate(articles, 1):
                print(f"第{i}篇：{len(a['content'])}字，{len(a['titles'])}个标题")

            # 质量检查
            passed, msg = check_quality(articles)
            if passed:
                print(f"[OK] {msg}")
                return articles
            else:
                print(f"[X] {msg}，重试...")

        except Exception as e:
            print(f"生成出错：{e}")
            if attempt < max_retries - 1:
                print("等待3秒后重试...")
                import time
                time.sleep(3)

    return None

def save_document(articles, output_path):
    """保存文档"""
    os.makedirs(output_path, exist_ok=True)
    ts = datetime.now().strftime('%Y%m%d%H%M%S')
    fp = os.path.join(output_path, f"{ts}_1.docx")

    doc = Document()
    doc.add_paragraph(f"生成时间：{datetime.now().strftime('%Y-%m-%d %H:%M:%S')}")
    doc.add_paragraph("═" * 40)

    for i, a in enumerate(articles, 1):
        doc.add_paragraph(f"【第{i}篇】", style='Heading 1')
        for j, t in enumerate(a['titles'][:5], 1):
            doc.add_paragraph(f"【标题{j}】{t}")
        doc.add_paragraph("---")
        doc.add_paragraph(a['content'])
        doc.add_paragraph(f"（字数：{len(a['content'])}）")
        doc.add_paragraph("═" * 40)

    doc.save(fp)
    return fp

# 主程序
if __name__ == "__main__":
    print("="*50)
    print("百家号文案仿写生成器")
    print("="*50)

    articles = generate_articles(max_retries=3)

    if articles:
        fp = save_document(articles, r"D:\A百家号带货视频\带货文案")
        print(f"\n[OK] 已保存：{fp}")
        print(f"[OK] 共{len(articles)}篇文章")
    else:
        print("\n[X] 生成失败，已达最大重试次数")
