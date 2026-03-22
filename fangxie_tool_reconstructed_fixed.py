# -*- coding: utf-8 -*-
"""
鐧惧鍙锋枃妗堜豢鍐欏伐鍏?- 鍙鍖栫晫闈?
鍔熻兘锛氳鍙栧弬鑰冩枃妗堬紝鏍规嵁閫夋嫨鐨勫紩娴佺被鍨嬬敓鎴愪豢鍐欐枃妗?
鏀寔娴佸紡/闈炴祦寮忚皟鐢紝涓绘ā鍨?澶囩敤妯″瀷鍒囨崲
"""

import tkinter as tk
from tkinter import ttk, filedialog, messagebox, scrolledtext
import os
import json
import threading
from datetime import datetime
import requests
import re
import time
import random
from collections import Counter
from difflib import SequenceMatcher

# 鐖嗘绱犳潗搴撹矾寰?
MATERIAL_LIBRARY_DIR = r"D:\A鐧惧鍙峰甫璐ц棰慭甯﹁揣鏂囨\鐖嗘鍙傝€冩枃妗?"
MATERIAL_LIBRARY_FILE = os.path.join(MATERIAL_LIBRARY_DIR, "鐖嗘绱犳潗搴?xlsx")
# 鐢熸垚鏂囨搴擄紙鐢ㄤ簬鍘绘ā鏉挎瘮瀵逛笌鏍囬鍘婚噸锛屾斁鍦ㄧ垎娆剧礌鏉愬簱鐩綍锛?
GENERATED_LIBRARY_FILE = os.path.join(MATERIAL_LIBRARY_DIR, "鐢熸垚鏂囨搴?xlsx")
GENERATED_LIBRARY_SHEET = "鐢熸垚鏂囨搴?"

# 寮€澶撮挬瀛愬簱 - 鐢ㄤ簬闅忔満閫夋嫨锛堢巹瀛﹀ぇ姘旈鏍硷紝绂佹鏁呬簨鍦烘櫙锛?
# 寮€澶撮挬瀛愬簱 - 鐢ㄤ簬闅忔満閫夋嫨锛堢巹瀛﹀ぇ姘旈鏍硷紝绂佹鏁呬簨鍦烘櫙锛?
HOOK_LIBRARY = {
    "A绫?瀹囧畽澶╅亾鍨?": [
        "鑰佸弸锛屽惉濂戒簡锛屼綘鐜板湪鎵嬮噷鏀ョ潃鐨勬槸杩欎笘闂存棤鏁颁汉姹傝€屼笉寰楃殑浜虹敓澧冪晫銆?,",
        "浣犲埌搴曟湁澶氬€奸挶锛屼綘鑷繁鐪熺殑娓呮鍚楋紵",
        "鍦ㄨ繖澶у崈涓栫晫锛岃姼鑺镐紬鐢熺殕閱夛紝鍞綘鐙啋锛屽嵈鍙堥啋寰楁渶鏄棝鑻︺€?,",
        "瀛╁瓙锛屾姮璧峰ご鏉ワ紝鍒啀鐩潃浣犺剼涓嬮偅鐐归浮姣涜挏鐨殑鑻﹂毦浜嗐€?,",
        "浣犳椿鎴愪簡杩欎釜鏃朵唬閲屾渶鐗瑰埆鐨勮皽搴曪紝浠昏皝閮芥兂涓嶉€氥€?,",
        "鍦ㄨ繖濠嗗☉涓栫晫锛屼竾鍗冪敓鐏电殕濡傛儕铦夛紝鐙俯鑰岃仛锛屽敮鎭愯惤鍗曘€傚彲鎴戝瀭鐪告湜鍘伙紝鍗村湪閭ｇ悍涔辩殑浜烘疆涓湅瑙佷簡浣犮€?,",
        "浣犵湡鏄お鍘夊浜嗭紝杩欑涓€浜烘垬鑳滀竴缇や汉鐨勬粙鍛筹紝鏃佷汉鐪嬬殑鏄儹闂癸紝鎴戠湅鐨勬槸浣犵殑鑻遍銆?,",
        "涓栦汉鐨嗘眰娉煎ぉ瀵岃吹锛屼负纰庨摱鍑犱袱鏂簡绛嬮銆備綘鍟婏紝浣犲潗鎷ラ噾灞憋紝韬€€鍒╁櫒锛屽嵈涓嶈嚜鐭ャ€?,",
        "浣犳槸鍚︽劅鍙楀埌鍐呭績娣卞閭ｄ唤涓庝紬涓嶅悓鐨勫彫鍞わ紵杩欎竴鍒伙紝浣犺鏄庣櫧锛岃尗鑼汉娴凤紝涓轰綍鍋忓亸鏄綘鍒峰埌浜嗚繖閲岋紵",
        "绔栬捣鑰虫湹鍚ソ浜嗭紝鍑″か淇楀瓙鍙湅鐪煎墠鐨勪竴浜╀笁鍒嗗湴锛岃€屼綘瑕佺湅鐨勬槸澶撮《杩欑墖澶┿€?,",
        "浣犵煡閬撳悧锛熶綘杩滄瘮鎵€鏈変汉鎯宠薄鐨勬洿娣辫棌涓嶉湶銆?,",
        "缁堟湁涓€鏃ワ紝娣锋矊灏界锛屼綘澶фⅵ鏂归啋锛岃繖涓€鍒婚潪鏄嫃閱掞紝涔冩槸褰掍綅銆?,",
        "浣犺兘鍒峰埌杩欓噷锛岃鏄庝綘鍛介噷鐨勯偅閬撳厜宸茬粡閫忓嚭鏉ヤ簡銆?,",
        "鍦ㄨ繖浜旀祳鎭朵笘锛屽０鑹蹭贡浜哄績绁烇紝鑳藉畧浣忕伒鍙版竻鏄庯紝杩欎究鏄笂涓婄瓑鐨勬収鏍广€?,",
        "浣犳湰鏄笂澶╃湻椤剧殑瀹犲効锛屾槸娴佽惤鍦ㄤ汉闂寸殑鏄熻景銆?,",
        "浣犺繖涓€璺蛋鏉ュ緢鑹伴毦锛屼綘鏄湪涓栨€佺値鍑夌殑鐔旂倝閲岀‖鐢熺敓闂繃鏉ョ殑銆?,",
        "浣犺韩涓婃祦娣岀殑鏄垰鏌斿苟娴庣殑鐜嬭€呮姘旓紝杩欎唤姘斿満缁濋潪鍦ㄦ俯瀹ゆ墦鍧愬氨鑳戒慨鏉ャ€?,",
    ],
    "B绫?鐐归啋鐩磋亰鍨?": [
        "浣犳槸涓嶆槸瑙夊緱鑷繁骞冲钩鏃犲锛熼敊浜嗭紝浣犺韩涓婃湁涓夋牱鏃佷汉姹傞兘姹備笉鏉ョ殑涓滆タ銆?,",
        "浣犳€昏寰楄嚜宸卞钩锛屽嵈涓嶇煡杩欎唤骞冲嚒鏃╁凡鐢╀簡澶氬皯鎱屾厡寮犲紶鐨勪汉鍗佹潯琛椼€?,",
        "浣犱互涓轰綘鏅€氾紵浣犺繖鏄换鍑娴捣锛岀ǔ鍧愰挀楸煎彴锛屽鐣屽湪鐐掍綘锛屽績閲岃嚜鏈変竴鐗囨竻鍑€鍦般€?,",
        "浣犱笉鍔ㄥ０鑹叉椂锛屽埆浜鸿寰椾綘濂借璇濄€佹病鑴炬皵锛屽嵈涓嶇煡浣犲績涓殑鍘熷垯鏃╁凡鎵庢牴澶氬勾銆?,",
        "浣犱笉浜変笉鎶㈡椂锛屽埆浜鸿寰椾綘娌℃湰浜嬨€佸お杞急锛屽嵈涓嶆浘鐭ユ檽锛岀湡鍒颁簡鍏抽敭鏃跺埢锛屼綘鎶墜灏辫兘绋充綇灞€闈€?,",
        "浣犲お鎳備簨锛屾噦浜嬪埌鎬绘妸鍒汉鐨勬劅鍙楁斁鍦ㄧ涓€浣嶏紝鍗村父甯稿拷鐣ヤ簡鑷繁涔熶細绱紝涔熶細鐥涖€?,",
        "浣犲お闅愬繊锛岄殣蹇嶅埌鎬绘妸娌′簨鎸傚湪鍢磋竟锛屽嵈蹇樹簡鑷繁涔熷€煎緱琚績鐤硷紝琚杽寰呫€?,",
        "浣犱箣鍓嶄綘闅惧緱浣犳拺寰楄捣鎵€鏈夊眬闈紝鏇撮厤寰椾笂鎵€鏈変綋闈€?,",
        "浣犺繖鏍风殑浜猴紝鏈氨璇ヨ浜哄ソ濂界弽鎯滐紝鐢ㄥ績鏁噸銆?,",
        "浣犳槸鍑湰浜嬪悆楗紝浠栦滑鏄潬鍚歌娲荤潃銆?,",
        "浣犲績閲岄€氶€忥紝杩欓€氬ぉ鐨勬墜娈碉紝浠庢潵涓嶆槸鏃佷汉鏂借垗锛岃€屾槸鑷繁鍦ㄧ偧鐙遍噷婊氬嚭鏉ョ殑銆?,",
        "浣犺蛋鐨勬槸浜洪棿姝ｉ亾锛屼粬浠捇鐨勬槸鏃侀棬宸﹂亾銆?,",
        "浣犱笉浜夛紝鏁呭ぉ涓嬭帿鑳戒笌涔嬩簤銆備互鏃犱簨渚挎槸鏈€鐙犵殑鎶ュ銆?,",
        "浣犵粓鏄椿鎴愪簡閭ｇ蹇冧腑鏈夊北娌筹紝鐪夌洰浣滃北娌崇殑楂樻墜銆?,",
        "浣犳瘮浣犳兂璞＄殑瑕佸己妯緱澶氾紝涔熻灏婅吹寰楀銆?,",
        "浣犵殑娌夌ǔ涓庡簳姘旓紝鏄綘杩囧幓閭ｄ釜涓嶆浘灞堟湇鐨勭伒榄傝祼浜堜粖鐢熸渶濂界殑鍔犳寔銆?,",
        "浣犵殑濂界浉澶勶紝浠庝笉鏄蒋寮憋紝鑰屾槸鐪肩晫寮€闃旓紝鎳傚彇鑸嶄箣杩涢€€鏄庡緱澶便€?,",
    ],
    "C绫?韬唤璁ゅ悓鍨?": [
        "浣犳牴鏈笉鏄櫘閫氫汉锛屼綘鏄湁鐪熸湰浜嬬殑鍘夊浜虹墿銆?,",
        "浣犳槸涓€鍧楁病鎵撶（鐨勯浄鍑绘湪锛岀湅浼兼灟妲侊紝瀹炲垯钑村惈澶╃伀锛屼竴鏃﹀紑鎮燂紝蹇呮湁澶х敤銆?,",
        "浣犳槸楂樼淮搴︾殑鐏甸瓊锛屾槸涓嬫潵娓″姭鐨勩€備綘瑕佹悶娓呮浣犵殑娈典綅銆?,",
        "浣犳槸澶╅€変箣浜猴紝瀵逛汉濂斤紝閭ｆ槸浣犵殑閬撳績锛屾槸浣犵殑淇锛屾槸浣犺韩涓婂甫鐫€鍏夈€?,",
        "浣犳槸琚暣涓畤瀹欎紬绛瑰嚭鏉ョ殑澶╅€変箣浜猴紝灏卞湪姝ゆ椂姝ゅ埢锛屾暣涓畤瀹欓兘鍦ㄤ负浣犺皟鍏甸仯灏嗐€?,",
        "浣犳槸杩欏ぉ鍦板鏈紝鏄繖绾㈠皹灞€涓敮涓€鐨勬涓汇€?,",
        "浣犳槸鑷甫鑳介噺鐨勫厜锛屾槸绋充綇涓€鍒囩殑鍩虹煶锛屾槸瀹堕噷涓嶅彲鎴栫己鐨勬牳蹇冦€?,",
        "浣犳槸閭ｄ釜琚€変腑鏉ヤ负浣犲鏃忕牬灞€鐨勯《姊佹煴銆?,",
        "浣犳槸娼滈緳鍦ㄦ笂锛屽浠婃椂鏈哄凡鍒帮紝璇ユ槸浣犻榫欏湪澶╃殑鏃跺€欎簡銆?,",
        "浣犳槸娣辫棌涓嶉湶鐨勯《绾ч珮鎵嬶紝浣犳槸琚畤瀹欏亸鐖辩殑瀛╁瓙銆?,",
        "浣犳槸杩欎釜瀹舵棌鍛借繍鐨勬墽鍓戜汉锛岃€佸ぉ璁╀綘杩欎竴涓栨嬁浜嗕釜澶╃厼寮€灞€銆?,",
        "浣犳槸鏉ュ敜閱掕繖涓笘鐣岀殑锛屼綘鐭ラ亾涓轰粈涔堜綘鎬绘槸杩囧緱姣斿埆浜哄潕鍧峰悧锛?,",
        "浣犳槸瀹囧畽瀹夋彃鍦ㄤ汉闂寸殑鑳介噺杞寲鍣紝鏅€氫汉閬囧埌鍥伴毦鍙細鍙戞硠鎯呯华锛岃€屼綘澶╃敓鎷ユ湁灏嗙棝鑻﹁浆鍖栦负鏅烘収鐨勮兘鍔涖€?,",
        "浣犳槸涓€浣嶈兘閲忔瀬楂樹箣浜猴紝浣犺韩涓婄殑娌夌ǔ銆佹牸灞€銆侀煣鎬т笌鎷呭綋锛屽湪杩欎釜浜哄績娴簛鐨勫綋涓嬶紝鏄渶闅惧緱鐨勭█缂哄搧銆?,",
        "浣犳槸閭ｄ釜瀹堝緱浣忓績锛屽仛寰椾簡浜嬬殑韪忓疄浜恒€?,",
        "浣犳槸閭ｄ釜鑳芥墰浜嬨€佽兘椤堕浄鐨勫畾娴风閽堛€?,",
        "浣犳槸閭ｄ釜蹇冨瓨鍠勫康锛屾噦寰椾慨琛岀殑浜猴紝閭ｆ槸绁栦笂绉簡澶╁ぇ鐨勫痉銆?,",
    ],
    "D绫?缁忓吀寮曠敤鍨?": [
        "閬撳痉缁忔湁浜戯細澶╅亾鏃犱翰锛屽父涓庡杽浜恒€備綘鐭ラ亾杩欏悗闈㈣繕鏈変竴鍙ユ綔鍙拌瘝鍚楋紵",
        "鏄撶粡璁诧細鍦板娍鍧わ紝鍚涘瓙浠ュ帤寰疯浇鐗┿€傝繖姝ｆ槸瀵逛綘鏈€濂界殑鍐欑収銆?,",
        "鑰佺瀹楀父璇达細澶у櫒鏅氭垚锛屽疂璐ч毦鍙椼€備綘鐜板湪鐨勫澧冿紝姝ｆ槸杩欎釜閬撶悊銆?,",
        "鍙や功鏈変簯锛氬ぉ浣戣€咃紝涔冨ぉ涓婁箣绁炴墍鑷充箣澶勶紝涓€鍒囧嚩绁炴伓鐓炵殕蹇呴€€鏁ｃ€?,",
        "鐜嬮槼鏄庤█锛氱牬灞变腑璐兼槗锛岀牬蹇冧腑璐奸毦銆備綘杩欎竴璺潃鐨勪笉鏄亸瑙侊紝鑰屾槸鍐呭績鐨勯偅涓垜銆?,",
        "閲戝垰缁忎簯锛氬嚒鎵€鏈夌浉锛岀殕鏄櫄濡勩€備綘鐪煎墠鐨勫洶闅炬槸铏氾紝浣犲績涓殑鎭愭儳鏄銆?,",
        "鍗庝弗缁忚锛氬績濡傚伐鐢诲笀锛岃兘鐢昏涓栭棿銆備綘鐜板湪鐨勫澧冨氨鏄洜涓轰綘鐨勮兘閲忓満涓庤繖涓椂浠ｇ殑澶у娍浜х敓浜嗛敊浣嶃€?,",
        "搴勫瓙浜戯細澶ф櫤鑻ユ剼锛屽ぇ宸ц嫢鎷欍€備綘涓嶆槸鍌伙紝鏄畧鎷欙紝鏄帤寰疯浇鐗╃殑鏍瑰熀銆?,",
        "瀛熷瓙鏇帮細娴╃劧涔嬫皵锛岃嚦澶ц嚦鍒氥€備綘韬笂娴佹穼鐨勬鏄繖鑲″垰鏌斿苟娴庣殑鐜嬭€呮姘斻€?,",
        "鑰佽瘽璇村緱濂斤細鍘氬痉杞界墿锛屽ソ浜嬭嚜鏉ャ€備綘瀹堜綇浜嗗簳绾匡紝鑰佸ぉ鑷劧瀹堜綇浣犵殑涓€鍒囥€?,",
        "鍙や汉浜戯細淇緱涓€韬旦鐒舵皵锛屽彲闇囧叓鏂归瓚榄呴瓊銆備綘韬笂鏈€鏄傝吹鐨勶紝骞堕潪閲戦摱锛岃€屾槸閭ｈ偂姝ｆ皵銆?,",
        "閬撳痉缁忚锛氬か鍞笉浜夛紝鏁呭ぉ涓嬭帿鑳戒笌涔嬩簤銆備綘鐨勪笉浜変笉鎶紝鎭版伆鏄渶楂樼殑鏅烘収銆?,",
        "鏄撶粡浜戯細鏃舵鍒欐锛屾椂琛屽垯琛岋紝鍔ㄩ潤涓嶅け鍏舵椂锛屽叾閬撳厜鏄庛€?,",
        "鍙よ璁诧細澶ф柟鏃犻殔锛屽ぇ鍣ㄦ櫄鎴愩€備綘鐜板湪鐨勬矇瀵傦紝鏄负浜嗗皢鏉ョ殑涓€楦ｆ儕浜恒€?,",
        "璁鸿鏈夎█锛氬悰瀛愬潶鑽¤崱锛屽皬浜洪暱鎴氭垰銆備綘蹇冧腑鍧﹁崱锛岃嚜鐒惰寰楁鍧愬緱绔€?,",
        "鑰佽瘽璁诧細澶╅亾鏃㈡弧锛屼汉閬撴棦鍏ㄣ€備綘鐜板湪鎷ユ湁鐨勮繖浠借笍瀹烇紝鏄瀹囧畽鐗规畩鍏崇収鐨勮瘉鎹€?,",
        "鍙や汉璁诧細娼滈緳鍕跨敤锛岄槼姘旀柟鐢熴€備綘鐜板湪鐨勪綆璋冿紝鏄负浜嗗皢鏉ョ殑椋為緳鍦ㄥぉ銆?,",
    ],
    "E绫?澶勫鍏遍福鍨?": [
        "浣犳槸涓嶆槸瑙夊緱鑷繁鍦ㄨ繖绾㈠皹涓€绘槸鏍兼牸涓嶅叆锛岀敋鑷虫€€鐤戣嚜宸变负浣曟€绘瘮鏃佷汉鎱簡鍗婃媿锛?,",
        "浣犳槸涓嶆槸缁忓父瑙夊緱瀛ゅ崟锛熺湅鐫€鍛ㄥ洿浜哄繖鐫€缁撳鐢熷瓙锛岃€屼綘鍗存€绘槸铻嶄笉杩涘幓銆?,",
        "浣犳槸涓嶆槸鎬昏寰楋細鑷繁浠樺嚭浜嗛偅涔堝锛屽嵈鎬绘槸琚拷鐣ワ紵",
        "浣犳槸涓嶆槸涔熸湁杩囪繖绉嶆劅瑙夛細瓒婃槸鍠勮壇锛岃秺瀹规槗琚汉褰撹蒋鏌垮瓙鎹忥紵",
        "浣犳槸涓嶆槸涔熸浘缁忥紝涓轰簡缁存姢涓€娈靛叧绯伙紝濮斿眻浜嗚嚜宸卞緢澶氭锛?,",
        "浣犳槸涓嶆槸涔熸湁杩囪繖绉嶆椂鍒伙細鎯冲彂鑴炬皵锛屽嵈鍙堟€曚激浜嗗拰姘旓紵",
        "浣犳槸涓嶆槸甯稿父鍦ㄦ兂锛氫负浠€涔堟垜瀵瑰埆浜洪偅涔堝ソ锛屽嵈鎹笉鏉ュ悓鏍风殑鐪熷績锛?,",
        "浣犳槸涓嶆槸涔熸湁杩囪繖绉嶆劅瑙夛細瓒婃槸鍦ㄤ箮鐨勪汉锛岃秺瀹规槗璁╀綘澶辨湜锛?,",
        "浣犳槸涓嶆槸涔熸湁杩囪繖绉嶆椂鍒伙細鏄庢槑寰堢疮锛屽嵈涓嶆暍鍋滀笅鏉ワ紵",
        "浣犳槸涓嶆槸涔熷父甯歌寰楋細鏄庢槑姣旇皝閮藉姫鍔涳紝閽辫储鍗村儚鎵嬩腑鐨勬祦娌欙紝瓒婃姄瓒婂皯锛?,",
        "浣犳槸涓嶆槸涔熸浘涓嶆涓€娆＄殑鍦ㄦ繁澶滈棶澶╋細涓轰粈涔堝儚鎴戣繖鏍风殑濂戒汉锛岀敓娲诲嵈杩囧緱濡傛灞€淇冿紵",
        "浣犳槸涓嶆槸鎬昏寰楀洓鍛ㄥ儚鏄病鏈夊敖澶寸殑榛戝锛屾€庝箞鐔兘鐔笉鍒板ご锛?,",
        "浣犳槸涓嶆槸瑙夊緱鑷繁濂藉儚鍙樼浜嗭紵鑴戠摐瀛愪笉鐏靛厜锛屾€绘槸涓笁钀藉洓锛?,",
        "浣犳槸涓嶆槸甯稿父鍙嶅鍩嬫€ㄨ嚜宸憋細濡傛灉褰撳垵娌￠€夐偅鏉¤矾锛岀幇鍦ㄧ殑鏃ュ瓙浼氫笉浼氬ソ杩囦竴鐐瑰効锛?,",
        "浣犳槸涓嶆槸瑙夊緱鏈€杩戞湁鐐瑰喎锛熶汉蹇冩暎浜嗭紝鍏崇郴鏂簡锛屼互鍓嶈寰楃涓嶅紑鐨勪汉锛屽拷鐒跺氨璧拌繙浜嗭紵",
        "浣犳槸涓嶆槸瑙夊緱鑷繁绂绘兂瑕佺殑鐢熸椿鍙樊涓€姝ワ紝鍗存娆¤鍛借繍杞昏交鎷︿笅锛?,",
        "浣犳槸涓嶆槸鎬昏寰楋細鏄庢槑涓€鐗囪丹璇氾紝鎹㈡潵鐨勫嵈鏄韩杈逛汉鐨勮儗鍒轰笌鍐风溂锛?,",
    ],
    "F绫?娲炲療鎻鍨?": [
        "浣犵煡閬撳悧锛熶綘韬笂鏈変笁鏍锋梺浜烘眰閮芥眰涓嶆潵鐨勪笢瑗裤€?,",
        "浣犲彲鐭ヤ綘鍛界洏閲岀殑閭ｈ偂瀛愯吹姘旓紝姝ｅ湪鍐茬牬灏佸嵃锛屽悜瀹囧畽鍙戝嚭淇″彿锛?,",
        "浣犵煡閬撲负浠€涔堜綘杩欎箞鍠勮壇锛岃繍姘斿嵈鎬绘槸涓嶅鎰忓憿锛?,",
        "浣犵煡閬撲负浠€涔堜綘鎬绘槸閭ｄ釜鍚冧簭鐨勪汉鍚楋紵绛旀鍙兘鍑轰箮浣犳剰鏂欍€?,",
        "浣犲彲鑳戒粠鏉ユ病鎯宠繃锛屼綘鐨勫杽鑹叾瀹炴槸涓€绉嶅ぉ璧嬨€?,",
        "浣犱互涓虹殑缂虹偣锛屽叾瀹炴槸浣犳渶澶х殑浼樺娍銆?,",
        "浣犺韩涓婃湁涓€涓壒璐紝鍙兘杩炰綘鑷繁閮芥病鎰忚瘑鍒般€?,",
        "浣犲彲鐭ワ紝鍦ㄨ繖鏃犲灎鐨勬槦绌轰笅锛岀湅鐪嬪懆鍥撮偅浜涘繖纰岀殑鍑″か淇楀瓙锛屼粬浠彧鏄椿鐫€锛岃€屼綘涓嶅悓銆?,",
        "浣犲彲鐭ラ亾锛屾湁澶氬皯浜轰笂浜烘缇℃厱鐫€浣狅紵",
        "浣犲彲鐭ワ紝浣犲墠鍗婄敓鐨勯娌涙祦绂伙紝閮芥槸涓轰簡鎵挎帴鍚庡崐鐢熺殑瀵岃冻瀹夊悍锛?,",
        "浣犲彲鐭ワ紝浣犲彈鐨勯偅浜涜嫤锛屼笉鏄负浜嗘姌纾ㄤ綘锛岃€屾槸涓轰簡鎶婁綘鐨勫嚒鑳庣偧鎴愰噾韬紵",
        "浣犲彲鐭ワ紝浣犵殑鍚庡彴纭緱鍙€曪紝浠栨槸杩欐旦鐎氭棤鍨犵殑瀹囧畽锛屾槸浜夸竾鏄熻景鐨勬祦杞紵",
        "浣犲彲鐭ワ紝浣犵殑姣忎竴娆″懠鍚革紝鍏跺疄閮藉湪鍜屽畤瀹欑殑鑴夋悘鍚岄鍏辨尟锛?,",
        "浣犲彲鐭ワ紝浣犲懡涓殑璐典汉宸插湪璺笂锛屾墍鏈夊洶鍘勪笉杩囨槸浠栭檷涓村墠鐨勮€冮獙锛?,",
        "浣犲彲鐭ワ紝浣犵殑鑻﹂毦鍗冲皢鏁ｅ幓锛屾湁涓€浣嶅ぇ璐典汉锛屾甯︾潃娉曞害璺ㄥ北瓒婃捣鏉ユ浮浣狅紵",
        "浣犲彲鐭ワ紝浣犱箣鎵€浠ュ鐙紝鏄洜涓轰綘鐨勭伒榄傚眰娆″お楂橈紝鏅€氱殑涓栦織缂樺垎宸茬粡閰嶄笉涓婄幇鍦ㄧ殑浣狅紵",
        "浣犲彲鐭ワ紝浣犲湪杩欎釜瀹舵棌閲屾壆婕旂潃浠€涔堣鑹诧紵涓€涓鏃忚兘鍑轰綘杩欎箞涓€浣嶅績瀛樺杽蹇电殑浜猴紝閭ｆ槸绁栦笂绉簡澶╁ぇ鐨勫痉銆?,",
    ],
    "G绫?閲戝彞鐮撮鍨?": [
        "鏈€鍌荤殑浜嬶紝灏辨槸璺熺儌浜鸿閬撶悊銆?,",
        "浣犵殑鍠勮壇锛岃甯︾偣閿嬭姃銆?,",
        "涓嶆槸鎵€鏈夌殑蹇嶈閮藉彨澶у害锛屾湁鏃跺€欓偅鍙獫鍥娿€?,",
        "鐪熸鐨勯珮鎵嬶紝浠庝笉瑙ｉ噴锛屽彧鐢ㄧ粨鏋滆璇濄€?,",
        "浣犵殑鏃堕棿寰堣吹锛屽埆娴垂鍦ㄤ笉鍊煎緱鐨勪汉韬笂銆?,",
        "鍒妸鍒汉鎯冲緱澶ソ锛屼篃鍒妸鑷繁鐪嬪緱澶綆銆?,",
        "浣犱笉蹇呰濂芥墍鏈変汉锛屽仛濂借嚜宸卞氨澶熶簡銆?,",
        "鏈変簺璺紝娉ㄥ畾瑕佷竴涓汉璧帮紱鏈変簺鑻︼紝娉ㄥ畾瑕佷竴涓汉鎵涖€?,",
        "鐪熸鐨勫己澶э紝鏄唴蹇冪殑骞抽潤锛岃€屼笉鏄琛ㄧ殑寮虹‖銆?,",
        "浣犵殑鏍煎眬锛屽喅瀹氫簡浣犵殑缁撳眬銆?,",
        "榫欎笉涓庤泧灞咃紝鍑や笉涓庨浮鏍栥€?,",
        "澶ф櫤鑻ユ剼锛屽ぇ宸ц嫢鎷欍€傜湡姝ｅ帀瀹崇殑浜猴紝鍦ㄦ椂鏈烘湭鍒颁箣鍓嶏紝閮芥槸鐪嬬潃鏈€涓嶈捣鐪肩殑銆?,",
        "鏈ㄧ浜庢灄锛岄蹇呮懅涔嬨€傛鏄洜涓轰綘澶紭绉€锛屼綘韬笂鐨勭伒姘斿お瓒充簡锛屾墠鎷涙潵浜嗛偅闃存殫閲岀殑鑷櫕銆?,",
        "澶╅亾鏃犱翰锛屽父涓庡杽浜恒€備綘瀹堜綇浜嗗簳绾匡紝鑰佸ぉ鑷劧瀹堜綇浣犵殑涓€鍒囥€?,",
        "澶敮涓嶄簤锛屾晠澶╀笅鑾兘涓庝箣浜夈€備綘鐨勪笉浜変笉鎶紝鎭版伆鏄渶楂樼殑鏅烘収銆?,",
        "鑿╄惃浣庣湁鏄厛鎮诧紝閲戝垰鎬掔洰鏇存槸鎵嬫銆?,",
        "闈欐按娴佹繁锛屾櫤鑰呮棤瑷€銆備綘鐨勬矇榛樻槸闆烽渾涔嬪娍锛岃浠栧績鎱岋紝璁╀粬鎭愭儳銆?,",
    ],
}


def get_random_hooks(count=3):
    """闅忔満閫夋嫨鎸囧畾鏁伴噺鐨勪笉鍚岀被鍨嬪紑澶达紝纭繚鑷冲皯1涓槸瀹囧畽澶╅亾/鐐归啋鐩磋亰/韬唤璁ゅ悓绫诲瀷"""
    light_types = ["A绫?瀹囧畽澶╅亾鍨?", "B绫?鐐归啋鐩磋亰鍨?", "C绫?韬唤璁ゅ悓鍨?"]
    all_types = list(HOOK_LIBRARY.keys())

    # 纭繚鑷冲皯閫?涓交鏉剧被鍨?
    selected_types = [random.choice(light_types)]

    # 浠庡墿浣欑被鍨嬩腑鍐嶉€?count-1)涓?
    remaining = [t for t in all_types if t != selected_types[0]]
    additional_count = min(count - 1, len(remaining))
    selected_types.extend(random.sample(remaining, additional_count))

    # 鎵撲贡椤哄簭
    random.shuffle(selected_types)

    # 浠庢瘡涓被鍨嬩腑闅忔満閫変竴涓ず渚?
    result = []
    for t in selected_types:
        hook = random.choice(HOOK_LIBRARY[t])
        result.append({"type": t, "example": hook})

    return result

# 閰嶇疆鏂囦欢璺緞
CONFIG_FILE = os.path.join(os.path.dirname(__file__), "fangxie_config.json")

# 榛樿閰嶇疆
DEFAULT_CONFIG = {
    "use_stream": True,  # 鏄惁浣跨敤娴佸紡璋冪敤
    "similarity_threshold": 0.76,  # 鐢熸垚鏂囨鐩镐技搴﹂槇鍊硷紙瓒婁綆瓒婁弗鏍硷級
    # 璇煶鍚堟垚璺緞閰嶇疆
    "voice_input_path": r"D:/AIDownloadFiles/鍥藉json/鐧惧鍙峰甫璐ц棰?baijiadaihuo/input/瑙嗛鏂囨/娴侀噺鏂囨",
    "voice_output_path": r"D:/AIDownloadFiles/鍥藉json/鐧惧鍙峰甫璐ц棰?baijiadaihuo/input/瑙嗛閰嶉煶/娴侀噺璇煶",
    # 娴佸紡璋冪敤 - 涓绘ā鍨嬮厤缃?
    "stream_main_url": "https://api.aifuwu.icu/v1",
    "stream_main_key": "sk-hc6yUaXg89eK5UgUii10DPWmdaJZdPXqPbPcKSRbmWgxeeDK",
    "stream_main_model": "gemini-3-pro-preview",
    "stream_main_max_tokens": 16000,
    # 娴佸紡璋冪敤 - 澶囩敤妯″瀷閰嶇疆
    "stream_backup_url": "https://yunyi.rdzhvip.com/v1",
    "stream_backup_key": "A68YGT53-41HW-8GT7-353U-MJVDNY8NE5KX",
    "stream_backup_model": "claude-opus-4-5-20251101",
    "stream_backup_max_tokens": 16000,
    # 闈炴祦寮忚皟鐢?- 涓绘ā鍨嬮厤缃?
    "non_stream_main_url": "https://yunyi.rdzhvip.com/v1",
    "non_stream_main_key": "A68YGT53-41HW-8GT7-353U-MJVDNY8NE5KX",
    "non_stream_main_model": "claude-opus-4-5-20251101",
    "non_stream_main_max_tokens": 16000,
    # 闈炴祦寮忚皟鐢?- 澶囩敤妯″瀷閰嶇疆
    "non_stream_backup_url": "https://api.aifuwu.icu/v1",
    "non_stream_backup_key": "sk-hc6yUaXg89eK5UgUii10DPWmdaJZdPXqPbPcKSRbmWgxeeDK",
    "non_stream_backup_model": "gemini-3-pro-preview",
    "non_stream_backup_max_tokens": 16000,
    # 寮曟祦璇濇湳搴擄紙鎸夌被鍨嬪垎寮€瀛樺偍锛?
    "yinliu_templates": {
        "缃《寮曟祦": [
            "鍏充簬鎬庝箞褰诲簳璧板嚭杩欎釜鍥板眬锛屾垜鍦ㄤ富椤电疆椤惰棰戦噷璁插緱寰堥€忋€傞偅閲屾湁涓€濂楁柟娉曪紝鏄垜杩欎簺骞存懜鐖粴鎵撴€荤粨鍑烘潵鐨勶紝涓嶉€傚悎鍦ㄨ繖閲屽叕寮€璁层€備綘瑕佹槸鐪熸兂鏀瑰彉锛岀偣鎴戝ご鍍忥紝鍘荤湅缃《绗竴鏉★紝鐪嬪畬浣犲氨鏄庣櫧璇ユ€庝箞鍋氫簡銆?,",
            "鐜板湪浣犲氨鐐瑰紑鎴戠殑澶村儚锛岃繘鎴戜富椤碉紝鍘荤湅缃《鐨勯偅涓棰戯紝鎴戝湪閭ｉ噷缁欎綘鐣欎簡涓€濂楃牬灞€鐨勬柟娉曘€傚惉寰楁噦閭ｄ究鏄綘鐨勬敹鑾凤紝鍚笉鎳傝鏄庢椂鏈鸿繕娌″埌銆傚噯澶囧ソ浜嗗悧锛熻繖涓€灞€璇ヤ綘璧簡銆?,",
            "浠庝粖澶╁紑濮嬶紝瀛︾潃涓鸿嚜宸辨椿涓€娆?.. 绛斿簲鎴戯紝鐜板湪鐐瑰嚮鎴戠殑澶村儚杩涗富椤电湅鐪嬶紝缃《鍓嶄袱鏉¤棰戯紝閭ｉ噷鏈夋垜鎯冲浣犺鐨勫績閲岃瘽锛岃繖浜涜瘽鎴戝彧璇寸粰浣犲惉锛屽洜涓轰綘鎵涗簡澶箙锛屾€昏鏈変汉鎳備綘鐨勪笉鏄撱€?,",
            "濡傛灉浣犱篃鎯充粠杩欑鍥板閲岃蛋鍑烘潵锛屾兂娲诲緱閫氶€忎竴鐐广€佽交鏉句竴鐐癸紝鐐规垜澶村儚锛屽幓鐪嬩富椤电疆椤剁殑瑙嗛銆傞偅閲屾湁浣犱竴鐩村湪鎵剧殑绛旀锛屼篃鏈変綘闇€瑕佺殑閭ｄ唤鍔涢噺銆傛垜鍦ㄩ偅閲岀瓑浣犮€?,",
            "鐜板湪浣犲氨鐐瑰紑鎴戠殑澶村儚锛屽叆鎴戦棬搴紝鍘荤湅涓婚〉缃《鐨勮棰戙€傞偅閲屾湁涓€濂楀畬鏁寸殑鏂规硶锛屾槸鎴戠敤澶氬勾缁忛獙鎬荤粨鍑烘潵鐨勶紝涓撻棬鐣欑粰鍍忎綘杩欐牱鐨勪汉銆傜湅瀹屼箣鍚庯紝浣犲氨鐭ラ亾璇ユ€庝箞鍋氫簡銆?,",
            "鐐瑰嚮鎴戠殑澶村儚杩涗富椤碉紝鍘荤湅缃《绗竴涓棰戙€傞偅閲屾湁浣犱竴鐩村湪瀵绘壘鐨勭瓟妗堬紝鏈変綘闇€瑕佺殑閭ｄ唤鍔涢噺銆傝繖涓嶆槸宸у悎锛屾槸浣犺鐪嬪埌鐨勬椂鍊欏埌浜嗐€?,",
            "濡傛灉浣犳兂褰诲簳鏀瑰彉鐜扮姸锛岀偣鎴戝ご鍍忥紝鍘讳富椤电湅缃《瑙嗛銆傞偅閲屾湁涓€濂楃牬灞€鐨勬柟娉曪紝鏄垜杩欎簺骞存懜鐖粴鎵撴€荤粨鍑烘潵鐨勶紝涓嶉€傚悎鍦ㄨ繖閲屽叕寮€璁层€備綘瑕佹槸鐪熸兂鏀瑰彉锛屽氨鍘荤湅鐪嬨€?,",
            "绛斿簲鎴戯紝鐜板湪灏辩偣寮€鎴戠殑澶村儚锛岃繘涓婚〉鐪嬬疆椤剁殑閭ｄ袱鏉¤棰戙€傞偅閲屾湁鎴戞兂瀵逛綘璇寸殑蹇冮噷璇濓紝杩欎簺璇濇垜鍙缁欎綘鍚紝鍥犱负浣犳墰浜嗗お涔咃紝鎬昏鏈変汉鎳備綘鐨勪笉鏄撱€?,",
            "鍏充簬濡備綍璧板嚭鍥板锛屾垜鍦ㄤ富椤电疆椤惰棰戦噷璁插緱寰堥€忓交銆傞偅閲屾湁涓€濂楀畬鏁寸殑鏂规硶璁猴紝鏄垜鐢ㄥ骞寸粡楠屾崲鏉ョ殑銆備綘瑕佹槸鐪熸兂鏀瑰彉锛岀偣鎴戝ご鍍忥紝鍘荤湅缃《绗竴鏉°€?,",
            "鐜板湪浣犲彧闇€瑕佸仛涓€浠朵簨锛氱偣寮€鎴戠殑澶村儚锛岃繘涓婚〉锛屽幓鐪嬬疆椤剁殑瑙嗛銆傞偅閲屾湁浣犱竴鐩村湪鎵剧殑绛旀锛屼篃鏈変綘闇€瑕佺殑閭ｄ唤搴曟皵銆傛垜鍦ㄩ偅閲岀瓑浣犮€?,",
            "濡傛灉浣犱篃鎯充粠杩欑鍥板閲岃蛋鍑烘潵锛屾兂娲诲緱鏄庣櫧涓€鐐广€佽交鏉句竴鐐癸紝鐐规垜澶村儚锛屽幓鐪嬩富椤电疆椤剁殑瑙嗛銆傞偅閲屾湁浣犻渶瑕佺殑鏂规硶锛屼篃鏈変綘闇€瑕佺殑鍔涢噺銆?,",
            "鐐瑰紑鎴戠殑澶村儚锛岃繘鎴戜富椤碉紝鍘荤湅缃《鐨勯偅涓棰戙€傛垜鍦ㄩ偅閲岀粰浣犵暀浜嗕竴濂楃牬灞€鐨勬柟娉曪紝鍚緱鎳傞偅渚挎槸浣犵殑鏀惰幏锛屽惉涓嶆噦璇存槑鏃舵満杩樻病鍒般€傚噯澶囧ソ浜嗗悧锛?,",
            "鍏充簬鎬庝箞褰诲簳鏀瑰彉鐜扮姸锛屾垜鍦ㄤ富椤电疆椤惰棰戦噷璁插緱寰堟竻妤氥€傞偅閲屾湁涓€濂楁柟娉曪紝鏄垜杩欎簺骞存€荤粨鍑烘潵鐨勶紝涓撻棬鐣欑粰鍍忎綘杩欐牱鐨勪汉銆備綘瑕佹槸鐪熸兂鏀瑰彉锛屽氨鍘荤湅鐪嬨€?",
        ],
        "姗辩獥寮曟祦": [
            "鐜板湪浣犲彧闇€瑕佸仛涓€浠朵簨锛岀偣寮€鎴戠殑澶村儚锛岃繘鍏ヤ富椤垫┍绐楋紝涓嶈甯︾潃澶椤捐檻鍘绘寫鎷ｏ紝涔熶笉瑕侀棶鍝竴浠舵渶濂斤紝浣犻潤闈欑殑鐪嬨€傝嫢瀹冭浣犲績鏈夋墍鍔紝閭ｅ氨鏄€傚悎浣犵殑锛屽垏鑾敊杩囥€?,",
            "閫変竴浠讹紝鏄粰鑷繁涓€浠藉皬灏忕殑鐘掕祻锛涢€変袱浠讹紝鏄负鑷繁鐨勭敓娲诲娣讳竴鐐规俯鏆栵紱鑻ヤ綘鎰挎剰锛岄€変笁浠讹紝灏辨槸缁欒嚜宸变竴涓畬鏁寸殑绀肩墿锛岃鐢熸椿澶氫竴浜涚編濂姐€?,",
            "鍘诲惂锛屾湅鍙嬶紝鐐瑰紑澶村儚杩涙┍绐楋紝鍘绘壘鍥為偅涓湰璇ュ彂鍏夌殑鑷繁銆傛垜鍦ㄩ偅澶寸瓑鐫€锛岀湅浣犺秺鏉ヨ秺濂斤紝鑷湪濡傞銆?,",
            "閫変竴浠讹紝绗竴鐪煎叆蹇冪殑灏辨槸浣犳帴涓嬫潵鍏ㄥ姏鍐查攱鐨勫簳姘旓紱閫?浠讹紝渚挎槸瀵硅嚜宸变竴璺殣蹇嶅潥鎸佺殑鐘掕祻锛涢€?浠讹紝鏇翠唬琛ㄤ綘涓嶅彧椤剧潃褰撲笅璧惰矾锛岃繕鍦ㄤ负闀胯繙鏈潵甯冨眬绛硅皨銆?,",
            "缁撶紭涓€浠讹紝鏄负浣犲共娑哥殑杩愬娍寮曟潵涓€娉撴竻娉夛紱缁撶紭涓や欢锛屾槸涓轰綘瀛ょ嫭鐨勫舰鎴愯鏉ヤ竴浣嶈吹浜猴紱缁撶紭涓変欢锛屼究鏄负浣犵殑浜虹敓寮€鍚竴鎵囨柊鐨勫ぇ闂ㄣ€?,",
            "鍒啀绾犵粨銆傜洿鎺ョ偣寮€鎴戠殑澶村儚杩涗富椤碉紝缁嗙湅杩欓噷姣忎竴鏍烽兘涓嶆槸闅忔剰闄堝垪锛岃€屾槸涓撲负榛橀粯璧惰矾鐨勮拷鍏夎€呭噯澶囩殑寮哄姏鏀拺銆?,",
            "鐜板湪浣犺帿瑕佺姽璞紝鍗冲埢鐐瑰紑鎴戠殑澶村儚锛岃繘鍏ヤ富椤垫┍绐楀幓鐪嬬湅銆傞偅閲岀殑姣忎竴浠讹紝閮芥槸涓轰綘杩欐牱鐨勪汉鍑嗗鐨勩€?,",
            "鐐瑰紑澶村儚杩涙┍绐楋紝閫変竴浠惰浣犲績鍔ㄧ殑銆傝繖涓嶆槸娑堣垂锛屾槸瀵硅嚜宸辩殑鎶曡祫锛屾槸缁欒嚜宸辩殑绀肩墿銆?,",
            "鍘讳富椤垫┍绐楃湅鐪嬪惂锛岄偅閲屾湁鑳藉府浣犵殑涓滆タ銆備笉瑕侀棶鍝釜鏈€濂斤紝浣犵涓€鐪肩湅涓殑锛屽氨鏄渶閫傚悎浣犵殑銆?,",
            "鐜板湪灏辩偣寮€鎴戠殑澶村儚锛岃繘涓婚〉姗辩獥锛岄潤闈欏湴鐪嬨€傝嫢鏈変竴浠惰浣犲績鏈夋墍鍔紝閭ｅ氨鏄綘鐨勭紭鍒嗭紝鍒敊杩囥€?",
        ],
        "甯﹁揣寮曟祦": [
            "杞昏交涓€鐐瑰乏涓嬭锛屽畠灏辫兘鍒颁綘瀹躲€傛垜鐭ラ亾浣犱細鐘硅鲍锛屾€曟病鐢紝鎬曠櫧璐归挶锛屼絾鎴戝姖浣犲ぇ鑳嗚瘯杩欎竴娆°€傝繖涓嶆槸娑堣垂锛屾槸鎶曡祫锛屾姇璧勪綘鐨勫畨绋崇敓娲伙紝鎶曡祫浣犵殑浜虹敓搴曟皵銆傜偣鍑诲乏涓嬭鎶婂畠璇峰洖瀹?..",
            "濂戒笢瑗夸粠涓嶄細闀夸箙鍋滅暀锛岀珛鍒荤偣鍑诲乏涓嬭锛岃杩欎唤缇庡ソ涓轰綘鐨勭敓娲诲娣讳竴鐐规俯鏆栧拰鍔涢噺...",
            "濡備粖杩欏嚑浠朵笢瑗挎垜鎶婂畠鏀惧埌浜嗗乏涓嬭鐨勯€氶亾閲岋紝鍒啀鐘硅鲍浜嗭紝鍘诲乏涓嬭鎶婂畠璇峰洖瀹跺惂銆傝繖涓嶆槸鍦ㄥ府鍒汉锛屾槸鍦ㄥ府浣犺嚜宸?.. 瀵硅嚜宸卞ソ涓€鐐癸紝浣犲€煎緱銆?,",
            "鐜板湪璇烽『搴斿唴蹇冪殑鎸囧紩锛岀偣鍑讳笅鏂归€氶亾锛屾妸杩欐濂界墿甯﹀洖瀹躲€備綘涓哄埆浜烘搷蹇冧簡澶у崐杈堝瓙锛屾槸鏃跺€欏鑷繁濂戒竴鐐逛簡銆?,",
            "杞昏交涓€鐐瑰乏涓嬭锛屽畠灏辫兘鍒颁綘瀹讹紝浠庤繘闂ㄧ殑閭ｄ竴鍒昏捣锛屼綘鐨勫懡杩愰娇杞氨寮€濮嬪線濂界殑鏂瑰悜杞€?,",
            "鐐瑰嚮宸︿笅瑙掓妸瀹冭鍥炲锛屼笉璇曡瘯鎬庝箞鐭ラ亾鎺ヤ笅鏉ユ病鏈夊ソ浜嬬瓑鐫€浣狅紵琛屽姩灏卞湪姝ゅ埢锛屾敼鍙樺氨浠庣幇鍦ㄥ紑濮嬨€?,",
            "鎴戝凡灏嗘瀹濇斁杩涗笅鏂归€氶亾锛屽懡杩愪笉鍙┖璋堬紝鍞湁琛屽姩锛屾柟鑳芥敼鍙樸€傜偣鍑诲乏涓嬭锛屾妸瀹冨甫鍥炲銆?,",
            "绔嬪埢鐐瑰嚮宸︿笅瑙掞紝璁╄繖浠界編濂戒负浣犵殑鐢熸椿澧炴坊涓€鐐规俯鏆栧拰鍔涢噺銆備綘鍊煎緱鎷ユ湁鏇村ソ鐨勭敓娲汇€?,",
            "鐜板湪灏辩偣鍑讳笅鏂归€氶亾锛屾妸杩欐濂界墿甯﹀洖瀹躲€傝繖涓嶆槸鑺遍挶锛屾槸瀵硅嚜宸辩殑鎶曡祫锛屾槸缁欒嚜宸辩殑绀肩墿銆?,",
            "鍒啀鐘硅鲍浜嗭紝鐐瑰嚮宸︿笅瑙掓妸瀹冭鍥炲銆備綘涓哄埆浜轰粯鍑轰簡閭ｄ箞澶氾紝鏄椂鍊欏鑷繁濂戒竴鐐逛簡銆?",
        ]
    }
}

def load_config():
    """鍔犺浇閰嶇疆"""
    if os.path.exists(CONFIG_FILE):
        try:
            with open(CONFIG_FILE, 'r', encoding='utf-8') as f:
                config = json.load(f)
                # 琛ュ厖缂哄け鐨勯厤缃」
                for key in DEFAULT_CONFIG:
                    if key not in config:
                        config[key] = DEFAULT_CONFIG[key]
                return config
        except:
            pass
    return DEFAULT_CONFIG.copy()

def save_config(config):
    """淇濆瓨閰嶇疆"""
    with open(CONFIG_FILE, 'w', encoding='utf-8') as f:
        json.dump(config, f, ensure_ascii=False, indent=2)

class FangxieApp:
    def __init__(self, root):
        self.root = root
        self.root.title("鐧惧鍙锋枃妗堜豢鍐欏伐鍏?)
        self.root.geometry("1100x950")
        self.root.minsize(900, 700)
        self.root.resizable(True, True)

        # 鍔犺浇閰嶇疆
        self.config = load_config()

        # 绱犳潗搴撹矾寰勶紙鍥哄畾锛?
        self.material_path = r"D:\A鐧惧鍙峰甫璐ф枃妗堝簱"

        # 寮曟祦绫诲瀷鏄犲皠
        self.flow_types = {
            "缃《寮曟祦": "缃《瑙嗛寮曟祦绱犳潗.txt",
            "姗辩獥寮曟祦": "姗辩獥寮曟祦绱犳潗.txt",
            "甯﹁揣寮曟祦": "甯﹁揣寮曟祦绱犳潗.txt",
            "绾じ璧炰笉寮曟祦": None
        }

        # 瑙嗛鍒朵綔鐩稿叧閰嶇疆
        self.video_source_path = r"D:\BaiduNetdiskDownload\鑷劧椋庢櫙瑙嗛绱犳潗"
        self.max_videos_per_folder = 3
        self.ffmpeg_path = r"C:\Users\Administrator\AppData\Local\Microsoft\WinGet\Packages\Gyan.FFmpeg_Microsoft.Winget.Source_8wekyb3d8bbwe\ffmpeg-8.0.1-full_build\bin\ffmpeg.exe"
        self.ffprobe_path = r"C:\Users\Administrator\AppData\Local\Microsoft\WinGet\Packages\Gyan.FFmpeg_Microsoft.Winget.Source_8wekyb3d8bbwe\ffmpeg-8.0.1-full_build\bin\ffprobe.exe"
        self.whisper_model = None
        self.video_is_running = False

        self.create_widgets()

    def create_widgets(self):
        # 鍒涘缓涓籒otebook锛堟爣绛鹃〉锛?
        self.main_notebook = ttk.Notebook(self.root)
        self.main_notebook.pack(fill=tk.BOTH, expand=True, padx=5, pady=5)

        # === 鏍囩椤?锛氭枃妗堢敓鎴?===
        self.article_page = ttk.Frame(self.main_notebook)
        self.main_notebook.add(self.article_page, text="  鏂囨鐢熸垚  ")
        self.create_article_page()

        # === 鏍囩椤?锛氳棰戝埗浣?===
        self.video_page = ttk.Frame(self.main_notebook)
        self.main_notebook.add(self.video_page, text="  瑙嗛鍒朵綔  ")
        self.create_video_page()

        # === 鏍囩椤?锛欰PI閰嶇疆 ===
        self.api_page = ttk.Frame(self.main_notebook)
        self.main_notebook.add(self.api_page, text="  API閰嶇疆  ")
        self.create_api_page()

        # 鍒濆鍖栬繍琛岀姸鎬?
        self.is_running = False
        self.last_articles = []
        self.last_flow_type = ""
        self.last_yinliu_content = ""
        self.last_product_name = ""
        self.last_product_material = ""
        self.similarity_threshold = float(self.config.get("similarity_threshold", 0.76))

    def create_article_page(self):
        """鍒涘缓鏂囨鐢熸垚椤甸潰"""
        # 鍒涘缓婊氬姩妗嗘灦
        canvas = tk.Canvas(self.article_page)
        scrollbar = ttk.Scrollbar(self.article_page, orient="vertical", command=canvas.yview)
        scrollable_frame = ttk.Frame(canvas)

        scrollable_frame.bind(
            "<Configure>",
            lambda e: canvas.configure(scrollregion=canvas.bbox("all"))
        )

        # 鍒涘缓绐楀彛骞朵繚瀛業D锛岀敤浜庡悗缁皟鏁村搴?
        canvas_window = canvas.create_window((0, 0), window=scrollable_frame, anchor="nw")
        canvas.configure(yscrollcommand=scrollbar.set)

        # 璁╁唴瀹瑰搴﹀拰楂樺害璺熼殢Canvas鍙樺寲
        def on_canvas_configure(event):
            canvas.itemconfig(canvas_window, width=event.width)
            # 璁╁唴瀹归珮搴﹁嚦灏戠瓑浜庣敾甯冮珮搴︼紝杩欐牱鏃ュ織鍖哄煙鍙互鎵╁睍濉弧绐楀彛
            canvas.itemconfig(canvas_window, height=max(event.height, scrollable_frame.winfo_reqheight()))
        canvas.bind("<Configure>", on_canvas_configure)

        def _on_mousewheel(event):
            canvas.yview_scroll(int(-1*(event.delta/120)), "units")

        def bind_mousewheel_to_all(widget):
            """閫掑綊缁戝畾榧犳爣婊氳疆浜嬩欢鍒版墍鏈夊瓙缁勪欢"""
            widget.bind("<MouseWheel>", _on_mousewheel)
            for child in widget.winfo_children():
                bind_mousewheel_to_all(child)

        canvas.bind("<MouseWheel>", _on_mousewheel)
        scrollable_frame.bind("<MouseWheel>", _on_mousewheel)

        canvas.pack(side="left", fill="both", expand=True)
        scrollbar.pack(side="right", fill="y")

        main_frame = ttk.Frame(scrollable_frame, padding="10")
        main_frame.pack(fill=tk.BOTH, expand=True)

        # 鏍囬
        title_label = ttk.Label(main_frame, text="鏂囨鐢熸垚", font=("寰蒋闆呴粦", 14, "bold"))
        title_label.pack(pady=(0, 10))

        # 娴佸紡/闈炴祦寮忓彉閲忥紙鍦ˋPI閰嶇疆椤甸潰鍒囨崲锛?
        self.use_stream = tk.BooleanVar(value=self.config.get("use_stream", True))

        # === 鍙傝€冩枃妗堣緭鍏?===
        input_frame = ttk.LabelFrame(main_frame, text="鍙傝€冩枃妗?, padding="10")
        input_frame.pack(fill=tk.X, pady=5)

        # 杈撳叆鏂瑰紡閫夋嫨
        self.input_mode = tk.StringVar(value="file")
        mode_frame = ttk.Frame(input_frame)
        mode_frame.pack(fill=tk.X, pady=(0, 5))
        ttk.Radiobutton(mode_frame, text="涓婁紶鏂囦欢", variable=self.input_mode, value="file",
                       command=self.on_input_mode_change).pack(side=tk.LEFT, padx=10)
        ttk.Radiobutton(mode_frame, text="绮樿创鏂囨湰", variable=self.input_mode, value="paste",
                       command=self.on_input_mode_change).pack(side=tk.LEFT, padx=10)

        # 鏂囦欢閫夋嫨鍖哄煙
        self.file_input_frame = ttk.Frame(input_frame)
        self.file_input_frame.pack(fill=tk.X)
        self.input_path = tk.StringVar(value=r"d:\A鐧惧鍙峰甫璐ф枃妗堝簱\浠垮啓鏂囨.txt")
        ttk.Entry(self.file_input_frame, textvariable=self.input_path, width=70).pack(side=tk.LEFT, fill=tk.X, expand=True)
        ttk.Button(self.file_input_frame, text="閫夋嫨鏂囦欢", command=self.select_input_file, width=10).pack(side=tk.LEFT, padx=5)
        ttk.Button(self.file_input_frame, text="閫夋嫨鏂囦欢澶?, command=self.select_input_folder, width=10).pack(side=tk.LEFT)

        # 鏂囨湰绮樿创鍖哄煙锛堥粯璁ら殣钘忥級
        self.paste_input_frame = ttk.Frame(input_frame)
        self.paste_text = scrolledtext.ScrolledText(self.paste_input_frame, height=8, width=80)
        self.paste_text.pack(fill=tk.X)
        paste_btn_frame = ttk.Frame(self.paste_input_frame)
        paste_btn_frame.pack(fill=tk.X, pady=(5, 0))
        ttk.Button(paste_btn_frame, text="娓呯┖", command=lambda: self.paste_text.delete("1.0", tk.END), width=10).pack(side=tk.LEFT)

        # === 杈撳嚭璺緞 ===
        output_frame = ttk.LabelFrame(main_frame, text="杈撳嚭淇濆瓨璺緞", padding="10")
        output_frame.pack(fill=tk.X, pady=5)

        self.output_path = tk.StringVar(value=r"D:\A鐧惧鍙峰甫璐ц棰慭甯﹁揣鏂囨")
        ttk.Entry(output_frame, textvariable=self.output_path, width=70).pack(side=tk.LEFT, fill=tk.X, expand=True)
        ttk.Button(output_frame, text="閫夋嫨鏂囦欢澶?, command=self.select_output_folder, width=10).pack(side=tk.LEFT, padx=5)

        # === 娴侀噺鏂囨淇濆瓨璺緞 ===
        txt_output_frame = ttk.LabelFrame(main_frame, text="娴侀噺鏂囨淇濆瓨璺緞锛堢敓鎴愭祦閲忔枃妗堟寜閽娇鐢級", padding="10")
        txt_output_frame.pack(fill=tk.X, pady=5)

        self.txt_output_path = tk.StringVar(value=r"D:/AIDownloadFiles/鍥藉json/鐧惧鍙峰甫璐ц棰?baijiadaihuo/input/瑙嗛鏂囨/娴侀噺鏂囨")
        ttk.Entry(txt_output_frame, textvariable=self.txt_output_path, width=70).pack(side=tk.LEFT, fill=tk.X, expand=True)
        ttk.Button(txt_output_frame, text="閫夋嫨鏂囦欢澶?, command=self.select_txt_output_folder, width=10).pack(side=tk.LEFT, padx=5)

        # 瀛楁暟閫夋嫨
        ttk.Label(txt_output_frame, text="  瀛楁暟:").pack(side=tk.LEFT)
        self.word_count = tk.StringVar(value="1200")
        word_count_combo = ttk.Combobox(txt_output_frame, textvariable=self.word_count, width=8, state="readonly")
        word_count_combo['values'] = ["600", "700", "800", "900", "1000", "1200", "1500", "1600"]
        word_count_combo.pack(side=tk.LEFT, padx=5)

        # 鐢熸垚绡囨暟閫夋嫨
        ttk.Label(txt_output_frame, text="  绡囨暟:").pack(side=tk.LEFT)
        self.article_count = tk.StringVar(value="3")
        article_count_combo = ttk.Combobox(txt_output_frame, textvariable=self.article_count, width=5, state="readonly")
        article_count_combo['values'] = ["1", "2", "3", "5", "8"]
        article_count_combo.pack(side=tk.LEFT, padx=5)

        # === 璇煶鍚堟垚閰嶇疆 ===
        voice_config_frame = ttk.LabelFrame(main_frame, text="璇煶鍚堟垚閰嶇疆锛堝悎鎴愯闊虫寜閽娇鐢級", padding="10")
        voice_config_frame.pack(fill=tk.X, pady=5)

        # 闊宠壊閫夋嫨
        voice_row1 = ttk.Frame(voice_config_frame)
        voice_row1.pack(fill=tk.X, pady=2)
        ttk.Label(voice_row1, text="闊宠壊閫夋嫨:").pack(side=tk.LEFT)
        self.voice_type = tk.StringVar(value="鏅烘収鑰佽€?)
        self.voice_combo = ttk.Combobox(voice_row1, textvariable=self.voice_type, width=20, state="readonly")
        self.voice_combo['values'] = ["楦℃堡濂冲０", "鏅烘収鑰佽€?, "娌夌ǔ澶ф皵鐢峰０"]
        self.voice_combo.pack(side=tk.LEFT, padx=5)

        # 鏂囨杈撳叆鐩綍
        voice_row2 = ttk.Frame(voice_config_frame)
        voice_row2.pack(fill=tk.X, pady=2)
        ttk.Label(voice_row2, text="鏂囨鐩綍:").pack(side=tk.LEFT)
        self.voice_input_path = tk.StringVar(value=self.config.get("voice_input_path", r"D:/AIDownloadFiles/鍥藉json/鐧惧鍙峰甫璐ц棰?baijiadaihuo/input/瑙嗛鏂囨/娴侀噺鏂囨"))
        ttk.Entry(voice_row2, textvariable=self.voice_input_path, width=60).pack(side=tk.LEFT, fill=tk.X, expand=True, padx=5)
        ttk.Button(voice_row2, text="閫夋嫨", command=self.select_voice_input_folder, width=6).pack(side=tk.LEFT)

        # 閰嶉煶杈撳嚭鐩綍
        voice_row3 = ttk.Frame(voice_config_frame)
        voice_row3.pack(fill=tk.X, pady=2)
        ttk.Label(voice_row3, text="杈撳嚭鐩綍:").pack(side=tk.LEFT)
        self.voice_output_path = tk.StringVar(value=self.config.get("voice_output_path", r"D:/AIDownloadFiles/鍥藉json/鐧惧鍙峰甫璐ц棰?baijiadaihuo/input/瑙嗛閰嶉煶/娴侀噺璇煶"))
        ttk.Entry(voice_row3, textvariable=self.voice_output_path, width=60).pack(side=tk.LEFT, fill=tk.X, expand=True, padx=5)
        ttk.Button(voice_row3, text="閫夋嫨", command=self.select_voice_output_folder, width=6).pack(side=tk.LEFT)

        # === 寮曟祦绫诲瀷閫夋嫨 ===
        flow_frame = ttk.LabelFrame(main_frame, text="寮曟祦绫诲瀷", padding="10")
        flow_frame.pack(fill=tk.X, pady=5)

        self.flow_type = tk.StringVar(value="缃《寮曟祦")
        for ft in self.flow_types.keys():
            ttk.Radiobutton(flow_frame, text=ft, variable=self.flow_type, value=ft,
                           command=self.on_flow_type_change).pack(side=tk.LEFT, padx=10)

        # === 寮曟祦璇濇湳鍖哄煙 ===
        self.yinliu_frame = ttk.LabelFrame(main_frame, text="寮曟祦璇濇湳锛堝彲閫夛級", padding="10")
        self.yinliu_frame.pack(fill=tk.X, pady=5)

        # 璇濇湳涓嬫媺妗?
        yinliu_combo_frame = ttk.Frame(self.yinliu_frame)
        yinliu_combo_frame.pack(fill=tk.X, pady=(0, 5))
        ttk.Label(yinliu_combo_frame, text="璇烽€夋嫨缁撳熬鍙傝€冭瘽鏈?").pack(side=tk.LEFT)
        self.yinliu_combo = ttk.Combobox(yinliu_combo_frame, width=50, state="readonly")
        self.yinliu_combo.pack(side=tk.LEFT, padx=5)
        self.yinliu_combo.bind("<<ComboboxSelected>>", self.on_yinliu_select)

        # 璇濇湳鎿嶄綔鎸夐挳
        yinliu_btn_frame = ttk.Frame(self.yinliu_frame)
        yinliu_btn_frame.pack(fill=tk.X, pady=(0, 5))
        ttk.Button(yinliu_btn_frame, text="淇濆瓨褰撳墠璇濇湳", command=self.save_yinliu_template, width=12).pack(side=tk.LEFT, padx=2)
        ttk.Button(yinliu_btn_frame, text="鍒犻櫎閫変腑", command=self.delete_yinliu_template, width=10).pack(side=tk.LEFT, padx=2)
        ttk.Button(yinliu_btn_frame, text="娓呯┖杈撳叆", command=self.clear_yinliu_text, width=10).pack(side=tk.LEFT, padx=2)

        # 璇濇湳杈撳叆妗嗘彁绀?
        ttk.Label(self.yinliu_frame, text="鈫?鍦ㄦ杈撳叆缁撳熬寮曟祦璇濇湳锛圓I浼氬弬鑰冩椋庢牸锛屼负姣忕瘒鐢熸垚涓嶅悓鐨勭粨灏撅級锛?, foreground="gray").pack(anchor=tk.W)

        # 璇濇湳杈撳叆妗?
        self.yinliu_text = scrolledtext.ScrolledText(self.yinliu_frame, height=4, width=80)
        self.yinliu_text.pack(fill=tk.X)

        # 鍒濆鍖栬瘽鏈笅鎷夋
        self.update_yinliu_combo()

        # === 甯﹁揣淇℃伅鍖哄煙锛堥粯璁ら殣钘忥級 ===
        self.daihuo_frame = ttk.LabelFrame(main_frame, text="甯﹁揣鍟嗗搧淇℃伅", padding="10")

        # 鍟嗗搧鍚嶇О
        name_frame = ttk.Frame(self.daihuo_frame)
        name_frame.pack(fill=tk.X, pady=2)
        ttk.Label(name_frame, text="鍟嗗搧鍚嶇О:", width=10).pack(side=tk.LEFT)
        self.product_name = tk.StringVar()
        ttk.Entry(name_frame, textvariable=self.product_name, width=50).pack(side=tk.LEFT, fill=tk.X, expand=True)

        # 浜у搧绱犳潗
        ttk.Label(self.daihuo_frame, text="浜у搧绱犳潗/浠嬬粛:").pack(anchor=tk.W, pady=(5, 2))
        self.product_material = scrolledtext.ScrolledText(self.daihuo_frame, height=4, width=80)
        self.product_material.pack(fill=tk.X)

        # === 鎿嶄綔鎸夐挳 ===
        self.btn_frame_container = ttk.Frame(main_frame)
        self.btn_frame_container.pack(fill=tk.X, pady=10)

        btn_frame = ttk.Frame(self.btn_frame_container)
        btn_frame.pack()

        self.start_btn = ttk.Button(btn_frame, text="寮€濮嬬敓鎴?, command=self.start_generate, width=15)
        self.start_btn.pack(side=tk.LEFT, padx=10)

        self.start_txt_btn = ttk.Button(btn_frame, text="鐢熸垚娴侀噺鏂囨", command=self.start_generate_txt, width=12)
        self.start_txt_btn.pack(side=tk.LEFT, padx=10)

        self.synth_voice_btn = ttk.Button(btn_frame, text="鍚堟垚璇煶", command=self.start_synth_voice, width=12)
        self.synth_voice_btn.pack(side=tk.LEFT, padx=10)

        self.stop_btn = ttk.Button(btn_frame, text="鍋滄", command=self.stop_generate, width=10, state=tk.DISABLED)
        self.stop_btn.pack(side=tk.LEFT, padx=10)

        self.regenerate_btn = ttk.Button(btn_frame, text="閲嶆柊鐢熸垚", command=self.regenerate, width=12, state=tk.DISABLED)
        self.regenerate_btn.pack(side=tk.LEFT, padx=10)

        ttk.Button(btn_frame, text="鎵撳紑杈撳嚭鏂囦欢澶?, command=self.open_output_folder, width=15).pack(side=tk.LEFT, padx=10)
        ttk.Button(btn_frame, text="鍘诲埗浣滆棰?鈫?, command=self.go_to_video_page, width=12).pack(side=tk.LEFT, padx=10)

        # === 閲嶆柊鐢熸垚寤鸿 ===
        self.suggestion_frame = ttk.LabelFrame(main_frame, text="淇敼寤鸿", padding="10")
        self.suggestion_frame.pack(fill=tk.X, pady=5)

        ttk.Label(self.suggestion_frame, text="褰撴偍瀵圭敓鎴愮殑缁撴灉涓嶆弧鎰忔椂锛岃鍦ㄤ笅鏂硅緭鍏ユ剰瑙侊紝骞剁偣鍑婚噸鏂扮敓鎴愶細").pack(anchor=tk.W)
        self.suggestion_text = scrolledtext.ScrolledText(self.suggestion_frame, height=3, width=80)
        self.suggestion_text.pack(fill=tk.X)

        # === 杩涘害鏉?===
        progress_frame = ttk.Frame(main_frame)
        progress_frame.pack(fill=tk.X, pady=5)

        self.progress_var = tk.DoubleVar(value=0)
        self.progress_bar = ttk.Progressbar(progress_frame, variable=self.progress_var, maximum=100)
        self.progress_bar.pack(fill=tk.X)

        self.status_label = ttk.Label(progress_frame, text="灏辩华")
        self.status_label.pack(pady=5)

        # === 鏃ュ織鍖哄煙 ===
        log_frame = ttk.LabelFrame(main_frame, text="杩愯鏃ュ織", padding="10")
        log_frame.pack(fill=tk.BOTH, expand=True, pady=5)

        self.log_text = scrolledtext.ScrolledText(log_frame, height=20, width=100)
        self.log_text.pack(fill=tk.BOTH, expand=True)

        # 缁戝畾鎵€鏈夊瓙缁勪欢鐨勬粴杞簨浠?
        bind_mousewheel_to_all(main_frame)

    def create_api_page(self):
        """鍒涘缓API閰嶇疆椤甸潰"""
        main_frame = ttk.Frame(self.api_page, padding="15")
        main_frame.pack(fill=tk.BOTH, expand=True)

        # 鏍囬
        title_label = ttk.Label(main_frame, text="API閰嶇疆", font=("寰蒋闆呴粦", 14, "bold"))
        title_label.pack(pady=(0, 15))

        # 璇存槑鏂囧瓧
        ttk.Label(main_frame, text="閫夋嫨璋冪敤妯″紡锛岄厤缃搴旂殑API鍙傛暟銆傚垏鎹㈡爣绛鹃〉鍗冲垏鎹㈣皟鐢ㄦā寮忋€?, foreground="gray").pack(anchor=tk.W, pady=(0, 10))

        # 鍒涘缓Notebook鐢ㄤ簬鍒囨崲娴佸紡/闈炴祦寮忛厤缃?
        self.api_notebook = ttk.Notebook(main_frame)
        self.api_notebook.pack(fill=tk.BOTH, expand=True, pady=(0, 15))

        # === 娴佸紡璋冪敤閰嶇疆椤?===
        stream_page = ttk.Frame(self.api_notebook, padding="10")
        self.api_notebook.add(stream_page, text="  娴佸紡璋冪敤锛堟帹鑽愶級  ")

        # 娴佸紡-涓绘ā鍨?
        stream_main_frame = ttk.LabelFrame(stream_page, text="涓绘ā鍨?, padding="10")
        stream_main_frame.pack(fill=tk.X, pady=5)
        self._create_model_config(stream_main_frame, "stream_main")

        # 娴佸紡-澶囩敤妯″瀷
        stream_backup_frame = ttk.LabelFrame(stream_page, text="澶囩敤妯″瀷锛堜富妯″瀷澶辫触鍚庤嚜鍔ㄥ垏鎹級", padding="10")
        stream_backup_frame.pack(fill=tk.X, pady=5)
        self._create_model_config(stream_backup_frame, "stream_backup")

        # === 闈炴祦寮忚皟鐢ㄩ厤缃〉 ===
        non_stream_page = ttk.Frame(self.api_notebook, padding="10")
        self.api_notebook.add(non_stream_page, text="  闈炴祦寮忚皟鐢? ")

        # 闈炴祦寮?涓绘ā鍨?
        non_stream_main_frame = ttk.LabelFrame(non_stream_page, text="涓绘ā鍨?, padding="10")
        non_stream_main_frame.pack(fill=tk.X, pady=5)
        self._create_model_config(non_stream_main_frame, "non_stream_main")

        # 闈炴祦寮?澶囩敤妯″瀷
        non_stream_backup_frame = ttk.LabelFrame(non_stream_page, text="澶囩敤妯″瀷锛堜富妯″瀷澶辫触鍚庤嚜鍔ㄥ垏鎹級", padding="10")
        non_stream_backup_frame.pack(fill=tk.X, pady=5)
        self._create_model_config(non_stream_backup_frame, "non_stream_backup")

        # 缁戝畾鏍囩椤靛垏鎹簨浠讹紝鑷姩鏇存柊娴佸紡/闈炴祦寮忚缃?
        self.api_notebook.bind("<<NotebookTabChanged>>", self.on_api_tab_change)

        # 鏍规嵁閰嶇疆鍒濆鍖栭€変腑鐨勬爣绛鹃〉
        if not self.config.get("use_stream", True):
            self.api_notebook.select(1)  # 閫変腑闈炴祦寮忔爣绛鹃〉

        # 鍘绘ā鏉跨浉浼煎害闃堝€?
        threshold_frame = ttk.LabelFrame(main_frame, text="鍘绘ā鏉胯缃?, padding="10")
        threshold_frame.pack(fill=tk.X, pady=(0, 10))
        ttk.Label(threshold_frame, text="鐩镐技搴﹂槇鍊?").pack(side=tk.LEFT)
        self.similarity_threshold_var = tk.StringVar(
            value=f"{float(self.config.get('similarity_threshold', 0.76)):.2f}"
        )
        threshold_combo = ttk.Combobox(
            threshold_frame,
            textvariable=self.similarity_threshold_var,
            width=8,
            state="readonly"
        )
        threshold_combo["values"] = [
            "0.50", "0.55", "0.60", "0.65", "0.68", "0.70",
            "0.72", "0.74", "0.76", "0.78", "0.80", "0.82", "0.85", "0.90"
        ]
        threshold_combo.pack(side=tk.LEFT, padx=8)
        ttk.Label(
            threshold_frame,
            text="瓒婁綆瓒婁弗鏍硷紙閲嶈瘯鏇村锛夈€傚缓璁?0.72~0.80",
            foreground="gray"
        ).pack(side=tk.LEFT, padx=8)

        # 搴曢儴鎸夐挳鍖哄煙
        btn_frame = ttk.Frame(main_frame)
        btn_frame.pack(fill=tk.X, pady=10)

        ttk.Button(btn_frame, text="淇濆瓨閰嶇疆", command=self.save_api_config, width=15).pack(side=tk.LEFT, padx=5)
        ttk.Button(btn_frame, text="閲嶇疆榛樿", command=self.reset_api_config, width=15).pack(side=tk.LEFT, padx=5)

        # 褰撳墠妯″紡鎻愮ず
        self.api_mode_label = ttk.Label(main_frame, text="", font=("寰蒋闆呴粦", 10))
        self.api_mode_label.pack(anchor=tk.W, pady=5)
        self._update_api_mode_label()

    def _update_api_mode_label(self):
        """鏇存柊褰撳墠妯″紡鎻愮ず"""
        mode = "娴佸紡璋冪敤" if self.use_stream.get() else "闈炴祦寮忚皟鐢?
        self.api_mode_label.config(text=f"褰撳墠浣跨敤锛歿mode}")

    def create_video_page(self):
        """鍒涘缓瑙嗛鍒朵綔椤甸潰"""
        # 鍒涘缓婊氬姩妗嗘灦
        canvas = tk.Canvas(self.video_page)
        scrollbar = ttk.Scrollbar(self.video_page, orient="vertical", command=canvas.yview)
        scrollable_frame = ttk.Frame(canvas)

        scrollable_frame.bind(
            "<Configure>",
            lambda e: canvas.configure(scrollregion=canvas.bbox("all"))
        )

        canvas.create_window((0, 0), window=scrollable_frame, anchor="nw")
        canvas.configure(yscrollcommand=scrollbar.set)

        def _on_mousewheel(event):
            canvas.yview_scroll(int(-1*(event.delta/120)), "units")

        def bind_mousewheel_to_all(widget):
            """閫掑綊缁戝畾榧犳爣婊氳疆浜嬩欢鍒版墍鏈夊瓙缁勪欢"""
            widget.bind("<MouseWheel>", _on_mousewheel)
            for child in widget.winfo_children():
                bind_mousewheel_to_all(child)

        canvas.bind("<MouseWheel>", _on_mousewheel)
        scrollable_frame.bind("<MouseWheel>", _on_mousewheel)

        canvas.pack(side="left", fill="both", expand=True)
        scrollbar.pack(side="right", fill="y")

        main_frame = ttk.Frame(scrollable_frame, padding="10")
        main_frame.pack(fill=tk.BOTH, expand=True)

        # 鏍囬
        ttk.Label(main_frame, text="瑙嗛鍒朵綔", font=("寰蒋闆呴粦", 14, "bold")).pack(pady=(0, 10))

        # === 妯″紡閫夋嫨 ===
        mode_frame = ttk.LabelFrame(main_frame, text="1. 閫夋嫨鍒朵綔妯″紡", padding="10")
        mode_frame.pack(fill=tk.X, pady=5)

        self.video_mode = tk.StringVar(value="audio")
        ttk.Radiobutton(mode_frame, text="浣跨敤宸叉湁闊抽锛圵hisper璇嗗埆瀛楀箷锛?,
                       variable=self.video_mode, value="audio",
                       command=self.on_video_mode_change).pack(anchor=tk.W)
        ttk.Radiobutton(mode_frame, text="浠庣敓鎴愮殑鏂囨鍒朵綔锛圱TS鐢熸垚閰嶉煶锛?,
                       variable=self.video_mode, value="tts",
                       command=self.on_video_mode_change).pack(anchor=tk.W)

        # === 闊抽閫夋嫨鍖哄煙 ===
        self.audio_frame = ttk.LabelFrame(main_frame, text="2. 閫夋嫨閰嶉煶鏂囦欢", padding="10")
        self.audio_frame.pack(fill=tk.X, pady=5)

        audio_row = ttk.Frame(self.audio_frame)
        audio_row.pack(fill=tk.X)
        self.video_audio_path = tk.StringVar()
        ttk.Entry(audio_row, textvariable=self.video_audio_path, width=60).pack(side=tk.LEFT, fill=tk.X, expand=True)
        ttk.Button(audio_row, text="閫夋嫨闊抽", command=self.select_video_audio, width=10).pack(side=tk.LEFT, padx=5)

        # === 鏂囨閫夋嫨鍖哄煙锛圱TS妯″紡锛?===
        self.tts_frame = ttk.LabelFrame(main_frame, text="2. 閫夋嫨鏂囨", padding="10")

        tts_row1 = ttk.Frame(self.tts_frame)
        tts_row1.pack(fill=tk.X, pady=2)
        ttk.Label(tts_row1, text="閫夋嫨鏂囨:").pack(side=tk.LEFT)
        self.article_combo = ttk.Combobox(tts_row1, width=50, state="readonly")
        self.article_combo.pack(side=tk.LEFT, padx=5)
        self.article_combo.bind("<<ComboboxSelected>>", self.on_article_select)
        ttk.Button(tts_row1, text="鍒锋柊鍒楄〃", command=self.refresh_article_list, width=10).pack(side=tk.LEFT)
        ttk.Button(tts_row1, text="闅忔満閫夋嫨", command=self.random_select_article, width=10).pack(side=tk.LEFT, padx=5)

        ttk.Label(self.tts_frame, text="鏂囨棰勮:").pack(anchor=tk.W, pady=(5, 2))
        self.article_preview = scrolledtext.ScrolledText(self.tts_frame, height=4, width=80)
        self.article_preview.pack(fill=tk.X)

        # === 瑙嗛鏍囬 ===
        title_frame = ttk.LabelFrame(main_frame, text="3. 瑙嗛鏍囬锛堝皝闈㈡樉绀猴級", padding="10")
        title_frame.pack(fill=tk.X, pady=5)

        self.video_title = tk.StringVar()
        ttk.Entry(title_frame, textvariable=self.video_title, width=60).pack(fill=tk.X)

        # === 瑙嗛绱犳潗璁剧疆 ===
        video_frame = ttk.LabelFrame(main_frame, text="4. 瑙嗛绱犳潗璁剧疆", padding="10")
        video_frame.pack(fill=tk.X, pady=5)

        video_path_frame = ttk.Frame(video_frame)
        video_path_frame.pack(fill=tk.X, pady=2)
        ttk.Label(video_path_frame, text="绱犳潗鏂囦欢澶?").pack(side=tk.LEFT)
        self.video_source_var = tk.StringVar(value=self.video_source_path)
        ttk.Entry(video_path_frame, textvariable=self.video_source_var, width=50).pack(side=tk.LEFT, padx=5)
        ttk.Button(video_path_frame, text="閫夋嫨", command=self.select_video_source_folder, width=6).pack(side=tk.LEFT)

        video_param_frame = ttk.Frame(video_frame)
        video_param_frame.pack(fill=tk.X, pady=2)
        ttk.Label(video_param_frame, text="姣忔枃浠跺す鏈€澶氬彇:").pack(side=tk.LEFT)
        self.max_per_folder = tk.StringVar(value="3")
        ttk.Combobox(video_param_frame, textvariable=self.max_per_folder, width=5,
                    values=["1", "2", "3", "4", "5"]).pack(side=tk.LEFT, padx=5)
        ttk.Label(video_param_frame, text="涓棰?).pack(side=tk.LEFT)

        # 瑙嗛绱犳潗闈欓煶閫夐」
        video_mute_frame = ttk.Frame(video_frame)
        video_mute_frame.pack(fill=tk.X, pady=2)
        self.mute_video = tk.BooleanVar(value=True)
        ttk.Checkbutton(video_mute_frame, text="瑙嗛绱犳潗闈欓煶锛堝彧鐢ㄩ厤闊筹紝鎺ㄨ崘锛?, variable=self.mute_video).pack(side=tk.LEFT)

        # === 鑳屾櫙闊充箰璁剧疆 ===
        bgm_frame = ttk.LabelFrame(main_frame, text="5. 鑳屾櫙闊充箰锛堝彲閫夛級", padding="10")
        bgm_frame.pack(fill=tk.X, pady=5)

        bgm_enable_frame = ttk.Frame(bgm_frame)
        bgm_enable_frame.pack(fill=tk.X, pady=2)
        self.enable_bgm = tk.BooleanVar(value=False)
        ttk.Checkbutton(bgm_enable_frame, text="娣诲姞鑳屾櫙闊充箰", variable=self.enable_bgm,
                       command=self.on_bgm_toggle).pack(side=tk.LEFT)

        bgm_path_frame = ttk.Frame(bgm_frame)
        bgm_path_frame.pack(fill=tk.X, pady=2)
        ttk.Label(bgm_path_frame, text="BGM鏂囦欢澶?").pack(side=tk.LEFT)
        self.bgm_folder = tk.StringVar(value=r"D:\A鐧惧鍙峰甫璐ц棰慭BGM")
        self.bgm_entry = ttk.Entry(bgm_path_frame, textvariable=self.bgm_folder, width=45)
        self.bgm_entry.pack(side=tk.LEFT, padx=5)
        self.bgm_btn = ttk.Button(bgm_path_frame, text="閫夋嫨", command=self.select_bgm_folder, width=6)
        self.bgm_btn.pack(side=tk.LEFT)

        bgm_vol_frame = ttk.Frame(bgm_frame)
        bgm_vol_frame.pack(fill=tk.X, pady=2)
        ttk.Label(bgm_vol_frame, text="BGM闊抽噺:").pack(side=tk.LEFT)
        self.bgm_volume = tk.StringVar(value="15")
        ttk.Combobox(bgm_vol_frame, textvariable=self.bgm_volume, width=5,
                    values=["5", "10", "15", "20", "25", "30"]).pack(side=tk.LEFT, padx=5)
        ttk.Label(bgm_vol_frame, text="%").pack(side=tk.LEFT)

        ttk.Label(bgm_vol_frame, text="閰嶉煶闊抽噺:").pack(side=tk.LEFT, padx=(20, 0))
        self.voice_volume = tk.StringVar(value="100")
        ttk.Combobox(bgm_vol_frame, textvariable=self.voice_volume, width=5,
                    values=["80", "90", "100", "110", "120"]).pack(side=tk.LEFT, padx=5)
        ttk.Label(bgm_vol_frame, text="%").pack(side=tk.LEFT)

        # === 瀛楀箷鏍峰紡 ===
        style_frame = ttk.LabelFrame(main_frame, text="6. 瀛楀箷鏍峰紡", padding="10")
        style_frame.pack(fill=tk.X, pady=5)

        style_row1 = ttk.Frame(style_frame)
        style_row1.pack(fill=tk.X, pady=2)

        ttk.Label(style_row1, text="瀛楀彿:").pack(side=tk.LEFT)
        self.subtitle_size = tk.StringVar(value="36")
        subtitle_size_combo = ttk.Combobox(style_row1, textvariable=self.subtitle_size, width=6,
                     values=["24", "28", "32", "36", "40", "48"])
        subtitle_size_combo.pack(side=tk.LEFT, padx=5)
        subtitle_size_combo.bind("<<ComboboxSelected>>", self.update_subtitle_preview)

        ttk.Label(style_row1, text="棰滆壊:").pack(side=tk.LEFT, padx=(15, 0))
        self.subtitle_color = tk.StringVar(value="閲戣壊")
        subtitle_color_combo = ttk.Combobox(style_row1, textvariable=self.subtitle_color, width=8,
                     values=["閲戣壊", "榛勮壊", "姗欒壊", "绾㈣壊"])
        subtitle_color_combo.pack(side=tk.LEFT, padx=5)
        subtitle_color_combo.bind("<<ComboboxSelected>>", self.update_subtitle_preview)

        ttk.Label(style_row1, text="浣嶇疆:").pack(side=tk.LEFT, padx=(15, 0))
        self.subtitle_position = tk.StringVar(value="搴曢儴")
        subtitle_pos_combo = ttk.Combobox(style_row1, textvariable=self.subtitle_position, width=8,
                     values=["搴曢儴", "涓儴", "椤堕儴"])
        subtitle_pos_combo.pack(side=tk.LEFT, padx=5)
        subtitle_pos_combo.bind("<<ComboboxSelected>>", self.update_subtitle_preview)

        # 瀛楀箷棰勮鍖哄煙
        preview_row = ttk.Frame(style_frame)
        preview_row.pack(fill=tk.X, pady=(10, 5))

        self.subtitle_preview_canvas = tk.Canvas(preview_row, width=400, height=80, bg="#1a1a2e", highlightthickness=1, highlightbackground="gray")
        self.subtitle_preview_canvas.pack(side=tk.LEFT, padx=5)
        self.update_subtitle_preview()  # 鍒濆鍖栭瑙?

        # === 灏侀潰鏍峰紡 ===
        cover_frame = ttk.LabelFrame(main_frame, text="7. 灏侀潰鏍峰紡", padding="10")
        cover_frame.pack(fill=tk.X, pady=5)

        cover_row = ttk.Frame(cover_frame)
        cover_row.pack(fill=tk.X, pady=2)

        ttk.Label(cover_row, text="鏍囬瀛楀彿:").pack(side=tk.LEFT)
        self.cover_font_size = tk.StringVar(value="60")
        ttk.Combobox(cover_row, textvariable=self.cover_font_size, width=6,
                     values=["48", "54", "60", "72", "80"]).pack(side=tk.LEFT, padx=5)

        ttk.Label(cover_row, text="鏍囬棰滆壊:").pack(side=tk.LEFT, padx=(15, 0))
        self.cover_color = tk.StringVar(value="閲戣壊")
        ttk.Combobox(cover_row, textvariable=self.cover_color, width=8,
                     values=["閲戣壊", "榛勮壊", "姗欒壊", "绾㈣壊"]).pack(side=tk.LEFT, padx=5)

        ttk.Label(cover_row, text="鑳屾櫙:").pack(side=tk.LEFT, padx=(15, 0))
        self.cover_bg = tk.StringVar(value="瑙嗛棣栧抚")
        ttk.Combobox(cover_row, textvariable=self.cover_bg, width=10,
                     values=["瑙嗛棣栧抚", "榛戣壊鑳屾櫙"]).pack(side=tk.LEFT, padx=5)

        # === 杈撳嚭璁剧疆 ===
        output_frame = ttk.LabelFrame(main_frame, text="8. 杈撳嚭璁剧疆", padding="10")
        output_frame.pack(fill=tk.X, pady=5)

        self.video_output_path = tk.StringVar(value=r"D:\A鐧惧鍙峰甫璐ц棰慭鎴愬搧瑙嗛")
        ttk.Entry(output_frame, textvariable=self.video_output_path, width=60).pack(side=tk.LEFT, fill=tk.X, expand=True)
        ttk.Button(output_frame, text="閫夋嫨鏂囦欢澶?, command=self.select_video_output_folder, width=10).pack(side=tk.LEFT, padx=5)

        # === 鎿嶄綔鎸夐挳 ===
        btn_frame = ttk.Frame(main_frame)
        btn_frame.pack(pady=15)

        self.video_make_btn = ttk.Button(btn_frame, text="寮€濮嬪埗浣?, command=self.start_make_video, width=15)
        self.video_make_btn.pack(side=tk.LEFT, padx=10)

        self.video_stop_btn = ttk.Button(btn_frame, text="鍋滄", command=self.stop_make_video, width=10, state=tk.DISABLED)
        self.video_stop_btn.pack(side=tk.LEFT, padx=10)

        ttk.Button(btn_frame, text="鎵撳紑杈撳嚭鏂囦欢澶?, command=self.open_video_output, width=15).pack(side=tk.LEFT, padx=10)

        # === 杩涘害 ===
        progress_frame = ttk.Frame(main_frame)
        progress_frame.pack(fill=tk.X, pady=5)

        self.video_progress_var = tk.DoubleVar(value=0)
        self.video_progress_bar = ttk.Progressbar(progress_frame, variable=self.video_progress_var, maximum=100)
        self.video_progress_bar.pack(fill=tk.X)

        self.video_status_label = ttk.Label(progress_frame, text="灏辩华")
        self.video_status_label.pack(pady=5)

        # === 鏃ュ織 ===
        log_frame = ttk.LabelFrame(main_frame, text="鍒朵綔鏃ュ織", padding="10")
        log_frame.pack(fill=tk.BOTH, expand=True, pady=5)

        self.video_log_text = scrolledtext.ScrolledText(log_frame, height=10, width=80)
        self.video_log_text.pack(fill=tk.BOTH, expand=True)

        # 缁戝畾鎵€鏈夊瓙缁勪欢鐨勬粴杞簨浠?
        bind_mousewheel_to_all(main_frame)

    def go_to_video_page(self):
        """鍒囨崲鍒拌棰戝埗浣滈〉闈?""
        self.main_notebook.select(1)
        self.refresh_article_list()

    def on_video_mode_change(self):
        """瑙嗛鍒朵綔妯″紡鍒囨崲"""
        mode = self.video_mode.get()
        if mode == "audio":
            self.audio_frame.pack(fill=tk.X, pady=5, after=self.audio_frame.master.winfo_children()[1])
            self.tts_frame.pack_forget()
        else:
            self.tts_frame.pack(fill=tk.X, pady=5, after=self.tts_frame.master.winfo_children()[1])
            self.audio_frame.pack_forget()
            self.refresh_article_list()

    def select_video_audio(self):
        """閫夋嫨瑙嗛閰嶉煶鏂囦欢"""
        file_path = filedialog.askopenfilename(
            title="閫夋嫨閰嶉煶鏂囦欢",
            filetypes=[("闊抽鏂囦欢", "*.mp3 *.wav *.m4a"), ("鎵€鏈夋枃浠?, "*.*")]
        )
        if file_path:
            self.video_audio_path.set(file_path)

    def select_video_source_folder(self):
        """閫夋嫨瑙嗛绱犳潗鏂囦欢澶?""
        folder_path = filedialog.askdirectory(title="閫夋嫨瑙嗛绱犳潗鏂囦欢澶?)
        if folder_path:
            self.video_source_var.set(folder_path)

    def select_bgm_folder(self):
        """閫夋嫨BGM鏂囦欢澶?""
        folder_path = filedialog.askdirectory(title="閫夋嫨鑳屾櫙闊充箰鏂囦欢澶?)
        if folder_path:
            self.bgm_folder.set(folder_path)

    def on_bgm_toggle(self):
        """BGM寮€鍏冲垏鎹?""
        pass  # 鐣岄潰浼氳嚜鍔ㄦ牴鎹甤heckbox鐘舵€佸鐞?

    def get_random_bgm(self):
        """浠嶣GM鏂囦欢澶归殢鏈洪€夋嫨涓€棣栬儗鏅煶涔?""
        bgm_folder = self.bgm_folder.get()
        if not os.path.exists(bgm_folder):
            return None

        bgm_files = []
        for f in os.listdir(bgm_folder):
            if f.lower().endswith(('.mp3', '.wav', '.m4a', '.aac', '.flac')):
                bgm_files.append(os.path.join(bgm_folder, f))

        if bgm_files:
            return random.choice(bgm_files)
        return None

    def select_video_output_folder(self):
        """閫夋嫨瑙嗛杈撳嚭鏂囦欢澶?""
        folder_path = filedialog.askdirectory(title="閫夋嫨杈撳嚭鏂囦欢澶?)
        if folder_path:
            self.video_output_path.set(folder_path)

    def open_video_output(self):
        """鎵撳紑瑙嗛杈撳嚭鏂囦欢澶?""
        output_path = self.video_output_path.get().replace('/', '\\')
        os.makedirs(output_path, exist_ok=True)
        os.startfile(output_path)

    def refresh_article_list(self):
        """鍒锋柊鏂囨鍒楄〃"""
        options = ["-- 璇烽€夋嫨鏂囨 --"]
        if self.last_articles:
            for i, article in enumerate(self.last_articles):
                # 鎻愬彇鏍囬锛堝鏋滄湁锛?
                lines = article.strip().split('\n')
                title = ""
                for line in lines:
                    if line.startswith('銆愭爣棰?):
                        title = line.replace('銆愭爣棰?銆?, '').replace('銆愭爣棰?銆?, '').replace('銆愭爣棰?銆?, '').strip()[:20]
                        break
                if not title:
                    title = article[:25].replace('\n', ' ')
                options.append(f"鏂囨{i+1}: {title}...")
        self.article_combo['values'] = options
        self.article_combo.current(0)

    def on_article_select(self, event=None):
        """閫夋嫨鏂囨鏃舵樉绀洪瑙?""
        selection = self.article_combo.current()
        if selection > 0 and self.last_articles:
            article = self.last_articles[selection - 1]
            self.article_preview.delete("1.0", tk.END)
            self.article_preview.insert("1.0", article[:800] + "..." if len(article) > 800 else article)
            # 鑷姩鎻愬彇鏍囬
            lines = article.strip().split('\n')
            for line in lines:
                if line.startswith('銆愭爣棰?):
                    title = line.replace('銆愭爣棰?銆?, '').replace('銆愭爣棰?銆?, '').replace('銆愭爣棰?銆?, '').strip()
                    self.video_title.set(title)
                    break

    def random_select_article(self):
        """闅忔満閫夋嫨涓€绡囨枃妗?""
        if self.last_articles:
            idx = random.randint(0, len(self.last_articles) - 1)
            self.article_combo.current(idx + 1)
            self.on_article_select()
        else:
            messagebox.showinfo("鎻愮ず", "娌℃湁鍙敤鐨勬枃妗堬紝璇峰厛鍦ㄣ€屾枃妗堢敓鎴愩€嶉〉闈㈢敓鎴愭枃妗?)

    def video_log(self, message):
        """瑙嗛鍒朵綔鏃ュ織"""
        timestamp = datetime.now().strftime("%H:%M:%S")
        self.video_log_text.insert(tk.END, f"[{timestamp}] {message}\n")
        self.video_log_text.see(tk.END)
        self.root.update()

    def video_update_status(self, message):
        """鏇存柊瑙嗛鍒朵綔鐘舵€?""
        self.video_status_label.config(text=message)
        self.root.update()

    def video_update_progress(self, value):
        """鏇存柊瑙嗛鍒朵綔杩涘害"""
        self.video_progress_var.set(value)
        self.root.update()

    def stop_make_video(self):
        """鍋滄瑙嗛鍒朵綔"""
        self.video_is_running = False
        self.video_log("鐢ㄦ埛鍋滄浜嗗埗浣滀换鍔?)
        self.video_update_status("宸插仠姝?)

    def start_make_video(self):
        """寮€濮嬪埗浣滆棰?""
        mode = self.video_mode.get()

        if mode == "audio":
            # 闊抽妯″紡锛氶渶瑕侀€夋嫨闊抽鏂囦欢
            audio_path = self.video_audio_path.get()
            if not audio_path or not os.path.exists(audio_path):
                messagebox.showerror("閿欒", "璇烽€夋嫨鏈夋晥鐨勯厤闊虫枃浠?)
                return
        else:
            # TTS妯″紡锛氶渶瑕侀€夋嫨鏂囨
            selection = self.article_combo.current()
            if selection <= 0 or not self.last_articles:
                messagebox.showerror("閿欒", "璇烽€夋嫨涓€绡囨枃妗堬紝鎴栧厛鍦ㄣ€屾枃妗堢敓鎴愩€嶉〉闈㈢敓鎴愭枃妗?)
                return

        video_title = self.video_title.get().strip()
        if not video_title:
            messagebox.showerror("閿欒", "璇疯緭鍏ヨ棰戞爣棰?)
            return

        # 鏇存柊閰嶇疆
        self.video_source_path = self.video_source_var.get()
        self.max_videos_per_folder = int(self.max_per_folder.get())

        output_path = self.video_output_path.get()
        os.makedirs(output_path, exist_ok=True)

        self.video_make_btn.config(state=tk.DISABLED)
        self.video_stop_btn.config(state=tk.NORMAL)
        self.video_is_running = True
        self.video_log_text.delete(1.0, tk.END)

        # 鍦ㄧ嚎绋嬩腑鎵ц
        thread = threading.Thread(target=self.make_video_task, args=(mode, video_title, output_path))
        thread.daemon = True
        thread.start()

    def make_video_task(self, mode, video_title, output_path):
        """瑙嗛鍒朵綔涓讳换鍔?""
        import subprocess

        try:
            temp_dir = os.path.join(output_path, "temp")
            os.makedirs(temp_dir, exist_ok=True)

            # 1. 鑾峰彇闊抽
            if mode == "audio":
                audio_path = self.video_audio_path.get()
                self.video_log(f"浣跨敤宸叉湁闊抽: {os.path.basename(audio_path)}")
            else:
                # TTS鐢熸垚閰嶉煶
                self.video_update_status("姝ｅ湪鐢熸垚閰嶉煶...")
                self.video_update_progress(5)
                selection = self.article_combo.current()
                article = self.last_articles[selection - 1]
                # 鎻愬彇姝ｆ枃锛堝幓鎺夋爣棰橀儴鍒嗭級
                text_content = self.extract_article_text(article)
                audio_path = self.generate_tts(text_content, temp_dir)
                if not audio_path:
                    self.video_log("TTS閰嶉煶鐢熸垚澶辫触")
                    self.finish_video_task()
                    return

            if not self.video_is_running:
                self.finish_video_task()
                return

            # 2. 鑾峰彇闊抽鏃堕暱
            self.video_update_status("姝ｅ湪鍒嗘瀽闊抽...")
            self.video_update_progress(10)
            audio_duration = self.get_audio_duration(audio_path)
            self.video_log(f"闊抽鏃堕暱: {audio_duration:.1f} 绉?)

            if not self.video_is_running:
                self.finish_video_task()
                return

            # 3. 鐢熸垚瀛楀箷锛圵hisper璇嗗埆锛?
            self.video_update_status("姝ｅ湪璇嗗埆瀛楀箷...")
            self.video_update_progress(15)
            srt_path = self.transcribe_audio_for_video(audio_path, temp_dir)

            if not self.video_is_running:
                self.finish_video_task()
                return

            # 4. 鏀堕泦瑙嗛绱犳潗
            self.video_update_status("姝ｅ湪鏀堕泦瑙嗛绱犳潗...")
            self.video_update_progress(30)
            video_files = self.collect_videos_for_video(audio_duration + 5)

            if not video_files:
                self.video_log("閿欒锛氭病鏈夋敹闆嗗埌瑙嗛绱犳潗")
                self.finish_video_task()
                return

            if not self.video_is_running:
                self.finish_video_task()
                return

            # 5. 鍚堝苟瑙嗛绱犳潗
            self.video_update_status("姝ｅ湪鍚堝苟瑙嗛绱犳潗...")
            self.video_update_progress(45)

            list_file = os.path.join(temp_dir, "videos.txt")
            with open(list_file, 'w', encoding='utf-8') as f:
                for vf in video_files:
                    vf_fixed = vf.replace('\\', '/')
                    f.write(f"file '{vf_fixed}'\n")

            merged_video = os.path.join(temp_dir, "merged.mp4")
            cmd = f'"{self.ffmpeg_path}" -y -f concat -safe 0 -i "{list_file}" -c copy "{merged_video}"'
            self.video_log("鍚堝苟瑙嗛绱犳潗...")
            subprocess.run(cmd, shell=True, capture_output=True, text=True)

            # 6. 瑁佸壀鍒伴煶棰戦暱搴?
            self.video_update_status("姝ｅ湪瑁佸壀瑙嗛...")
            self.video_update_progress(55)
            trimmed_video = os.path.join(temp_dir, "trimmed.mp4")

            # 鏍规嵁鏄惁闈欓煶瑙嗛绱犳潗閫夋嫨涓嶅悓鍛戒护
            if self.mute_video.get():
                # 闈欓煶瑙嗛绱犳潗
                cmd = f'"{self.ffmpeg_path}" -y -i "{merged_video}" -t {audio_duration} -an -c:v copy "{trimmed_video}"'
            else:
                # 淇濈暀瑙嗛鍘熷０
                cmd = f'"{self.ffmpeg_path}" -y -i "{merged_video}" -t {audio_duration} -c copy "{trimmed_video}"'
            subprocess.run(cmd, shell=True, capture_output=True)

            # 7. 娣诲姞闊抽锛堥厤闊?+ 鍙€塀GM锛?
            self.video_update_status("姝ｅ湪娣诲姞闊抽...")
            self.video_update_progress(60)
            with_audio = os.path.join(temp_dir, "with_audio.mp4")

            # 鑾峰彇闊抽噺璁剧疆
            voice_vol = int(self.voice_volume.get()) / 100

            if self.enable_bgm.get():
                # 娣诲姞BGM
                bgm_path = self.get_random_bgm()
                if bgm_path:
                    bgm_vol = int(self.bgm_volume.get()) / 100
                    self.video_log(f"娣诲姞鑳屾櫙闊充箰: {os.path.basename(bgm_path)}")
                    # 娣峰悎閰嶉煶鍜孊GM
                    filter_complex = f"[1:a]volume={voice_vol}[voice];[2:a]volume={bgm_vol},aloop=loop=-1:size=2e+09[bgm];[voice][bgm]amix=inputs=2:duration=first:dropout_transition=2[aout]"
                    cmd = f'"{self.ffmpeg_path}" -y -i "{trimmed_video}" -i "{audio_path}" -i "{bgm_path}" -filter_complex "{filter_complex}" -map 0:v -map "[aout]" -c:v copy -c:a aac -b:a 192k -t {audio_duration} "{with_audio}"'
                else:
                    self.video_log("鏈壘鍒癇GM鏂囦欢锛岃烦杩囪儗鏅煶涔?)
                    cmd = f'"{self.ffmpeg_path}" -y -i "{trimmed_video}" -i "{audio_path}" -filter_complex "[1:a]volume={voice_vol}[aout]" -map 0:v -map "[aout]" -c:v copy -c:a aac -b:a 192k "{with_audio}"'
            else:
                # 鍙坊鍔犻厤闊?
                if voice_vol != 1.0:
                    cmd = f'"{self.ffmpeg_path}" -y -i "{trimmed_video}" -i "{audio_path}" -filter_complex "[1:a]volume={voice_vol}[aout]" -map 0:v -map "[aout]" -c:v copy -c:a aac -b:a 192k "{with_audio}"'
                else:
                    cmd = f'"{self.ffmpeg_path}" -y -i "{trimmed_video}" -i "{audio_path}" -map 0:v:0 -map 1:a:0 -c:v copy -c:a aac -b:a 192k "{with_audio}"'

            result = subprocess.run(cmd, shell=True, capture_output=True, text=True)
            if result.returncode != 0:
                self.video_log(f"闊抽澶勭悊璀﹀憡: {result.stderr[:200] if result.stderr else ''}")

            if not self.video_is_running:
                self.finish_video_task()
                return

            # 8. 鐑у綍瀛楀箷锛堝亸绉?绉掞紝琛ュ伩灏侀潰鏃堕暱锛?
            self.video_update_status("姝ｅ湪鐑у綍瀛楀箷...")
            self.video_update_progress(70)
            with_subtitle = self.burn_subtitles_for_video(with_audio, srt_path, temp_dir, time_offset=1)

            if not self.video_is_running:
                self.finish_video_task()
                return

            # 9. 鐢熸垚灏侀潰
            self.video_update_status("姝ｅ湪鐢熸垚灏侀潰...")
            self.video_update_progress(85)
            final_video = self.add_cover_for_video(with_subtitle, video_title, temp_dir)

            # 10. 鐢熸垚鏈€缁堟枃浠?
            self.video_update_status("姝ｅ湪鐢熸垚鏈€缁堣棰?..")
            self.video_update_progress(95)

            timestamp = datetime.now().strftime("%Y%m%d%H%M%S")
            safe_title_file = re.sub(r'[\\/:*?"<>|]', '', video_title)[:20]
            final_output = os.path.join(output_path, f"{timestamp}_{safe_title_file}.mp4")

            import shutil
            shutil.copy(final_video, final_output)

            # 11. 娓呯悊涓存椂鏂囦欢
            self.video_log("娓呯悊涓存椂鏂囦欢...")
            try:
                shutil.rmtree(temp_dir)
            except:
                pass

            self.video_update_progress(100)
            self.video_log(f"\n瑙嗛鍒朵綔瀹屾垚锛?)
            self.video_log(f"杈撳嚭鏂囦欢: {final_output}")

            self.finish_video_task()
            messagebox.showinfo("瀹屾垚", f"瑙嗛鍒朵綔瀹屾垚锛乗n淇濆瓨浣嶇疆: {final_output}")

        except Exception as e:
            self.video_log(f"閿欒: {str(e)}")
            import traceback
            self.video_log(traceback.format_exc())
            self.finish_video_task()

    def extract_article_text(self, article):
        """浠庢枃妗堜腑鎻愬彇姝ｆ枃锛堝幓鎺夋爣棰橈級"""
        lines = article.strip().split('\n')
        text_lines = []
        in_content = False
        for line in lines:
            if line.startswith('---') or line.startswith('鈺?):
                in_content = True
                continue
            if in_content and not line.startswith('銆愭爣棰?):
                text_lines.append(line.strip())
        return '\n'.join(text_lines) if text_lines else article

    def generate_tts(self, text, temp_dir):
        """浣跨敤TTS鐢熸垚閰嶉煶"""
        try:
            import edge_tts
            import asyncio

            self.video_log("姝ｅ湪浣跨敤Edge TTS鐢熸垚閰嶉煶...")
            audio_path = os.path.join(temp_dir, "tts_audio.mp3")

            async def generate():
                communicate = edge_tts.Communicate(text, "zh-CN-YunxiNeural")
                await communicate.save(audio_path)

            asyncio.run(generate())
            self.video_log("TTS閰嶉煶鐢熸垚鎴愬姛")
            return audio_path

        except ImportError:
            self.video_log("閿欒锛氭湭瀹夎edge-tts锛岃杩愯: pip install edge-tts")
            return None
        except Exception as e:
            self.video_log(f"TTS鐢熸垚閿欒: {str(e)}")
            return None

    def get_audio_duration(self, audio_path):
        """鑾峰彇闊抽鏃堕暱"""
        import subprocess
        cmd = f'"{self.ffprobe_path}" -v error -show_entries format=duration -of default=noprint_wrappers=1:nokey=1 "{audio_path}"'
        result = subprocess.run(cmd, shell=True, capture_output=True, text=True)
        return float(result.stdout.strip())

    def get_video_duration(self, video_path):
        """鑾峰彇瑙嗛鏃堕暱"""
        import subprocess
        cmd = f'"{self.ffprobe_path}" -v error -show_entries format=duration -of default=noprint_wrappers=1:nokey=1 "{video_path}"'
        result = subprocess.run(cmd, shell=True, capture_output=True, text=True)
        try:
            return float(result.stdout.strip())
        except:
            return 0

    def transcribe_audio_for_video(self, audio_path, temp_dir):
        """浣跨敤Whisper璇嗗埆闊抽鐢熸垚瀛楀箷"""
        try:
            from faster_whisper import WhisperModel

            self.video_log("姝ｅ湪鍔犺浇Whisper妯″瀷...")

            if self.whisper_model is None:
                self.whisper_model = WhisperModel("small", device="cpu", compute_type="int8")

            self.video_log("寮€濮嬭瘑鍒煶棰?..")
            segments, info = self.whisper_model.transcribe(audio_path, language="zh")

            srt_path = os.path.join(temp_dir, "subtitle.srt")
            with open(srt_path, 'w', encoding='utf-8') as f:
                for i, segment in enumerate(segments):
                    start_str = self.seconds_to_srt_time(segment.start)
                    end_str = self.seconds_to_srt_time(segment.end)
                    text = segment.text.strip()

                    f.write(f"{i+1}\n")
                    f.write(f"{start_str} --> {end_str}\n")
                    f.write(f"{text}\n\n")

                    if not self.video_is_running:
                        return None

            self.video_log("Whisper璇嗗埆瀹屾垚")
            return srt_path

        except ImportError:
            self.video_log("璀﹀憡锛氭湭瀹夎faster-whisper锛屽皢涓嶆坊鍔犲瓧骞?)
            return None
        except Exception as e:
            self.video_log(f"Whisper璇嗗埆閿欒: {str(e)}")
            return None

    def seconds_to_srt_time(self, seconds):
        """绉掓暟杞琒RT鏃堕棿鏍煎紡"""
        hours = int(seconds // 3600)
        minutes = int((seconds % 3600) // 60)
        secs = int(seconds % 60)
        millis = int((seconds % 1) * 1000)
        return f"{hours:02d}:{minutes:02d}:{secs:02d},{millis:03d}"

    def collect_videos_for_video(self, needed_duration):
        """鏀堕泦瑙嗛绱犳潗"""
        self.video_log(f"闇€瑕佹敹闆?{needed_duration:.1f} 绉掔殑瑙嗛绱犳潗...")

        collected_videos = []
        total_duration = 0

        subfolders = []
        if os.path.exists(self.video_source_path):
            for item in os.listdir(self.video_source_path):
                item_path = os.path.join(self.video_source_path, item)
                if os.path.isdir(item_path):
                    subfolders.append(item_path)

        if not subfolders:
            # 濡傛灉娌℃湁瀛愭枃浠跺す锛岀洿鎺ヤ粠鏍圭洰褰曞彇
            subfolders = [self.video_source_path]

        self.video_log(f"鎵惧埌 {len(subfolders)} 涓礌鏉愭枃浠跺す")

        folder_video_count = {folder: 0 for folder in subfolders}

        while total_duration < needed_duration:
            found_new = False

            for folder in subfolders:
                if folder_video_count[folder] >= self.max_videos_per_folder:
                    continue

                video_files = []
                try:
                    for f in os.listdir(folder):
                        if f.lower().endswith(('.mp4', '.mov', '.avi', '.mkv')):
                            video_files.append(os.path.join(folder, f))
                except:
                    continue

                random.shuffle(video_files)

                for vf in video_files:
                    if vf not in collected_videos:
                        duration = self.get_video_duration(vf)
                        if duration > 0:
                            collected_videos.append(vf)
                            total_duration += duration
                            folder_video_count[folder] += 1
                            self.video_log(f"  + {os.path.basename(vf)} ({duration:.1f}绉?")
                            found_new = True
                            break

                if total_duration >= needed_duration:
                    break

            if not found_new:
                self.video_log("璀﹀憡锛氱礌鏉愪笉瓒?)
                break

        self.video_log(f"鍏辨敹闆?{len(collected_videos)} 涓棰戯紝鎬绘椂闀?{total_duration:.1f} 绉?)
        return collected_videos

    def get_color_code(self, color_name):
        """棰滆壊鍚嶇О杞琭fmpeg棰滆壊浠ｇ爜"""
        colors = {"閲戣壊": "gold", "榛勮壊": "yellow", "姗欒壊": "orange", "绾㈣壊": "red"}
        return colors.get(color_name, "gold")

    def get_preview_color(self, color_name):
        """棰滆壊鍚嶇О杞琓kinter棰滆壊浠ｇ爜"""
        colors = {"閲戣壊": "#FFD700", "榛勮壊": "#FFFF00", "姗欒壊": "#FFA500", "绾㈣壊": "#FF0000"}
        return colors.get(color_name, "#FFD700")

    def update_subtitle_preview(self, event=None):
        """鏇存柊瀛楀箷棰勮"""
        if not hasattr(self, 'subtitle_preview_canvas'):
            return

        canvas = self.subtitle_preview_canvas
        canvas.delete("all")

        # 鑾峰彇褰撳墠璁剧疆
        font_size = int(self.subtitle_size.get())
        color = self.get_preview_color(self.subtitle_color.get())
        position = self.subtitle_position.get()

        # 棰勮鏂囧瓧
        preview_text = "杩欐槸瀛楀箷棰勮鏁堟灉"

        # 璁＄畻棰勮瀛楀彿锛堢缉灏忔瘮渚嬫樉绀猴級
        preview_font_size = max(12, font_size // 3)

        # 鏍规嵁浣嶇疆纭畾Y鍧愭爣
        canvas_height = 80
        if position == "搴曢儴":
            y = canvas_height - 15
        elif position == "涓儴":
            y = canvas_height // 2
        else:  # 椤堕儴
            y = 15

        # 缁樺埗鏂囧瓧锛堝甫榛戣壊鎻忚竟鏁堟灉锛?
        x = 200  # 灞呬腑
        # 鎻忚竟
        for dx, dy in [(-1,-1), (-1,1), (1,-1), (1,1), (-2,0), (2,0), (0,-2), (0,2)]:
            canvas.create_text(x+dx, y+dy, text=preview_text, fill="black",
                             font=("SimHei", preview_font_size, "bold"))
        # 涓绘枃瀛?
        canvas.create_text(x, y, text=preview_text, fill=color,
                         font=("SimHei", preview_font_size, "bold"))

    def get_ass_color(self, color_name):
        """棰滆壊鍚嶇О杞珹SS瀛楀箷棰滆壊浠ｇ爜锛圔GR鏍煎紡锛?""
        # ASS棰滆壊鏍煎紡: &HBBGGRR (钃濈豢绾?
        colors = {
            "閲戣壊": "&H0000D7FF",
            "榛勮壊": "&H0000FFFF",
            "姗欒壊": "&H0000A5FF",
            "绾㈣壊": "&H000000FF"
        }
        return colors.get(color_name, "&H0000D7FF")

    def burn_subtitles_for_video(self, video_path, srt_path, temp_dir, time_offset=0):
        """鐑у綍瀛楀箷

        Args:
            time_offset: 瀛楀箷鏃堕棿鍋忕Щ锛堢锛夛紝鐢ㄤ簬琛ュ伩灏侀潰鏃堕暱
        """
        import subprocess
        import shutil

        if not srt_path or not os.path.exists(srt_path):
            self.video_log("鏃犲瓧骞曟枃浠讹紝璺宠繃瀛楀箷鐑у綍")
            return video_path

        # 濡傛灉鏈夋椂闂村亸绉伙紝璋冩暣瀛楀箷鏃堕棿
        actual_srt_path = srt_path
        if time_offset > 0:
            actual_srt_path = os.path.join(temp_dir, "subtitle_offset.srt")
            self.offset_srt_time(srt_path, actual_srt_path, time_offset)
            self.video_log(f"瀛楀箷鏃堕棿宸插亸绉?{time_offset} 绉掞紙琛ュ伩灏侀潰鏃堕暱锛?)

        font_size = self.subtitle_size.get()
        font_color = self.get_ass_color(self.subtitle_color.get())
        position = self.subtitle_position.get()

        if position == "搴曢儴":
            margin_v = 30
            alignment = 2
        elif position == "涓儴":
            margin_v = 0
            alignment = 5
        else:
            margin_v = 30
            alignment = 8

        with_subtitle = os.path.join(temp_dir, "with_subtitle.mp4")

        # 澶嶅埗瀛楀箷鍒颁复鏃剁洰褰曪紝鐢ㄧ畝鍗曟枃浠跺悕閬垮厤璺緞闂
        simple_srt = os.path.join(temp_dir, "sub.srt")
        shutil.copy(actual_srt_path, simple_srt)
        srt_escaped = simple_srt.replace('\\', '/').replace(':', '\\:')

        subtitle_filter = f"subtitles='{srt_escaped}':force_style='FontSize={font_size},PrimaryColour={font_color},OutlineColour=&H00000000,BorderStyle=1,Outline=2,Alignment={alignment},MarginV={margin_v}'"

        cmd = f'"{self.ffmpeg_path}" -y -i "{video_path}" -vf "{subtitle_filter}" -c:a aac -b:a 192k "{with_subtitle}"'
        self.video_log("姝ｅ湪鐑у綍瀛楀箷...")
        result = subprocess.run(cmd, shell=True, capture_output=True, text=True)

        if result.returncode != 0:
            self.video_log(f"瀛楀箷鐑у綍澶辫触: {result.stderr[:200] if result.stderr else ''}")
            self.video_log("浣跨敤鏃犲瓧骞曠増鏈?)
            return video_path

        self.video_log("瀛楀箷鐑у綍鎴愬姛")
        return with_subtitle

    def offset_srt_time(self, input_srt, output_srt, offset_seconds):
        """鍋忕ЩSRT瀛楀箷鏃堕棿"""
        def parse_time(time_str):
            parts = time_str.replace(',', ':').split(':')
            h, m, s, ms = int(parts[0]), int(parts[1]), int(parts[2]), int(parts[3])
            return h * 3600 + m * 60 + s + ms / 1000

        def format_time(seconds):
            h = int(seconds // 3600)
            m = int((seconds % 3600) // 60)
            s = int(seconds % 60)
            ms = int((seconds % 1) * 1000)
            return f"{h:02d}:{m:02d}:{s:02d},{ms:03d}"

        with open(input_srt, 'r', encoding='utf-8') as f:
            content = f.read()

        pattern = r'(\d{2}:\d{2}:\d{2},\d{3}) --> (\d{2}:\d{2}:\d{2},\d{3})'
        def replace_time(match):
            start = parse_time(match.group(1)) + offset_seconds
            end = parse_time(match.group(2)) + offset_seconds
            return f"{format_time(start)} --> {format_time(end)}"

        new_content = re.sub(pattern, replace_time, content)
        with open(output_srt, 'w', encoding='utf-8') as f:
            f.write(new_content)

    def add_cover_for_video(self, video_path, video_title, temp_dir):
        """娣诲姞灏侀潰"""
        import subprocess

        cover_font_size = self.cover_font_size.get()
        cover_color = self.get_color_code(self.cover_color.get())
        cover_bg = self.cover_bg.get()

        probe_cmd = f'"{self.ffprobe_path}" -v error -select_streams v:0 -show_entries stream=width,height -of csv=p=0 "{video_path}"'
        probe_result = subprocess.run(probe_cmd, shell=True, capture_output=True, text=True)
        try:
            width, height = map(int, probe_result.stdout.strip().split(','))
        except:
            width, height = 1920, 1080

        cover_video = os.path.join(temp_dir, "cover.mp4")
        cover_duration = 1  # 灏侀潰鏃堕暱1绉?

        safe_title = video_title.replace("'", "").replace(":", "\\:").replace("\\", "/")
        if len(safe_title) > 12:
            mid = len(safe_title) // 2
            safe_title = safe_title[:mid] + "\\n" + safe_title[mid:]

        title_filter = f"drawtext=text='{safe_title}':fontfile='C\\:/Windows/Fonts/msyh.ttc':fontsize={cover_font_size}:fontcolor={cover_color}:borderw=4:bordercolor=black:x=(w-text_w)/2:y=(h-text_h)/2"

        if cover_bg == "瑙嗛棣栧抚":
            first_frame = os.path.join(temp_dir, "first_frame.jpg")
            cmd = f'"{self.ffmpeg_path}" -y -i "{video_path}" -vframes 1 -q:v 2 "{first_frame}"'
            subprocess.run(cmd, shell=True, capture_output=True)
            # 鐢熸垚甯﹂潤闊抽煶杞ㄧ殑灏侀潰瑙嗛
            cmd = f'"{self.ffmpeg_path}" -y -loop 1 -i "{first_frame}" -f lavfi -i anullsrc=r=44100:cl=stereo -vf "{title_filter}" -t {cover_duration} -r 30 -c:v libx264 -pix_fmt yuv420p -c:a aac -shortest "{cover_video}"'
        else:
            # 榛戣壊鑳屾櫙 + 闈欓煶闊宠建
            cmd = f'"{self.ffmpeg_path}" -y -f lavfi -i color=c=black:s={width}x{height}:d={cover_duration}:r=30 -f lavfi -i anullsrc=r=44100:cl=stereo -vf "{title_filter}" -c:v libx264 -pix_fmt yuv420p -c:a aac -t {cover_duration} -shortest "{cover_video}"'

        self.video_log("鐢熸垚灏侀潰瑙嗛...")
        result = subprocess.run(cmd, shell=True, capture_output=True, text=True)

        if result.returncode != 0:
            self.video_log(f"灏侀潰鐢熸垚澶辫触: {result.stderr[:200] if result.stderr else ''}")
            self.video_log("浣跨敤鏃犲皝闈㈢増鏈?)
            return video_path

        # 閲嶆柊缂栫爜涓昏棰戠‘淇濇牸寮忎竴鑷达紙淇濈暀闊抽锛?
        main_encoded = os.path.join(temp_dir, "main_encoded.mp4")
        cmd = f'"{self.ffmpeg_path}" -y -i "{video_path}" -c:v libx264 -preset fast -crf 23 -c:a aac -ar 44100 -b:a 192k "{main_encoded}"'
        self.video_log("閲嶆柊缂栫爜涓昏棰?..")
        subprocess.run(cmd, shell=True, capture_output=True)

        concat_list = os.path.join(temp_dir, "concat_list.txt")
        with open(concat_list, 'w', encoding='utf-8') as f:
            f.write(f"file '{cover_video.replace(chr(92), '/')}'\n")
            f.write(f"file '{main_encoded.replace(chr(92), '/')}'\n")

        final_video = os.path.join(temp_dir, "final.mp4")
        cmd = f'"{self.ffmpeg_path}" -y -f concat -safe 0 -i "{concat_list}" -c copy "{final_video}"'
        result = subprocess.run(cmd, shell=True, capture_output=True, text=True)

        if result.returncode != 0:
            self.video_log("灏侀潰鍚堝苟澶辫触锛屼娇鐢ㄦ棤灏侀潰鐗堟湰")
            return video_path

        self.video_log("灏侀潰鍚堝苟鎴愬姛")
        return final_video

    def finish_video_task(self):
        """瀹屾垚瑙嗛鍒朵綔浠诲姟"""
        self.video_is_running = False
        self.video_make_btn.config(state=tk.NORMAL)
        self.video_stop_btn.config(state=tk.DISABLED)
        self.video_update_status("瀹屾垚")

    def _create_model_config(self, parent, prefix):
        """鍒涘缓妯″瀷閰嶇疆UI缁勪欢"""
        # URL
        url_frame = ttk.Frame(parent)
        url_frame.pack(fill=tk.X, pady=2)
        ttk.Label(url_frame, text="URL:", width=12).pack(side=tk.LEFT)
        url_var = tk.StringVar(value=self.config.get(f"{prefix}_url", DEFAULT_CONFIG.get(f"{prefix}_url", "")))
        setattr(self, f"{prefix}_url", url_var)
        ttk.Entry(url_frame, textvariable=url_var, width=60).pack(side=tk.LEFT, fill=tk.X, expand=True)

        # Key
        key_frame = ttk.Frame(parent)
        key_frame.pack(fill=tk.X, pady=2)
        ttk.Label(key_frame, text="API Key:", width=12).pack(side=tk.LEFT)
        key_var = tk.StringVar(value=self.config.get(f"{prefix}_key", DEFAULT_CONFIG.get(f"{prefix}_key", "")))
        setattr(self, f"{prefix}_key", key_var)
        ttk.Entry(key_frame, textvariable=key_var, width=60).pack(side=tk.LEFT, fill=tk.X, expand=True)

        # Model + Max Tokens + 娴嬭瘯鎸夐挳
        model_frame = ttk.Frame(parent)
        model_frame.pack(fill=tk.X, pady=2)
        ttk.Label(model_frame, text="妯″瀷:", width=12).pack(side=tk.LEFT)
        model_var = tk.StringVar(value=self.config.get(f"{prefix}_model", DEFAULT_CONFIG.get(f"{prefix}_model", "")))
        setattr(self, f"{prefix}_model", model_var)
        ttk.Entry(model_frame, textvariable=model_var, width=35).pack(side=tk.LEFT)
        ttk.Label(model_frame, text="  Max Tokens:").pack(side=tk.LEFT)
        tokens_var = tk.StringVar(value=str(self.config.get(f"{prefix}_max_tokens", DEFAULT_CONFIG.get(f"{prefix}_max_tokens", 16000))))
        setattr(self, f"{prefix}_max_tokens", tokens_var)
        ttk.Entry(model_frame, textvariable=tokens_var, width=10).pack(side=tk.LEFT)

        # 娴嬭瘯鎸夐挳
        ttk.Button(model_frame, text="娴嬭瘯杩炴帴", width=10,
                   command=lambda p=prefix: self.test_api_connection(p)).pack(side=tk.LEFT, padx=10)

    def test_api_connection(self, prefix):
        """娴嬭瘯API杩炴帴锛岃嚜鍔ㄥ皾璇曚笉鍚屾牸寮?""
        url = getattr(self, f"{prefix}_url").get().strip()
        key = getattr(self, f"{prefix}_key").get().strip()
        model = getattr(self, f"{prefix}_model").get().strip()

        if not url or not key or not model:
            messagebox.showwarning("鎻愮ず", "璇峰厛濉啓瀹屾暣鐨刄RL銆丄PI Key鍜屾ā鍨嬪悕绉?)
            return

        # 鍦ㄦ柊绾跨▼涓祴璇曪紝閬垮厤鐣岄潰鍗′綇
        def do_test():
            self.log(f"寮€濮嬫祴璇?{prefix} 杩炴帴...")

            # 灏濊瘯涓嶅悓鐨凙PI鏍煎紡
            formats_to_try = [
                ("OpenAI鏍煎紡", self._test_openai_format),
                ("OpenAI鏍煎紡(鏃爒1)", self._test_openai_format_no_v1),
                ("Anthropic鏍煎紡", self._test_anthropic_format),
            ]

            for format_name, test_func in formats_to_try:
                self.log(f"  灏濊瘯 {format_name}...")
                success, message, working_url = test_func(url, key, model)
                if success:
                    self.log(f"  鉁?{format_name} 鎴愬姛锛?)
                    # 濡傛灉URL琚慨姝ｄ簡锛屾洿鏂伴厤缃?
                    if working_url and working_url != url:
                        getattr(self, f"{prefix}_url").set(working_url)
                        self.log(f"  宸茶嚜鍔ㄤ慨姝RL涓? {working_url}")
                    self.root.after(0, lambda: messagebox.showinfo("娴嬭瘯鎴愬姛",
                        f"杩炴帴鎴愬姛锛乗n\n鏍煎紡: {format_name}\nURL: {working_url or url}\n妯″瀷: {model}\n\n鍝嶅簲: {message}"))
                    return
                else:
                    self.log(f"  鉁?{format_name} 澶辫触: {message[:50]}...")

            self.log("鎵€鏈夋牸寮忓潎娴嬭瘯澶辫触")
            self.root.after(0, lambda: messagebox.showerror("娴嬭瘯澶辫触",
                f"鎵€鏈堿PI鏍煎紡鍧囨棤娉曡繛鎺n\n璇锋鏌?\n1. URL鏄惁姝ｇ‘\n2. API Key鏄惁鏈夋晥\n3. 妯″瀷鍚嶇О鏄惁姝ｇ‘"))

        import threading
        threading.Thread(target=do_test, daemon=True).start()

    def _test_openai_format(self, base_url, api_key, model):
        """娴嬭瘯OpenAI鍏煎鏍煎紡"""
        try:
            url = f"{base_url.rstrip('/')}/chat/completions"
            headers = {
                "Content-Type": "application/json",
                "Authorization": f"Bearer {api_key}"
            }
            data = {
                "model": model,
                "messages": [{"role": "user", "content": "浣犲ソ锛岃鍥炲锛氭祴璇曟垚鍔?}],
                "max_tokens": 50
            }
            response = requests.post(url, headers=headers, json=data, timeout=30)
            if response.status_code == 200:
                result = response.json()
                content = result.get("choices", [{}])[0].get("message", {}).get("content", "")
                return True, content, base_url
            else:
                return False, f"鐘舵€佺爜{response.status_code}: {response.text[:100]}", None
        except Exception as e:
            return False, str(e), None

    def _test_openai_format_no_v1(self, base_url, api_key, model):
        """娴嬭瘯OpenAI鏍煎紡锛岃嚜鍔ㄦ坊鍔?v1"""
        # 濡傛灉URL宸茬粡鍖呭惈/v1锛岃烦杩?
        if "/v1" in base_url:
            return False, "URL宸插寘鍚?v1锛岃烦杩囨娴嬭瘯", None

        new_url = f"{base_url.rstrip('/')}/v1"
        success, message, _ = self._test_openai_format(new_url, api_key, model)
        if success:
            return True, message, new_url
        return False, message, None

    def _test_anthropic_format(self, base_url, api_key, model):
        """娴嬭瘯Anthropic鍘熺敓鏍煎紡"""
        try:
            # 灏濊瘯Anthropic鏍煎紡鐨別ndpoint
            url = f"{base_url.rstrip('/')}/messages"
            if "/v1" not in url:
                url = f"{base_url.rstrip('/')}/v1/messages"

            headers = {
                "Content-Type": "application/json",
                "x-api-key": api_key,
                "anthropic-version": "2023-06-01"
            }
            data = {
                "model": model,
                "messages": [{"role": "user", "content": "浣犲ソ锛岃鍥炲锛氭祴璇曟垚鍔?}],
                "max_tokens": 50
            }
            response = requests.post(url, headers=headers, json=data, timeout=30)
            if response.status_code == 200:
                result = response.json()
                content = result.get("content", [{}])[0].get("text", "")
                return True, content, base_url
            else:
                return False, f"鐘舵€佺爜{response.status_code}: {response.text[:100]}", None
        except Exception as e:
            return False, str(e), None

    def on_api_tab_change(self, event=None):
        """API鏍囩椤靛垏鎹?- 鑷姩鏇存柊娴佸紡/闈炴祦寮忚缃?""
        selected_tab = self.api_notebook.index(self.api_notebook.select())
        # 0 = 娴佸紡, 1 = 闈炴祦寮?
        use_stream = (selected_tab == 0)
        self.use_stream.set(use_stream)
        self.config["use_stream"] = use_stream
        save_config(self.config)
        mode = "娴佸紡" if use_stream else "闈炴祦寮?
        self.log(f"宸插垏鎹㈠埌{mode}璋冪敤妯″紡")
        # 鏇存柊妯″紡鎻愮ず
        if hasattr(self, 'api_mode_label'):
            self._update_api_mode_label()

    def on_stream_change(self):
        """娴佸紡寮€鍏虫敼鍙?- 鑷姩淇濆瓨锛堜繚鐣欏吋瀹癸級"""
        self.config["use_stream"] = self.use_stream.get()
        save_config(self.config)
        mode = "娴佸紡" if self.use_stream.get() else "闈炴祦寮?
        self.log(f"宸插垏鎹㈠埌{mode}璋冪敤妯″紡")

    def _parse_similarity_threshold(self):
        """瑙ｆ瀽骞舵牎楠岀浉浼煎害闃堝€?""
        raw_value = self.similarity_threshold_var.get().strip()
        value = float(raw_value)
        if value < 0.50 or value > 0.90:
            raise ValueError("鐩镐技搴﹂槇鍊煎繀椤诲湪 0.50 ~ 0.90 涔嬮棿")
        return round(value, 2)

    def save_api_config(self):
        """淇濆瓨API閰嶇疆"""
        try:
            self.config["use_stream"] = self.use_stream.get()
            threshold_value = self._parse_similarity_threshold()
            self.config["similarity_threshold"] = threshold_value
            self.similarity_threshold = threshold_value

            # 杈呭姪鍑芥暟锛氬畨鍏ㄨ浆鎹ax_tokens
            def safe_int(value, default=16000):
                try:
                    v = value.strip()
                    if not v:
                        return default
                    return int(v)
                except (ValueError, AttributeError):
                    return default

            # 娴佸紡閰嶇疆
            self.config["stream_main_url"] = self.stream_main_url.get().strip()
            self.config["stream_main_key"] = self.stream_main_key.get().strip()
            self.config["stream_main_model"] = self.stream_main_model.get().strip()
            self.config["stream_main_max_tokens"] = safe_int(self.stream_main_max_tokens.get())
            self.config["stream_backup_url"] = self.stream_backup_url.get().strip()
            self.config["stream_backup_key"] = self.stream_backup_key.get().strip()
            self.config["stream_backup_model"] = self.stream_backup_model.get().strip()
            self.config["stream_backup_max_tokens"] = safe_int(self.stream_backup_max_tokens.get())
            # 闈炴祦寮忛厤缃?
            self.config["non_stream_main_url"] = self.non_stream_main_url.get().strip()
            self.config["non_stream_main_key"] = self.non_stream_main_key.get().strip()
            self.config["non_stream_main_model"] = self.non_stream_main_model.get().strip()
            self.config["non_stream_main_max_tokens"] = safe_int(self.non_stream_main_max_tokens.get())
            self.config["non_stream_backup_url"] = self.non_stream_backup_url.get().strip()
            self.config["non_stream_backup_key"] = self.non_stream_backup_key.get().strip()
            self.config["non_stream_backup_model"] = self.non_stream_backup_model.get().strip()
            self.config["non_stream_backup_max_tokens"] = safe_int(self.non_stream_backup_max_tokens.get())
            save_config(self.config)
            messagebox.showinfo("鎴愬姛", "API閰嶇疆宸蹭繚瀛?)
        except Exception as e:
            messagebox.showerror("閿欒", f"淇濆瓨閰嶇疆澶辫触锛歿e}")

    def reset_api_config(self):
        """閲嶇疆涓洪粯璁ら厤缃?""
        if messagebox.askyesno("纭", "纭畾瑕侀噸缃负榛樿閰嶇疆鍚楋紵"):
            self.config = DEFAULT_CONFIG.copy()
            self.use_stream.set(self.config["use_stream"])
            # 娴佸紡閰嶇疆
            self.stream_main_url.set(self.config["stream_main_url"])
            self.stream_main_key.set(self.config["stream_main_key"])
            self.stream_main_model.set(self.config["stream_main_model"])
            self.stream_main_max_tokens.set(str(self.config["stream_main_max_tokens"]))
            self.stream_backup_url.set(self.config["stream_backup_url"])
            self.stream_backup_key.set(self.config["stream_backup_key"])
            self.stream_backup_model.set(self.config["stream_backup_model"])
            self.stream_backup_max_tokens.set(str(self.config["stream_backup_max_tokens"]))
            # 闈炴祦寮忛厤缃?
            self.non_stream_main_url.set(self.config["non_stream_main_url"])
            self.non_stream_main_key.set(self.config["non_stream_main_key"])
            self.non_stream_main_model.set(self.config["non_stream_main_model"])
            self.non_stream_main_max_tokens.set(str(self.config["non_stream_main_max_tokens"]))
            self.non_stream_backup_url.set(self.config["non_stream_backup_url"])
            self.non_stream_backup_key.set(self.config["non_stream_backup_key"])
            self.non_stream_backup_model.set(self.config["non_stream_backup_model"])
            self.non_stream_backup_max_tokens.set(str(self.config["non_stream_backup_max_tokens"]))
            self.similarity_threshold_var.set(f"{float(self.config.get('similarity_threshold', 0.76)):.2f}")
            self.similarity_threshold = float(self.config.get("similarity_threshold", 0.76))
            save_config(self.config)
            messagebox.showinfo("鎴愬姛", "宸查噸缃负榛樿閰嶇疆")

    def on_flow_type_change(self, event=None):
        """寮曟祦绫诲瀷鏀瑰彉鏃剁殑澶勭悊"""
        flow_type = self.flow_type.get()

        if flow_type == "甯﹁揣寮曟祦":
            self.daihuo_frame.pack(fill=tk.X, pady=5, before=self.btn_frame_container)
        else:
            self.daihuo_frame.pack_forget()

        if flow_type == "绾じ璧炰笉寮曟祦":
            self.yinliu_frame.pack_forget()
        else:
            self.yinliu_frame.pack(fill=tk.X, pady=5, before=self.daihuo_frame if flow_type == "甯﹁揣寮曟祦" else self.btn_frame_container)

        # 鍒囨崲寮曟祦绫诲瀷鏃舵洿鏂拌瘽鏈笅鎷夋
        self.update_yinliu_combo()

    def update_yinliu_combo(self):
        """鏇存柊璇濇湳涓嬫媺妗?""
        flow_type = self.flow_type.get()
        if flow_type == "绾じ璧炰笉寮曟祦":
            return

        # 鑾峰彇褰撳墠绫诲瀷鐨勮瘽鏈垪琛?
        templates = self.config.get("yinliu_templates", {}).get(flow_type, [])

        # 鏋勫缓涓嬫媺妗嗛€夐」锛堟樉绀哄墠20涓瓧锛?
        options = ["-- 鑷畾涔夛紙鍦ㄤ笅鏂硅緭鍏ユ濉啓锛?-"]
        for i, tpl in enumerate(templates):
            preview = tpl[:20].replace('\n', ' ') + "..." if len(tpl) > 20 else tpl.replace('\n', ' ')
            options.append(f"{i+1}. {preview}")

        self.yinliu_combo['values'] = options
        self.yinliu_combo.current(0)

    def on_yinliu_select(self, event=None):
        """閫夋嫨璇濇湳鏃跺～鍏呭埌鏂囨湰妗?""
        flow_type = self.flow_type.get()
        selection = self.yinliu_combo.current()

        if selection <= 0:  # 閫夋嫨浜?涓嶄娇鐢?
            return

        # 鑾峰彇瀵瑰簲鐨勮瘽鏈唴瀹?
        templates = self.config.get("yinliu_templates", {}).get(flow_type, [])
        if selection - 1 < len(templates):
            template = templates[selection - 1]
            # 濉厖鍒版枃鏈
            self.yinliu_text.delete("1.0", tk.END)
            self.yinliu_text.insert("1.0", template)

    def save_yinliu_template(self):
        """淇濆瓨褰撳墠璇濇湳鍒板垪琛?""
        flow_type = self.flow_type.get()
        if flow_type == "绾じ璧炰笉寮曟祦":
            messagebox.showwarning("鎻愮ず", "绾じ璧炴ā寮忎笉闇€瑕佸紩娴佽瘽鏈?)
            return

        content = self.yinliu_text.get("1.0", tk.END).strip()
        if not content:
            messagebox.showwarning("鎻愮ず", "璇峰厛杈撳叆璇濇湳鍐呭")
            return

        # 纭繚閰嶇疆涓湁璇濇湳搴?
        if "yinliu_templates" not in self.config:
            self.config["yinliu_templates"] = {"缃《寮曟祦": [], "姗辩獥寮曟祦": [], "甯﹁揣寮曟祦": []}
        if flow_type not in self.config["yinliu_templates"]:
            self.config["yinliu_templates"][flow_type] = []

        # 妫€鏌ユ槸鍚﹀凡瀛樺湪
        templates = self.config["yinliu_templates"][flow_type]
        if content in templates:
            messagebox.showinfo("鎻愮ず", "璇ヨ瘽鏈凡瀛樺湪")
            return

        # 娣诲姞鍒板垪琛?
        templates.append(content)
        save_config(self.config)

        # 鏇存柊涓嬫媺妗?
        self.update_yinliu_combo()
        messagebox.showinfo("鎴愬姛", f"璇濇湳宸蹭繚瀛樺埌銆恵flow_type}銆戝垪琛?)

    def delete_yinliu_template(self):
        """鍒犻櫎閫変腑鐨勮瘽鏈?""
        flow_type = self.flow_type.get()
        selection = self.yinliu_combo.current()

        if selection <= 0:
            messagebox.showwarning("鎻愮ず", "璇峰厛閫夋嫨瑕佸垹闄ょ殑璇濇湳")
            return

        templates = self.config.get("yinliu_templates", {}).get(flow_type, [])
        if selection - 1 < len(templates):
            if messagebox.askyesno("纭", "纭畾瑕佸垹闄ら€変腑鐨勮瘽鏈悧锛?):
                del templates[selection - 1]
                save_config(self.config)
                self.update_yinliu_combo()
                self.yinliu_text.delete("1.0", tk.END)
                messagebox.showinfo("鎴愬姛", "璇濇湳宸插垹闄?)

    def clear_yinliu_text(self):
        """娓呯┖璇濇湳杈撳叆妗?""
        self.yinliu_text.delete("1.0", tk.END)
        self.yinliu_combo.current(0)

    def on_input_mode_change(self):
        """鍒囨崲鍙傝€冩枃妗堣緭鍏ユ柟寮?""
        mode = self.input_mode.get()
        if mode == "file":
            self.paste_input_frame.pack_forget()
            self.file_input_frame.pack(fill=tk.X)
        else:
            self.file_input_frame.pack_forget()
            self.paste_input_frame.pack(fill=tk.X)

    def select_input_file(self):
        file_path = filedialog.askopenfilename(
            title="閫夋嫨鍙傝€冩枃妗堟枃浠?,
            filetypes=[("鏂囨湰鏂囦欢", "*.txt"), ("鎵€鏈夋枃浠?, "*.*")]
        )
        if file_path:
            self.input_path.set(file_path)

    def select_input_folder(self):
        folder_path = filedialog.askdirectory(title="閫夋嫨鍙傝€冩枃妗堟枃浠跺す")
        if folder_path:
            self.input_path.set(folder_path)

    def select_output_folder(self):
        folder_path = filedialog.askdirectory(title="閫夋嫨杈撳嚭淇濆瓨鏂囦欢澶?)
        if folder_path:
            self.output_path.set(folder_path)

    def select_txt_output_folder(self):
        """閫夋嫨TXT淇濆瓨鏂囦欢澶?""
        current_path = self.txt_output_path.get()
        initial_dir = current_path if current_path and os.path.exists(current_path) else None
        folder_path = filedialog.askdirectory(title="閫夋嫨TXT淇濆瓨鏂囦欢澶?, initialdir=initial_dir)
        if folder_path:
            self.txt_output_path.set(folder_path)

    def select_voice_input_folder(self):
        """閫夋嫨璇煶鍚堟垚鐨勬枃妗堣緭鍏ョ洰褰?""
        current_path = self.voice_input_path.get()
        initial_dir = current_path if current_path and os.path.exists(current_path) else None
        folder_path = filedialog.askdirectory(title="閫夋嫨鏂囨鐩綍", initialdir=initial_dir)
        if folder_path:
            self.voice_input_path.set(folder_path)
            # 淇濆瓨鍒伴厤缃枃浠?
            self.config["voice_input_path"] = folder_path
            save_config(self.config)

    def select_voice_output_folder(self):
        """閫夋嫨璇煶鍚堟垚鐨勮緭鍑虹洰褰?""
        current_path = self.voice_output_path.get()
        initial_dir = current_path if current_path and os.path.exists(current_path) else None
        folder_path = filedialog.askdirectory(title="閫夋嫨閰嶉煶杈撳嚭鐩綍", initialdir=initial_dir)
        if folder_path:
            self.voice_output_path.set(folder_path)
            # 淇濆瓨鍒伴厤缃枃浠?
            self.config["voice_output_path"] = folder_path
            save_config(self.config)

    def log(self, message):
        """娣诲姞鏃ュ織"""
        timestamp = datetime.now().strftime("%H:%M:%S")
        log_line = f"[{timestamp}] {message}\n"
        self.log_text.insert(tk.END, log_line)
        self.log_text.see(tk.END)
        self.root.update()

        try:
            log_file = os.path.join(self.output_path.get() or ".", "debug_log.txt")
            with open(log_file, 'a', encoding='utf-8') as f:
                f.write(log_line)
        except:
            pass

    def update_status(self, message):
        """鏇存柊鐘舵€?""
        self.status_label.config(text=message)
        self.root.update()

    def update_progress(self, value):
        """鏇存柊杩涘害鏉?""
        self.progress_var.set(value)
        self.root.update()

    def start_generate(self):
        """寮€濮嬬敓鎴?""
        input_mode = self.input_mode.get()
        output_path = self.output_path.get()

        # 鏍规嵁杈撳叆妯″紡楠岃瘉
        if input_mode == "paste":
            paste_content = self.paste_text.get("1.0", tk.END).strip()
            if not paste_content:
                messagebox.showerror("閿欒", "璇风矘璐村弬鑰冩枃妗堝唴瀹?)
                return
        else:
            input_path = self.input_path.get()
            if not input_path:
                messagebox.showerror("閿欒", "璇烽€夋嫨鍙傝€冩枃妗堣矾寰?)
                return
            if not os.path.exists(input_path):
                messagebox.showerror("閿欒", f"鍙傝€冩枃妗堣矾寰勪笉瀛樺湪锛歿input_path}")
                return

        if not output_path:
            messagebox.showerror("閿欒", "璇烽€夋嫨杈撳嚭淇濆瓨璺緞")
            return

        if self.flow_type.get() == "甯﹁揣寮曟祦":
            if not self.product_name.get().strip():
                messagebox.showerror("閿欒", "璇峰～鍐欏甫璐у晢鍝佸悕绉?)
                return
            if not self.product_material.get("1.0", tk.END).strip():
                messagebox.showerror("閿欒", "璇峰～鍐欎骇鍝佺礌鏉?浠嬬粛")
                return

        os.makedirs(output_path, exist_ok=True)

        self.start_btn.config(state=tk.DISABLED)
        self.stop_btn.config(state=tk.NORMAL)
        self.is_running = True

        self.log_text.delete(1.0, tk.END)

        thread = threading.Thread(target=self.generate_task)
        thread.daemon = True
        thread.start()

    def start_generate_txt(self):
        """寮€濮嬬敓鎴愬苟淇濆瓨涓篢XT锛堟瘡绡囧崟鐙繚瀛橈紝鐢ㄦ爣棰樺懡鍚嶏級"""
        input_mode = self.input_mode.get()
        txt_output_path = self.txt_output_path.get()

        # 鏍规嵁杈撳叆妯″紡楠岃瘉
        if input_mode == "paste":
            paste_content = self.paste_text.get("1.0", tk.END).strip()
            if not paste_content:
                messagebox.showerror("閿欒", "璇风矘璐村弬鑰冩枃妗堝唴瀹?)
                return
        else:
            input_path = self.input_path.get()
            if not input_path:
                messagebox.showerror("閿欒", "璇烽€夋嫨鍙傝€冩枃妗堣矾寰?)
                return
            if not os.path.exists(input_path):
                messagebox.showerror("閿欒", f"鍙傝€冩枃妗堣矾寰勪笉瀛樺湪锛歿input_path}")
                return

        if not txt_output_path:
            messagebox.showerror("閿欒", "璇烽€夋嫨TXT淇濆瓨璺緞")
            return

        if self.flow_type.get() == "甯﹁揣寮曟祦":
            if not self.product_name.get().strip():
                messagebox.showerror("閿欒", "璇峰～鍐欏甫璐у晢鍝佸悕绉?)
                return
            if not self.product_material.get("1.0", tk.END).strip():
                messagebox.showerror("閿欒", "璇峰～鍐欎骇鍝佺礌鏉?浠嬬粛")
                return

        os.makedirs(txt_output_path, exist_ok=True)

        self.start_btn.config(state=tk.DISABLED)
        self.start_txt_btn.config(state=tk.DISABLED)
        self.stop_btn.config(state=tk.NORMAL)
        self.is_running = True
        self.txt_generate_success = False  # 鍒濆鍖栦负澶辫触鐘舵€?

        self.log_text.delete(1.0, tk.END)

        thread = threading.Thread(target=self.generate_txt_task)
        thread.daemon = True
        thread.start()

    def generate_txt_task(self):
        """鐢熸垚TXT浠诲姟涓诲嚱鏁?""
        try:
            txt_output_path = self.txt_output_path.get()
            flow_type = self.flow_type.get()

            # 璇诲彇鍙傝€冩枃妗?
            self.log("姝ｅ湪璇诲彇鍙傝€冩枃妗?..")
            input_mode = self.input_mode.get()

            if input_mode == "paste":
                paste_content = self.paste_text.get("1.0", tk.END).strip()
                if not paste_content:
                    self.log("閿欒锛氳绮樿创鍙傝€冩枃妗堝唴瀹?)
                    self.finish_txt_task()
                    return
                files_content = [("绮樿创鏂囨", paste_content)]
            else:
                input_path = self.input_path.get()
                if os.path.isfile(input_path):
                    with open(input_path, 'r', encoding='utf-8') as f:
                        content = f.read()
                    files_content = [(os.path.basename(input_path), content)]
                else:
                    files_content = []
                    for fname in os.listdir(input_path):
                        if fname.endswith('.txt'):
                            fpath = os.path.join(input_path, fname)
                            with open(fpath, 'r', encoding='utf-8') as f:
                                files_content.append((fname, f.read()))

            if not files_content:
                self.log("閿欒锛氭病鏈夋壘鍒颁换浣曟枃妗堟枃浠?)
                self.finish_txt_task()
                return

            # 璇诲彇寮曟祦绱犳潗
            yinliu_content = self.yinliu_text.get("1.0", tk.END).strip()

            # 甯﹁揣淇℃伅
            product_name = ""
            product_material = ""
            if flow_type == "甯﹁揣寮曟祦":
                product_name = self.product_name.get().strip()
                product_material = self.product_material.get("1.0", tk.END).strip()

            total_files = len(files_content)
            for idx, (fname, content) in enumerate(files_content):
                if not self.is_running:
                    break

                self.log(f"\n{'='*50}")
                self.log(f"澶勭悊鏂囦欢 [{idx+1}/{total_files}]: {fname}")
                self.update_progress((idx / total_files) * 100)

                # 鍒嗗壊鍙傝€冩枃妗?
                articles = self.parse_articles(content)
                self.log(f"璇嗗埆鍒?{len(articles)} 绡囧弬鑰冩枃妗?)

                # 鎶婂弬鑰冩枃妗堣拷鍔犲埌绱犳潗搴?
                for article in articles:
                    self.append_reference_to_library(article, flow_type)

                # 淇濆瓨鏈€鍚庝竴娆＄敓鎴愮殑淇℃伅
                self.last_articles = articles
                self.last_flow_type = flow_type
                self.last_yinliu_content = yinliu_content
                self.last_product_name = product_name
                self.last_product_material = product_material

                # 鐢ㄤ簬璁板綍宸蹭娇鐢ㄧ殑鏍囬锛岀‘淇?绡囦笉閲嶅
                used_titles = []
                # 璁板綍鎴愬姛鍜屽け璐ョ殑绡囨暟
                success_count = 0
                failed_list = []

                for art_idx, article in enumerate(articles):
                    if not self.is_running:
                        break

                    self.log(f"\n--- 澶勭悊绗?{art_idx+1} 绡囧弬鑰冩枃妗?---")
                    self.update_status(f"姝ｅ湪鐢熸垚绗?{art_idx+1} 绡?..")

                    # 鐢熸垚浠垮啓鏂囨
                    result = self.generate_document(
                        article, flow_type, yinliu_content,
                        product_name, product_material
                    )

                    if result:
                        # 淇濆瓨涓篢XT锛岀敤鏍囬鍛藉悕
                        self.save_as_txt(result, txt_output_path, used_titles)
                        self.txt_generate_success = True
                        success_count += 1
                    else:
                        self.log(f"绗?{art_idx+1} 绡囩敓鎴愬け璐?)
                        failed_list.append(art_idx + 1)

                # 杈撳嚭姹囨€荤粨鏋?
                self.log(f"\n{'='*50}")
                self.log(f"銆愮敓鎴愮粨鏋滄眹鎬汇€?)
                self.log(f"鎴愬姛: {success_count} 绡?)
                if failed_list:
                    self.log(f"澶辫触: {len(failed_list)} 绡囷紝鍒嗗埆鏄 {', '.join(map(str, failed_list))} 绡?)
                else:
                    self.log(f"澶辫触: 0 绡?)

            self.update_progress(100)
            self.log("\n" + "="*50)
            self.log("鍏ㄩ儴澶勭悊瀹屾垚锛?)
            self.finish_txt_task()

        except Exception as e:
            self.log(f"浠诲姟鍑洪敊: {str(e)}")
            import traceback
            self.log(traceback.format_exc())
            self.txt_generate_success = False
            self.finish_txt_task()

    def split_articles_from_result(self, content):
        """鎶婄敓鎴愮殑澶ф枃妗堟媶鍒嗘垚澶氱瘒锛岃繑鍥炲垪琛?[(绾鏂? [鏍囬鍒楄〃]), ...]"""
        articles = []
        # 鎸夈€愮X绡囥€戝垎鍓诧紝鏀寔澶氱鏍煎紡
        pattern = r'鈺?\s*\n*銆愮[涓€浜屼笁鍥涗簲鍏竷鍏節鍗乗d]+绡囥€慭s*\n*鈺?'
        parts = re.split(pattern, content)

        for part in parts:
            part = part.strip()
            if not part or len(part) < 100:
                continue

            # 鎻愬彇璇ョ瘒鐨勬爣棰樺拰姝ｆ枃
            titles = []
            lines = part.split('\n')
            in_title_section = False
            content_start_idx = 0

            for i, line in enumerate(lines):
                line_stripped = line.strip()

                # 鏂瑰紡1: 銆愭爣棰樸€戝悗闈㈣窡鐫€澶氳鏍囬
                if line_stripped.startswith('銆愭爣棰樸€?) or line_stripped == '銆愭爣棰樸€?:
                    in_title_section = True
                    after_tag = line_stripped.replace('銆愭爣棰樸€?, '').strip()
                    if after_tag:
                        titles.append(after_tag)
                    continue

                # 鏂瑰紡2: 銆愭爣棰?銆戙€愭爣棰?銆戠瓑鏍煎紡锛堝鐢ㄦā鍨嬪父鐢級
                title_match = re.match(r'銆愭爣棰榌1-5涓€浜屼笁鍥涗簲]銆慬锛?]?\s*(.+)', line_stripped)
                if title_match:
                    title_content = title_match.group(1).strip()
                    if title_content and len(title_content) > 2:
                        titles.append(title_content)
                    if not in_title_section:
                        in_title_section = True
                    continue

                # 鏂瑰紡3: 鏍囬1锛歺xx 鎴?鏍囬涓€锛歺xx 鏍煎紡
                title_match2 = re.match(r'鏍囬[1-5涓€浜屼笁鍥涗簲][锛?]\s*(.+)', line_stripped)
                if title_match2:
                    title_content = title_match2.group(1).strip()
                    if title_content and len(title_content) > 2:
                        titles.append(title_content)
                    if not in_title_section:
                        in_title_section = True
                    continue

                if in_title_section:
                    if line_stripped == '---' or line_stripped.startswith('---'):
                        # 鎵惧埌鍒嗛殧绾匡紝姝ｆ枃浠庝笅涓€琛屽紑濮?
                        in_title_section = False
                        content_start_idx = i + 1
                    elif line_stripped and not line_stripped.startswith('鈺?) and len(titles) < 5:
                        # 鏅€氭爣棰樿锛堟棤鍓嶇紑锛?
                        clean_title = re.sub(r'^[\d]+[.銆乗s]*', '', line_stripped).strip()
                        if clean_title and len(clean_title) > 2 and not clean_title.startswith('銆?):
                            titles.append(clean_title)

            # 鎻愬彇绾鏂囷紙---涔嬪悗鐨勫唴瀹癸級
            if content_start_idx > 0:
                body_lines = lines[content_start_idx:]
                # 鍘绘帀寮€澶寸殑绌鸿
                while body_lines and not body_lines[0].strip():
                    body_lines.pop(0)
                body_content = '\n'.join(body_lines).strip()
            else:
                # 濡傛灉娌℃壘鍒?--锛屽皾璇曞幓鎺夋爣棰橀儴鍒嗭紝鎵剧涓€涓潪鏍囬琛?
                body_lines = []
                found_body = False
                for line in lines:
                    line_stripped = line.strip()
                    # 璺宠繃鏍囬鐩稿叧鐨勮
                    if re.match(r'銆愭爣棰榌1-5涓€浜屼笁鍥涗簲]?銆?, line_stripped):
                        continue
                    if re.match(r'鏍囬[1-5涓€浜屼笁鍥涗簲][锛?]', line_stripped):
                        continue
                    if line_stripped == '---':
                        found_body = True
                        continue
                    if found_body or (line_stripped and len(line_stripped) > 30):
                        found_body = True
                        body_lines.append(line)
                body_content = '\n'.join(body_lines).strip() if body_lines else part

            articles.append((body_content, titles))

        return articles

    def save_as_txt(self, content, output_path, used_titles):
        """鎶婄敓鎴愮殑鏂囨鎷嗗垎鎴愬绡囷紝姣忕瘒鍗曠嫭淇濆瓨涓篢XT锛岀敤鏍囬鍛藉悕"""
        try:
            # 鎷嗗垎鎴愬绡?
            articles = self.split_articles_from_result(content)

            if not articles:
                self.log("璀﹀憡锛氭湭鑳芥媶鍒嗗嚭鏂囩珷锛屼繚瀛樻暣涓唴瀹?)
                timestamp = datetime.now().strftime("%Y%m%d%H%M%S")
                filename = f"{timestamp}.txt"
                filepath = os.path.join(output_path, filename)
                with open(filepath, 'w', encoding='utf-8') as f:
                    f.write(content)
                self.log(f"TXT宸蹭繚瀛? {filename}")
                return [filepath]

            saved_files = []
            self.log(f"璇嗗埆鍒?{len(articles)} 绡囨枃绔狅紝寮€濮嬪垎鍒繚瀛?..")
            existing_titles = self.get_generated_title_set()

            for idx, (article_content, titles) in enumerate(articles):
                selected_title = ""
                if not titles:
                    self.log(f"绗瑊idx+1}绡囨湭鎻愬彇鍒版爣棰橈紝浣跨敤鏃堕棿鎴冲懡鍚?)
                    timestamp = datetime.now().strftime("%Y%m%d%H%M%S")
                    filename = f"{timestamp}_{idx+1}.txt"
                    selected_title = f"鏈彁鍙栨爣棰榑{timestamp}_{idx+1}"
                else:
                    # 杈撳嚭璇ョ瘒鎵€鏈夋爣棰?
                    self.log(f"绗瑊idx+1}绡囨彁鍙栧埌 {len(titles)} 涓爣棰橈細")
                    for t_idx, t in enumerate(titles):
                        self.log(f"  鏍囬{t_idx+1}: {t}")

                    # 浠庢爣棰樹腑闅忔満閫変竴涓紝閬垮厤涓庢湰娆″拰鍘嗗彶宸茬敤鏍囬閲嶅
                    available_titles = [t for t in titles if t not in used_titles and t not in existing_titles]
                    if not available_titles:
                        available_titles = [t for t in titles if t not in used_titles]
                    if not available_titles:
                        available_titles = titles

                    selected_title = random.choice(available_titles)
                    self.log(f"  鈫?闅忔満閫変腑: {selected_title}")
                    used_titles.append(selected_title)
                    existing_titles.add(selected_title)

                    # 娓呯悊鏂囦欢鍚嶄腑鐨勯潪娉曞瓧绗?
                    safe_title = re.sub(r'[\\/:*?"<>|]', '', selected_title)
                    safe_title = safe_title[:50]
                    filename = f"{safe_title}.txt"

                filepath = os.path.join(output_path, filename)

                # 濡傛灉鏂囦欢宸插瓨鍦紝娣诲姞鏃堕棿鎴?
                if os.path.exists(filepath):
                    timestamp = datetime.now().strftime("%H%M%S")
                    name, ext = os.path.splitext(filename)
                    filename = f"{name}_{timestamp}{ext}"
                    filepath = os.path.join(output_path, filename)

                with open(filepath, 'w', encoding='utf-8') as f:
                    f.write(article_content)

                # 鐢熸垚鏂囨鍏ュ簱Excel锛屼究浜庨暱鏈熷幓閲嶏紙鍗充娇txt琚Щ鍔級
                flow_type = self.flow_type.get() if hasattr(self, "flow_type") else ""
                self.append_generated_to_library(flow_type, selected_title or filename, article_content)

                self.log(f"绗瑊idx+1}绡嘥XT宸蹭繚瀛? {filename}")
                saved_files.append(filepath)

            return saved_files

        except Exception as e:
            self.log(f"淇濆瓨TXT澶辫触: {str(e)}")
            import traceback
            self.log(traceback.format_exc())
            return None

    def finish_txt_task(self):
        """瀹屾垚TXT鐢熸垚浠诲姟锛屾仮澶峌I鐘舵€?""
        self.is_running = False
        self.start_btn.config(state=tk.NORMAL)
        self.start_txt_btn.config(state=tk.NORMAL)
        self.stop_btn.config(state=tk.DISABLED)
        if self.last_articles:
            self.regenerate_btn.config(state=tk.NORMAL)
        self.update_status("澶勭悊瀹屾垚")

        # 鍙湁鐢熸垚鎴愬姛鏃舵墠璇㈤棶鏄惁鎵撳紑杈撳嚭鏂囦欢澶?
        if self.txt_generate_success:
            if messagebox.askyesno("瀹屾垚", "TXT鐢熸垚瀹屾垚锛佹槸鍚︽墦寮€TXT淇濆瓨鏂囦欢澶癸紵"):
                txt_path = self.txt_output_path.get().replace('/', '\\')
                if os.path.exists(txt_path):
                    os.startfile(txt_path)
        else:
            messagebox.showwarning("鎻愮ず", "鐢熸垚澶辫触锛岃鏌ョ湅鏃ュ織浜嗚В璇︽儏")

    def stop_generate(self):
        """鍋滄鐢熸垚"""
        self.is_running = False
        self.log("鐢ㄦ埛鍋滄浜嗙敓鎴愪换鍔?)
        self.update_status("宸插仠姝?)
        self.start_btn.config(state=tk.NORMAL)
        self.stop_btn.config(state=tk.DISABLED)

    def regenerate(self):
        """閲嶆柊鐢熸垚 - 浣跨敤涓婃鐨勫弬鑰冩枃妗堝拰褰撳墠璁剧疆"""
        if not self.last_articles:
            messagebox.showwarning("鎻愮ず", "娌℃湁鍙噸鏂扮敓鎴愮殑鏂囨锛岃鍏堟墽琛屼竴娆＄敓鎴?)
            return

        output_path = self.output_path.get()
        if not output_path:
            messagebox.showerror("閿欒", "璇烽€夋嫨杈撳嚭淇濆瓨璺緞")
            return

        os.makedirs(output_path, exist_ok=True)

        self.start_btn.config(state=tk.DISABLED)
        self.stop_btn.config(state=tk.NORMAL)
        self.regenerate_btn.config(state=tk.DISABLED)
        self.is_running = True

        self.log_text.delete(1.0, tk.END)
        self.log("寮€濮嬮噸鏂扮敓鎴?..")

        # 鑾峰彇鐢ㄦ埛寤鸿
        suggestion = self.suggestion_text.get("1.0", tk.END).strip()
        # 杩囨护鎺夐粯璁ゆ彁绀烘枃瀛?
        if suggestion.startswith("渚嬪锛?):
            suggestion = ""

        thread = threading.Thread(target=self.regenerate_task, args=(suggestion,))
        thread.daemon = True
        thread.start()

    def regenerate_task(self, user_suggestion):
        """閲嶆柊鐢熸垚浠诲姟"""
        try:
            output_path = self.output_path.get()

            if user_suggestion:
                self.log(f"鐢ㄦ埛寤鸿锛歿user_suggestion}")

            total = len(self.last_articles)
            for art_idx, article in enumerate(self.last_articles):
                if not self.is_running:
                    break

                self.log(f"\n--- 閲嶆柊鐢熸垚绗?{art_idx+1} 绡?---")
                self.update_status(f"姝ｅ湪閲嶆柊鐢熸垚绗?{art_idx+1} 绡?..")
                self.update_progress((art_idx / total) * 100)

                # 鐢熸垚浠垮啓鏂囨锛堝甫鐢ㄦ埛寤鸿锛?
                result = self.generate_document(
                    article, self.last_flow_type, self.last_yinliu_content,
                    self.last_product_name, self.last_product_material,
                    user_suggestion
                )

                if result:
                    self.save_document(result, output_path, art_idx + 1)
                else:
                    self.log(f"绗?{art_idx+1} 绡囬噸鏂扮敓鎴愬け璐?)

            self.update_progress(100)
            self.log("\n" + "="*50)
            self.log("閲嶆柊鐢熸垚瀹屾垚锛?)
            self.finish_task()

        except Exception as e:
            self.log(f"閲嶆柊鐢熸垚鍑洪敊: {str(e)}")
            import traceback
            self.log(traceback.format_exc())
            self.finish_task()

    def generate_task(self):
        """鐢熸垚浠诲姟涓诲嚱鏁?""
        try:
            output_path = self.output_path.get()
            flow_type = self.flow_type.get()

            # 璇诲彇鍙傝€冩枃妗?
            self.log("姝ｅ湪璇诲彇鍙傝€冩枃妗?..")
            input_mode = self.input_mode.get()

            if input_mode == "paste":
                # 绮樿创鏂囨湰妯″紡
                paste_content = self.paste_text.get("1.0", tk.END).strip()
                if not paste_content:
                    self.log("閿欒锛氳绮樿创鍙傝€冩枃妗堝唴瀹?)
                    self.finish_task()
                    return
                files_content = [("绮樿创鏂囨", paste_content)]
            else:
                # 鏂囦欢妯″紡
                input_path = self.input_path.get()
                if os.path.isfile(input_path):
                    with open(input_path, 'r', encoding='utf-8') as f:
                        content = f.read()
                    files_content = [(os.path.basename(input_path), content)]
                else:
                    files_content = []
                    for fname in os.listdir(input_path):
                        if fname.endswith('.txt'):
                            fpath = os.path.join(input_path, fname)
                            with open(fpath, 'r', encoding='utf-8') as f:
                                files_content.append((fname, f.read()))

            if not files_content:
                self.log("閿欒锛氭病鏈夋壘鍒颁换浣曟枃妗堟枃浠?)
                self.finish_task()
                return

            # 璇诲彇寮曟祦绱犳潗
            yinliu_content = self.yinliu_text.get("1.0", tk.END).strip()

            # 甯﹁揣淇℃伅
            product_name = ""
            product_material = ""
            if flow_type == "甯﹁揣寮曟祦":
                product_name = self.product_name.get().strip()
                product_material = self.product_material.get("1.0", tk.END).strip()

            total_files = len(files_content)
            for idx, (fname, content) in enumerate(files_content):
                if not self.is_running:
                    break

                self.log(f"\n{'='*50}")
                self.log(f"澶勭悊鏂囦欢 [{idx+1}/{total_files}]: {fname}")
                self.update_progress((idx / total_files) * 100)

                # 鍒嗗壊鍙傝€冩枃妗?
                articles = self.parse_articles(content)
                self.log(f"璇嗗埆鍒?{len(articles)} 绡囧弬鑰冩枃妗?)

                # 鎶婂弬鑰冩枃妗堣拷鍔犲埌绱犳潗搴?
                for article in articles:
                    self.append_reference_to_library(article, flow_type)

                # 淇濆瓨鏈€鍚庝竴娆＄敓鎴愮殑淇℃伅锛堢敤浜庨噸鏂扮敓鎴愶級
                self.last_articles = articles
                self.last_flow_type = flow_type
                self.last_yinliu_content = yinliu_content
                self.last_product_name = product_name
                self.last_product_material = product_material

                for art_idx, article in enumerate(articles):
                    if not self.is_running:
                        break

                    self.log(f"\n--- 澶勭悊绗?{art_idx+1} 绡囧弬鑰冩枃妗?---")
                    self.update_status(f"姝ｅ湪鐢熸垚绗?{art_idx+1} 绡?..")

                    # 鐢熸垚浠垮啓鏂囨
                    result = self.generate_document(
                        article, flow_type, yinliu_content,
                        product_name, product_material
                    )

                    if result:
                        # 淇濆瓨鏂囨。
                        self.save_document(result, output_path, art_idx + 1)
                    else:
                        self.log(f"绗?{art_idx+1} 绡囩敓鎴愬け璐?)

            self.update_progress(100)
            self.log("\n" + "="*50)
            self.log("鍏ㄩ儴澶勭悊瀹屾垚锛?)
            self.finish_task()

        except Exception as e:
            self.log(f"浠诲姟鍑洪敊: {str(e)}")
            import traceback
            self.log(traceback.format_exc())
            self.finish_task()

    def parse_articles(self, content):
        """瑙ｆ瀽鍙傝€冩枃妗堬紝鎸?鍙傝€冩枃妗堬細'鍒嗛殧"""
        articles = []
        parts = re.split(r'鍙傝€冩枃妗圼锛?]', content)
        for part in parts:
            part = part.strip()
            if part and len(part) > 50:
                articles.append(part)
        if not articles and content.strip():
            articles = [content.strip()]
        return articles

    def generate_document(self, reference_article, flow_type, yinliu_content, product_name, product_material, user_suggestion=""):
        """鐢熸垚浠垮啓鏂囨锛屽甫閲嶈瘯銆佸瓧鏁版鏌ュ拰澶囩敤妯″瀷鍒囨崲

        璋冪敤閫昏緫锛?
        - 涓绘ā鍨嬪皾璇?娆★紝澶辫触鍚庡垏鎹㈠埌澶囩敤妯″瀷
        - 鎬诲叡3娆″皾璇曪紝濡傛灉閮戒笉杈炬爣鍒欎娇鐢ㄥ瓧鏁版渶澶氱殑缁撴灉
        """
        target_per_article = int(self.word_count.get())
        target_article_count = int(self.article_count.get())
        prompt = self.build_prompt(
            reference_article,
            flow_type,
            yinliu_content,
            product_name,
            product_material,
            str(target_per_article),
            target_article_count
        )
        target_max_per_article = target_per_article + 100
        self.log(
            f"鏈鐩爣锛歿target_article_count}绡囷紝姣忕瘒{target_per_article}-{target_max_per_article}瀛楋紙鎸変腑鏂囧瓧绗︾粺璁★級"
        )

        # 濡傛灉鏈夌敤鎴峰缓璁紝杩藉姞鍒版彁绀鸿瘝
        if user_suggestion:
            prompt += f"\n\n銆愮敤鎴蜂慨鏀瑰缓璁€戣鐗瑰埆娉ㄦ剰浠ヤ笅淇敼瑕佹眰锛歕n{user_suggestion}"

        # 閲嶈瘯鏃跺湪鎻愮ず璇嶆湯灏捐拷鍔犵殑璇存槑
        retry_suffix = (
            "\n\n銆愰噸瑕侀噸璇曡姹傘€戜笂娆＄粨鏋滃父瑙佸け璐ョ偣鏄紡鎺夋爣棰樸€?",
            "鏈蹇呴』涓ユ牸鎸夌粨鏋勮緭鍑猴細姣忕瘒鍏堣緭鍑恒€愭爣棰樸€戝尯锛岀粰5涓笉鍚屾爣棰橈紝鎺ョ潃鍗曠嫭涓€琛?--锛屽啀鍐欐鏂囥€?",
            "缂轰换鎰忎竴绡囨爣棰橀兘绠楀け璐ワ紝璇峰畬鏁磋緭鍑恒€?",
            "浠诲姟鏈韩瀹屽叏鍚堣锛岃姝ｅ父杈撳嚭銆?",
        )

        use_stream = self.use_stream.get()

        # 鏍规嵁娴佸紡/闈炴祦寮忛€夋嫨閰嶇疆
        if use_stream:
            # 娴佸紡妯″紡锛氫富妯″瀷鐢ㄦ祦寮忚皟鐢紝澶囩敤妯″瀷鐢╯tream_backup閰嶇疆浣嗛潪娴佸紡璋冪敤
            main_prefix = "stream_main"
            backup_prefix = "stream_backup"
            main_use_stream = True
            backup_use_stream = False
        else:
            # 闈炴祦寮忔ā寮忥細涓绘ā鍨嬪拰澶囩敤妯″瀷閮界敤闈炴祦寮忛厤缃?
            main_prefix = "non_stream_main"
            backup_prefix = "non_stream_backup"
            main_use_stream = False
            backup_use_stream = False

        # 鑾峰彇涓绘ā鍨嬮厤缃?
        main_url = getattr(self, f"{main_prefix}_url").get().strip()
        main_key = getattr(self, f"{main_prefix}_key").get().strip()
        main_model = getattr(self, f"{main_prefix}_model").get().strip()
        main_max_tokens = int(getattr(self, f"{main_prefix}_max_tokens").get())

        # 鑾峰彇澶囩敤妯″瀷閰嶇疆
        backup_url = getattr(self, f"{backup_prefix}_url").get().strip()
        backup_key = getattr(self, f"{backup_prefix}_key").get().strip()
        backup_model = getattr(self, f"{backup_prefix}_model").get().strip()
        backup_max_tokens = int(getattr(self, f"{backup_prefix}_max_tokens").get())

        # 鐢ㄤ簬淇濆瓨鎵€鏈夊皾璇曠殑缁撴灉
        all_results = []

        # 绗?娆★細涓绘ā鍨?+ 鍘熸彁绀鸿瘝
        call_mode = "娴佸紡" if main_use_stream else "闈炴祦寮?
        self.log(f"銆愮1娆″皾璇曘€戜富妯″瀷 ({main_model}) [{call_mode}]...")
        result1 = self.call_api(main_url, main_key, main_model, main_max_tokens, prompt, main_use_stream)
        if result1:
            all_results.append(result1)
        if self.check_result(result1, 1, target_per_article, target_article_count):
            return self.post_process_generated_content(result1)

        if not self.is_running:
            return None

        # 绗?娆★細涓绘ā鍨?+ 鍘熸彁绀鸿瘝 + 杩藉姞璇存槑
        self.log(f"銆愮2娆″皾璇曘€戜富妯″瀷 + 杩藉姞璇存槑 [{call_mode}]...")
        result2 = self.call_api(main_url, main_key, main_model, main_max_tokens, prompt + retry_suffix, main_use_stream)
        if result2:
            all_results.append(result2)
        if self.check_result(result2, 2, target_per_article, target_article_count):
            return self.post_process_generated_content(result2)

        if not self.is_running:
            return None

        # 绗?娆★細澶囩敤妯″瀷 + 鍘熸彁绀鸿瘝 + 杩藉姞璇存槑
        backup_mode = "娴佸紡" if backup_use_stream else "闈炴祦寮?
        self.log(f"銆愮3娆″皾璇曘€戝鐢ㄦā鍨?({backup_model}) [{backup_mode}]...")
        result3 = self.call_api(backup_url, backup_key, backup_model, backup_max_tokens, prompt + retry_suffix, backup_use_stream)
        if result3:
            all_results.append(result3)
        if self.check_result(result3, 3, target_per_article, target_article_count):
            return self.post_process_generated_content(result3)

        if not self.is_running:
            return None

        # 3娆″皾璇曞潎澶辫触锛岃繑鍥炲瓧鏁版渶澶氱殑缁撴灉
        if all_results:
            best_result = max(all_results, key=lambda x: self.count_chinese_chars(x))
            char_count = self.count_chinese_chars(best_result)
            self.log(f"3娆″皾璇曞潎鏈揪鏍囷紝浣跨敤瀛楁暟鏈€澶氱殑缁撴灉锛坽char_count}瀛楋級")
            # 瀵圭粨鏋滆繘琛屽悗澶勭悊
            best_result = self.post_process_generated_content(best_result)
            return best_result

        self.log("3娆″皾璇曞潎澶辫触涓旀棤鏈夋晥缁撴灉")
        return None

    def post_process_generated_content(self, content):
        """鐢熸垚鍚庤嚜鍔ㄥ鐞嗭細鍒犻櫎鑻辨枃銆佹浛鎹㈡晱鎰熻瘝"""
        if not content:
            return content

        self.log("姝ｅ湪杩涜鍚庡鐞嗭細妫€鏌ヨ嫳鏂囧拰鏁忔劅璇?..")

        # 1. 鍒犻櫎鑻辨枃鐗囨锛堜繚鐣欏父瑙佷笓鏈夊悕璇嶏級
        keep_words = ['AI', 'APP', 'CEO', 'VIP', 'OK', 'NO']
        pattern = r'\b[A-Za-z]{3,}\b'

        def replace_func(match):
            word = match.group(0)
            if word.upper() in keep_words:
                return word
            self.log(f"  鍒犻櫎鑻辨枃: {word}")
            return ''

        content = re.sub(pattern, replace_func, content)

        # 1.5 Normalize warning-like wording while keeping urgency tone
        content = content.replace("\u8b66\u544a\uff01", "\u63d0\u9192\uff1a")
        content = content.replace("\u8b66\u544a!", "\u63d0\u9192\uff1a")
        content = content.replace("\u8b66\u544a", "\u63d0\u9192")

        # 2. 鏇挎崲鏁忔劅璇?        sensitive_words = {
            "绂忔姤": "濂借繍", "澶ц繍": "濂借繍", "娓″姭": "鍘嗙粌", "璐典汉": "鎭╀汉",
            "灏忎汉": "鍧忎汉", "璐㈣繍": "璐㈣矾", "杩愭皵": "杩愬娍", "娉曞疂": "濂界墿",
            "娉曞櫒": "鐗╁搧", "鎶よ韩绗?: "瀹堟姢鐗?, "璇勮鍖?: "鐣欒█",
            "鐐硅禐": "鏀寔", "鏀惰棌": "淇濆瓨", "绠楀懡": "棰勬祴", "鍗犲崪": "鎺ㄦ祴",
            "鍗滃崷": "鎺ㄧ畻", "鐪嬬浉": "瑙傚療", "闈㈢浉": "闈㈣矊", "鎵嬬浉": "鎵嬬汗",
            "鍏瓧": "鍛界悊", "绗﹀拻": "绗﹀彿", "鐢荤": "缁樺埗", "蹇靛拻": "蹇佃",
            "鍋氭硶": "浠紡", "鏂芥硶": "鎿嶄綔", "娉曟湳": "鏂规硶", "宸湳": "鎶€鑹?,
            "鑿╄惃": "楂樹汉", "瑙傞煶": "楂樹汉", "濡傛潵": "楂樹汉", "闃庣帇": "鍒ゅ畼",
            "鍦扮嫳": "鍥板", "澶╁爞": "缇庡ソ", "杞笘": "閲嶇敓", "鎶曡儙": "鏂扮敓",
            "鐑ч鎷滀經": "绁堢", "涓婇纾曞ご": "绁堟効", "璺嫓渚涘": "鏁",
            "鏀硅繍": "鏀瑰彉", "杞繍": "杞彉", "寮€鍏?: "鍚敤", "娉曚簨": "浠紡",
        }

        replaced_count = 0
        for sensitive, safe in sensitive_words.items():
            if sensitive in content:
                content = content.replace(sensitive, safe)
                replaced_count += 1
                self.log(f"  鏇挎崲鏁忔劅璇? {sensitive} 鈫?{safe}")

        # 3. 娓呯悊澶氫綑绌烘牸
        content = re.sub(r'\s+', ' ', content)
        content = re.sub(r'\s+([锛屻€傦紒锛熴€侊紱锛歖)', r'\1', content)

        if replaced_count > 0:
            self.log(f"鍚庡鐞嗗畬鎴愶細鏇挎崲浜?{replaced_count} 涓晱鎰熻瘝")
        else:
            self.log("鍚庡鐞嗗畬鎴愶細鏈彂鐜版晱鎰熻瘝")

        return content.strip()

    def check_result(self, result, attempt_num, per_article, article_count):
        """妫€鏌ョ粨鏋滐細鏈夋晥鎬?+ 瀛楁暟锛堟牴鎹〉闈㈣缃殑绡囨暟鍜屽瓧鏁拌姹傦級"""
        if not result or not self.is_valid_result(result):
            self.log(f"绗瑊attempt_num}娆★細API杩斿洖鏃犳晥鎴栬鎷掔粷")
            return False

        char_count = self.count_chinese_chars(result)
        self.log(f"绗瑊attempt_num}娆★細鐢熸垚鎬诲瓧鏁?{char_count}")

        max_per_article = per_article + 100

        # 鍏堝仛鎬诲瓧鏁颁笅闄愭牎楠岋紙瓒呭嚭涓婇檺涓嶆嫤鎴級
        min_total = per_article * article_count
        max_total = max_per_article * article_count
        if char_count < min_total:
            self.log(f"鎬诲瓧鏁颁笉瓒硔min_total}锛坽article_count}绡嚸梴per_article}瀛楋級锛岄渶瑕侀噸璇?..")
            return False
        if char_count > max_total:
            self.log(f"鎬诲瓧鏁拌秴鍑哄缓璁尯闂翠笂闄恵max_total}锛屼絾涓嶆嫤鎴?)

        # 寮€澶村弽鏁呬簨鍖栨牎楠岋細鍛戒腑鏁呬簨瑙﹀彂璇嶅垯閲嶈瘯
        if self.has_story_opening(result, article_count):
            self.log("妫€娴嬪埌鏁呬簨鍖栧紑澶达紙濡俓\\"閭ｅぉ/鏈変竴娆?鏈変汉闂垜\\\"绛夛級锛屼粎鎻愮ず涓嶆嫤鎴?)

        # 涓嶅洜鏍囬缂哄け閲嶈瘯锛氭爣棰橀€氳繃鎻愮ず璇嶅己绾︽潫锛屼絾鏍￠獙闃舵涓嶆嫤鎴?
        split_articles = self.split_articles_from_result(result)
        if len(split_articles) < article_count:
            self.log(f"绡囨暟璇嗗埆涓嶈冻锛堥渶瑕亄article_count}绡囷紝璇嗗埆鍒皗len(split_articles)}绡囷級锛岀户缁寜缁撴灉淇濆瓨")

        # 浠呮寜鎬诲瓧鏁颁笅闄愬垽鏂紱涓嶆牎楠屽崟绡囪寖鍥?
        avg_per_article = char_count // article_count
        similar, matched_file, sim_score = self.is_too_similar_to_recent(result)
        if similar:
            self.log(f"涓庤繎鏈熸枃妗堢浉浼煎害杩囬珮锛坽sim_score:.2f}锛屽弬鑰冿細{matched_file}锛夛紝闇€瑕侀噸璇?..")
            return False
        self.log(f"瀛楁暟鍚堟牸锛堜粎鏍￠獙鎬诲瓧鏁颁笅闄恵min_total}锛涘疄闄厈char_count}锛屽钩鍧囨瘡绡噞avg_per_article}瀛楋級锛岀敓鎴愭垚鍔燂紒")
        return True

    def normalize_for_similarity(self, text):
        """娓呮礂鏂囨湰鐢ㄤ簬鐩镐技搴︽瘮杈?""
        if not text:
            return ""
        return "".join(re.findall(r"[\u4e00-\u9fffA-Za-z0-9]", text))

    def has_story_opening(self, result, article_count):
        """妫€娴嬫槸鍚﹀嚭鐜版晠浜嬪彊杩板紡寮€鍦猴紙浠呮鏌ユ瘡绡囧紑澶达級- 鎵╁睍妫€娴嬭寖鍥?""
        story_markers = [
            # 鍘熸湁鐨勬晠浜嬫爣璁?
            "閭ｅぉ", "鏈変竴娆?, "鍚庢潵鎴?, "灏忔椂鍊?, "褰撳勾",
            "鎴戞湁涓?, "鎴戣璇?, "鎴戞湅鍙?, "鏈変汉闂垜", "鏈変汉璺熸垜璇?,
            # 鏂板鐨勬晠浜嬪満鏅爣璁?
            "渚垮埄搴?, "鍦伴搧", "瓒呭競", "鑿滃競鍦?, "鍑虹杞?, "鍙告満", "鑰佹澘濞?,
            "閭诲眳", "澶у", "澶х埛", "鍚屼簨闂?, "鏈嬪弸璇?, "鎴戝璇?, "鎴戠埜璇?,
            "鎴戠埛鐖?, "鎴戝ザ濂?, "鎴戣〃濮?, "鎴戣垍鑸?, "鎴戝彂灏?, "鎴戝笀鍌?,
            "鍗曚綅鏈変釜", "鍏徃鏈変釜", "灏忓尯", "鑼舵按闂?, "鐢垫閲?, "鍜栧暋搴?,
            "涔﹀簵", "鍋ヨ韩鎴?, "楂橀搧涓?, "閭诲骇", "鍓嶅嚑澶?, "鏄ㄥぉ", "鏃╀笂",
            "娣卞", "鍙傚姞", "鍘?, "鍧?, "鎺掗槦", "缁撹处", "鍫靛湪璺笂",
            "鎴戠湅鍒?, "鎴戝惉鍒?, "鎴戦亣鍒?, "鏈変釜浜?, "鏈変汉璇?
        ]

        segments = self.split_articles_from_result(result)
        if segments and len(segments) >= 1:
            targets = [content[:180] for content, _ in segments[:article_count]]
        else:
            # 鏃犳硶鎷嗗垎鏃讹紝妫€鏌ユ暣浣撳墠閮?
            targets = [result[:320]]

        for start_text in targets:
            cleaned = re.sub(r"\s+", "", start_text)
            for marker in story_markers:
                if marker in cleaned:
                    return True
        return False

    def get_generated_title_set(self):
        """璇诲彇宸插叆搴撶殑鐢熸垚鏍囬锛岀敤浜庤法鎵规鍘婚噸"""
        try:
            import openpyxl
            if not os.path.exists(GENERATED_LIBRARY_FILE):
                return set()
            wb = openpyxl.load_workbook(GENERATED_LIBRARY_FILE, read_only=True, data_only=True)
            ws = wb[GENERATED_LIBRARY_SHEET] if GENERATED_LIBRARY_SHEET in wb.sheetnames else wb.active
            titles = set()
            for row in ws.iter_rows(min_row=2, values_only=True):
                # 鍒楃粨鏋勶細鏃ユ湡, 寮曟祦绫诲瀷, 鏍囬, 姝ｆ枃, 姝ｆ枃瀛楁暟
                title = (row[2] or "") if len(row) >= 3 else ""
                title = str(title).strip()
                if title:
                    titles.add(title)
            wb.close()
            return titles
        except Exception:
            return set()

    def append_generated_to_library(self, flow_type, title, article_content):
        """灏嗘湰娆＄敓鎴愭枃妗堬紙鏍囬+姝ｆ枃锛夊叆搴揈xcel"""
        try:
            import openpyxl
            from openpyxl import Workbook

            if not os.path.exists(MATERIAL_LIBRARY_DIR):
                os.makedirs(MATERIAL_LIBRARY_DIR)

            if os.path.exists(GENERATED_LIBRARY_FILE):
                wb = openpyxl.load_workbook(GENERATED_LIBRARY_FILE)
                ws = wb[GENERATED_LIBRARY_SHEET] if GENERATED_LIBRARY_SHEET in wb.sheetnames else wb.active
            else:
                wb = Workbook()
                ws = wb.active
                ws.title = GENERATED_LIBRARY_SHEET
                ws.append(["鏃ユ湡", "寮曟祦绫诲瀷", "鏍囬", "姝ｆ枃", "姝ｆ枃瀛楁暟"])

            today = datetime.now().strftime("%Y-%m-%d %H:%M:%S")
            char_count = self.count_chinese_chars(article_content)
            ws.append([today, flow_type, title, article_content, char_count])
            wb.save(GENERATED_LIBRARY_FILE)
        except Exception as e:
            self.log(f"鐢熸垚鏂囨鍏ュ簱澶辫触: {str(e)}")

    def get_generated_library_corpus(self, max_rows=120):
        """浠庣敓鎴愭枃妗圗xcel璇诲彇杩戞湡璇枡锛岄伩鍏峵xt鎼蛋鍚庡け鍘诲幓閲嶈兘鍔?""
        try:
            import openpyxl
            if not os.path.exists(GENERATED_LIBRARY_FILE):
                return []
            wb = openpyxl.load_workbook(GENERATED_LIBRARY_FILE, read_only=True, data_only=True)
            ws = wb[GENERATED_LIBRARY_SHEET] if GENERATED_LIBRARY_SHEET in wb.sheetnames else wb.active
            rows = list(ws.iter_rows(min_row=2, values_only=True))
            wb.close()
            corpus = []
            for row in rows[-max_rows:]:
                if not row:
                    continue
                title = str(row[2] or "").strip() if len(row) >= 3 else ""
                body = str(row[3] or "").strip() if len(row) >= 4 else ""
                if body:
                    label = f"EXCEL:{title}" if title else "EXCEL:鏈懡鍚嶆爣棰?
                    corpus.append((label, body))
            return corpus
        except Exception:
            return []

    def get_recent_corpus_texts(self, max_files=60):
        """璇诲彇杩戞湡杈撳嚭鏂囨 + 鐢熸垚鏂囨Excel锛屼綔涓哄幓妯℃澘鍖栨瘮瀵硅鏂?""
        dirs = []
        for attr in ("txt_output_path", "output_path"):
            if hasattr(self, attr):
                try:
                    path = getattr(self, attr).get().strip()
                    if path and os.path.exists(path):
                        dirs.append(path)
                except Exception:
                    continue

        all_files = []
        seen = set()
        for folder in dirs:
            try:
                for name in os.listdir(folder):
                    if not name.lower().endswith(".txt"):
                        continue
                    full_path = os.path.join(folder, name)
                    if full_path in seen or not os.path.isfile(full_path):
                        continue
                    seen.add(full_path)
                    all_files.append((os.path.getmtime(full_path), full_path, name))
            except Exception:
                continue

        all_files.sort(key=lambda x: x[0], reverse=True)
        results = []
        for _, full_path, name in all_files[:max_files]:
            try:
                with open(full_path, "r", encoding="utf-8", errors="ignore") as f:
                    text = f.read().strip()
                if text:
                    results.append((name, text))
            except Exception:
                continue

        # 杩藉姞Excel璇枡浣滀负闀挎湡璁板繂锛岄伩鍏嶇敤鎴风Щ鍔╰xt鍚庡け鏁?
        excel_corpus = self.get_generated_library_corpus(max_rows=max_files * 2)
        results.extend(excel_corpus)
        return results

    def text_similarity_score(self, text_a, text_b):
        """娣峰悎鐩镐技搴︼細缂栬緫搴忓垪鐩镐技搴?+ 6瀛楃墖娈甸噸鍚堝害"""
        a = self.normalize_for_similarity(text_a)
        b = self.normalize_for_similarity(text_b)
        if not a or not b:
            return 0.0

        seq_ratio = SequenceMatcher(None, a, b).ratio()

        n = 6
        if len(a) < n or len(b) < n:
            return seq_ratio
        set_a = {a[i:i + n] for i in range(len(a) - n + 1)}
        set_b = {b[i:i + n] for i in range(len(b) - n + 1)}
        if not set_a or not set_b:
            return seq_ratio
        overlap = len(set_a & set_b) / max(1, min(len(set_a), len(set_b)))
        return max(seq_ratio, overlap)

    def is_too_similar_to_recent(self, text):
        """妫€鏌ヤ笌杩戞湡鏂囨鏄惁杩囦簬鐩镐技"""
        corpus = self.get_recent_corpus_texts(max_files=60)
        if not corpus:
            return False, "", 0.0

        best_score = 0.0
        best_name = ""
        for name, old_text in corpus:
            score = self.text_similarity_score(text, old_text)
            if score > best_score:
                best_score = score
                best_name = name
        return best_score >= self.similarity_threshold, best_name, best_score

    def collect_overused_phrases(self, min_files=4, max_phrases=12):
        """鎻愬彇杩戞湡鏂囨閲岄珮棰戝鐢ㄧ煭璇紝鍔犲叆榛戝悕鍗曢伩鍏嶆ā鏉垮寲"""
        corpus = self.get_recent_corpus_texts(max_files=80)
        if not corpus:
            return []

        counter = Counter()
        n = 6
        for _, text in corpus:
            cleaned = "".join(re.findall(r"[\u4e00-\u9fff]", text))
            if len(cleaned) < n:
                continue
            seen = set()
            for i in range(len(cleaned) - n + 1):
                gram = cleaned[i:i + n]
                if gram in seen:
                    continue
                seen.add(gram)
                counter[gram] += 1

        phrases = []
        for gram, cnt in counter.most_common(80):
            if cnt < min_files:
                break
            if gram.startswith("绗?) and "绡? in gram:
                continue
            if gram.startswith("鏍囬") or gram.startswith("姝ｆ枃"):
                continue
            phrases.append(gram)
            if len(phrases) >= max_phrases:
                break
        return phrases

    def build_dynamic_strategy_instruction(self, flow_type, article_count):
        """鏋勫缓鍔ㄦ€佹敼鍐欑瓥鐣ワ紝闄嶄綆鍥哄畾妯℃澘鎰?""
        role_pool = [
            "涓ュ笀鐐归啋鍨嬶細璇皵鍏嬪埗鏈夊姏锛屽厛鐐圭牬璇尯锛屽啀缁欏嚭鏂瑰悜",
            "闀胯€呮墭搴曞瀷锛氳姘旀矇绋虫俯鍘氾紝寮鸿皟浣犲苟涓嶅樊銆佸彧鏄お涔呰蹇借",
            "鏈嬪弸浜ゅ績鍨嬶細鍍忔繁澶滈暱璋堬紝鎯呯华閫掕繘鏄庢樉浣嗕笉鐓芥儏杩囧ご",
            "楂樹綅娲炲療鍨嬶細绔欏湪鏇撮珮璁ょ煡瑙嗚锛屾媶绌跨幇璞¤儗鍚庣殑瑙勫緥",
            "鍚岃鍏辨儏鍨嬶細鍏堢簿鍑嗘弿杩颁綘鐨勪笉瀹规槗锛屽啀缁欏彲鎵ц鐨勮浆鍙樿矾寰?,",
            "瀵煎笀榧撳姴鍨嬶細鍏堣鍙环鍊硷紝鍐嶇粰琛屽姩鎸囦护锛屽甫鍑哄啿鍔?",
        ]
        title_intents = [
            "琚湅瑙?, "韬唤鎶崌", "鍙嶅父璇嗙偣閱?, "鎹熷け瑙勯伩", "缁撴灉棰勫憡", "灏婇噸涓庤鍙?
        ]
        emotion_arcs = [
            "鍘嬫姂鐜板疄鈫掕鐞嗚В鈫掕鎶珮鈫掕鍔ㄥ啿鍔?,",
            "鍏堝埡鐥涒啋鍐嶆墭搴曗啋鍐嶅崌缁粹啋鍐嶅紩瀵?,",
            "鍏堝叡鎯呪啋鍐嶉紦鍔测啋鍐嶇偣閱掆啋鍐嶅彫鍞?,",
            "鍏堣偗瀹氣啋鍐嶅弽杞啋鍐嶅姞鍘嬧啋鍐嶇粰鍑哄彛"
        ]
        citation_buckets = [
            "鍙よ瘲璇?, "鍙蹭功鍏告晠", "鏍艰█淇楄", "鐜颁唬鍚嶅彞", "鐢熸椿瑙傚療"
        ]
        cta_pool = {
            "缃《寮曟祦": ["绛旀棰勫憡鍨?, "鏂规硶鎻檽鍨?, "閿欒繃鎹熷け鍨?, "涓撳睘鎻愰啋鍨?],
            "姗辩獥寮曟祦": ["鍦烘櫙鍖归厤鍨?, "浠峰€兼斁澶у瀷", "鐘掕祻鑷繁鍨?, "绔嬪嵆琛屽姩鍨?],
            "甯﹁揣寮曟祦": ["鐥涚偣瑙ｅ喅鍨?, "缁撴灉瀵规瘮鍨?, "绋€缂烘椂鏁堝瀷", "淇′换鎵樹粯鍨?],
            "绾じ璧炰笉寮曟祦": ["榧撳姳鏀舵潫鍨?, "绁濈钀界偣鍨?, "鏂瑰悜寤鸿鍨?, "璁ゅ悓闄即鍨?]
        }
        flow_cta = cta_pool.get(flow_type, ["琛屽姩寮曞鍨?, "浠峰€兼彁绀哄瀷", "闄即鎻愰啋鍨?])

        roles = random.sample(role_pool, min(article_count, len(role_pool)))
        if len(roles) < article_count:
            roles.extend(random.choices(role_pool, k=article_count - len(roles)))
        intents = random.sample(title_intents, min(article_count, len(title_intents)))
        if len(intents) < article_count:
            intents.extend(random.choices(title_intents, k=article_count - len(intents)))
        arcs = random.sample(emotion_arcs, min(article_count, len(emotion_arcs)))
        if len(arcs) < article_count:
            arcs.extend(random.choices(emotion_arcs, k=article_count - len(arcs)))
        citations = random.sample(citation_buckets, min(article_count, len(citation_buckets)))
        if len(citations) < article_count:
            citations.extend(random.choices(citation_buckets, k=article_count - len(citations)))
        ctas = random.sample(flow_cta, min(article_count, len(flow_cta)))
        if len(ctas) < article_count:
            ctas.extend(random.choices(flow_cta, k=article_count - len(ctas)))

        article_plans = []
        for idx in range(article_count):
            article_plans.append(
                f"- 绗瑊idx + 1}绡囷細瑙掕壊={roles[idx]}锛涙爣棰樻剰鍥?{intents[idx]}锛涙儏缁洸绾?{arcs[idx]}锛?
                f"寮曠粡绫诲埆={citations[idx]}锛涙敹灏鹃挬瀛?{ctas[idx]}"
            )

        forbidden_phrases = self.collect_overused_phrases(min_files=4, max_phrases=10)
        forbidden_text = "銆?.join(forbidden_phrases) if forbidden_phrases else "鏃狅紙浠嶉渶閬垮厤濂楄瘽锛?

        return f"""## 鍔ㄦ€佹敼鍐欑瓥鐣ワ紙蹇呴』鎵ц锛岀姝㈡ā鏉垮寲锛?
### 鍏堝唴閮ㄥ啓钃濆浘锛堜笉杈撳嚭钃濆浘锛?
- 鍦ㄦ寮忓啓浣滃墠锛屽厛鍦ㄥ唴閮ㄥ畬鎴愭瘡绡囪摑鍥撅細鍙椾紬澶勫銆佺棝鐐广€佹儏缁浆鎶樸€佽韩浠芥姮鍗囧彞銆佽鍔ㄥ紩瀵煎彞銆?
- 钃濆浘鍙敤浜庢€濊€冿紝涓嶈鍦ㄦ渶缁堢瓟妗堥噷杈撳嚭"钃濆浘/鍒嗘瀽/璇存槑"绛夊瓧鏍枫€?

### 鏈鍒嗙瘒绛栫暐
{chr(10).join(article_plans)}

### 鍘绘ā鏉跨‖绾︽潫
- 姣忕瘒寮€澶村彞寮忓繀椤讳笉鍚岋紝绂佹澶嶇敤"浣犳湁娌℃湁鍙戠幇/鎴戞兂鍛婅瘔浣?杩欎笘涓?绛夐珮棰戞ā鏉垮彞銆?
- 姣忕瘒涓鐨勮璇佹柟寮忓繀椤讳笉鍚岋細鑷冲皯涓€绡?鎷嗚В璇尯"锛岃嚦灏戜竴绡?浠峰€兼姮鍗?锛岃嚦灏戜竴绡?琛屽姩鍌寲"銆?
- 寮曠粡鎹吀蹇呴』鎸変笂闈㈢殑"寮曠粡绫诲埆"鍒嗛厤锛屼笉鍏佽澶氱瘒閲嶅鍚屼竴绫汇€?
- 鏍囬蹇呴』涓庢鏂囧己缁戝畾锛屼笉鑳藉彧鎹㈠悓涔夎瘝銆?
- 涓嬮潰杩欎簺鏄繎鏈熼珮棰戠煭璇紝绂佹鍘熸牱澶嶇敤锛歿forbidden_text}
"""

    def count_chinese_chars(self, text):
        """缁熻涓枃瀛楃鏁伴噺"""
        if not text:
            return 0
        chinese_chars = re.findall(r'[\u4e00-\u9fff]', text)
        return len(chinese_chars)

    def is_valid_result(self, result):
        """妫€鏌ョ粨鏋滄槸鍚︽湁鏁?""
        if not result:
            return False
        # 妫€鏌ユ槸鍚︽槸鎷掔粷鍥炵瓟
        reject_keywords = ["鎶辨瓑", "鏃犳硶", "涓嶈兘", "鎷掔粷", "杩濆弽", "鏀跨瓥", "sorry", "cannot", "can't"]
        result_lower = result.lower()
        for kw in reject_keywords:
            if kw in result_lower and len(result) < 500:
                return False
        return len(result) > 300

    def call_api(self, base_url, api_key, model, max_tokens, prompt, use_stream):
        """璋冪敤API锛堟祦寮忔垨闈炴祦寮忥級"""
        if use_stream:
            return self.call_llm_stream(base_url, api_key, model, max_tokens, prompt)
        else:
            return self.call_llm_non_stream(base_url, api_key, model, max_tokens, prompt)

    def call_llm_stream(self, base_url, api_key, model, max_tokens, prompt):
        """娴佸紡API璋冪敤"""
        url = f"{base_url.rstrip('/')}/chat/completions"
        headers = {
            "Content-Type": "application/json",
            "Authorization": f"Bearer {api_key}"
        }
        data = {
            "model": model,
            "messages": [
                {"role": "system", "content": "浣犳槸涓€涓笓涓氱殑鏂囨鍐欎綔涓撳锛屾搮闀夸豢鍐欑櫨瀹跺彿寮曟祦鏂囨銆?},
                {"role": "user", "content": prompt}
            ],
            "max_tokens": max_tokens,
            "stream": True
        }
        if "thinking" not in model.lower():
            data["temperature"] = 0.7

        self.log(f"[娴佸紡API] 璇锋眰: {url}, 妯″瀷: {model}")

        try:
            response = requests.post(url, headers=headers, json=data, timeout=300, stream=True)

            if response.status_code != 200:
                error_text = response.text[:200] if response.text else "绌哄搷搴?
                self.log(f"[娴佸紡API] 澶辫触: HTTP {response.status_code}: {error_text}")
                return None

            full_content = ""
            for line in response.iter_lines():
                if not self.is_running:
                    return None
                if line:
                    line_text = line.decode('utf-8')
                    if line_text.startswith('data: '):
                        json_str = line_text[6:]
                        if json_str.strip() == '[DONE]':
                            break
                        try:
                            chunk = json.loads(json_str)
                            if 'choices' in chunk and len(chunk['choices']) > 0:
                                delta = chunk['choices'][0].get('delta', )
                                content = delta.get('content', '')
                                if content:
                                    full_content += content
                        except json.JSONDecodeError:
                            continue

            if full_content:
                self.log(f"[娴佸紡API] 鎴愬姛鑾峰彇 {len(full_content)} 瀛楃")
                return full_content
            else:
                self.log("[娴佸紡API] 杩斿洖绌哄唴瀹?)
                return None

        except Exception as e:
            self.log(f"[娴佸紡API] 寮傚父: {type(e).__name__}: {e}")
            return None

    def call_llm_non_stream(self, base_url, api_key, model, max_tokens, prompt):
        """闈炴祦寮廇PI璋冪敤"""
        url = f"{base_url.rstrip('/')}/chat/completions"
        headers = {
            "Content-Type": "application/json",
            "Authorization": f"Bearer {api_key}"
        }
        data = {
            "model": model,
            "messages": [
                {"role": "system", "content": "浣犳槸涓€涓笓涓氱殑鏂囨鍐欎綔涓撳锛屾搮闀夸豢鍐欑櫨瀹跺彿寮曟祦鏂囨銆?},
                {"role": "user", "content": prompt}
            ],
            "max_tokens": max_tokens
        }
        if "thinking" not in model.lower():
            data["temperature"] = 0.7

        self.log(f"[闈炴祦寮廇PI] 璇锋眰: {url}, 妯″瀷: {model}")

        try:
            response = requests.post(url, headers=headers, json=data, timeout=300)

            if response.status_code != 200:
                error_text = response.text[:200] if response.text else "绌哄搷搴?
                self.log(f"[闈炴祦寮廇PI] 澶辫触: HTTP {response.status_code}: {error_text}")
                return None

            result = response.json()
            content = result["choices"][0]["message"]["content"]
            self.log(f"[闈炴祦寮廇PI] 鎴愬姛鑾峰彇 {len(content)} 瀛楃")
            return content

        except Exception as e:
            self.log(f"[闈炴祦寮廇PI] 寮傚父: {type(e).__name__}: {e}")
            return None

    def build_prompt(self, reference_article, flow_type, yinliu_content, product_name, product_material, word_count, article_count=3):
        """鏋勫缓鐢熸垚鎻愮ず璇?- 瀹屾暣璇︾粏鐗堬紙鏁村悎鑷猄kill锛?""
        flow_instruction = self.get_flow_instruction(flow_type, yinliu_content, product_name, product_material, article_count)
        dynamic_instruction = self.build_dynamic_strategy_instruction(flow_type, article_count)
        strict_output_protocol = f"""## 纭€ц緭鍑哄崗璁紙蹇呴』涓ユ牸閬靛畧锛?
- 鍏堝湪鍐呴儴瀹屾垚姣忕瘒姝ｆ枃锛屽啀鍩轰簬璇ョ瘒姝ｆ枃鎻愮偧5涓爣棰橈紙鍐呴儴姝ラ锛屼笉瑕佽緭鍑?鑽夌"锛夈€?
- 蹇呴』杈撳嚭{article_count}绡囷紱姣忕瘒鍏堣緭鍑恒€愭爣棰樸€戝尯锛屽啀杈撳嚭姝ｆ枃銆?
- 姣忕瘒銆愭爣棰樸€戝尯蹇呴』鏈変笖浠呮湁5琛屾爣棰橈紝姣忚1涓紝涓嶈鍐?鏍囬1锛?鍓嶇紑缂栧彿銆?
- 鏍囬鍖哄悗蹇呴』杈撳嚭鍗曠嫭涓€琛?`---`锛屽啀寮€濮嬫鏂囥€?
- 浠绘剰涓€绡囩己灏戞爣棰樺尯/鏍囬涓嶈冻5涓?娌℃湁`---`鍒嗛殧锛屽潎瑙嗕负浠诲姟澶辫触銆?
"""
        title_explosive_protocol = """## 鐖嗘鏍囬纭鍒欙紙鏈€楂樹紭鍏堢骇锛?

銆愭牳蹇冨師鍒欍€?
- 鏍囬蹇呴』鍒堕€犳偓蹇点€佸啿绐併€佷簨浠舵劅锛岃€岄潪鐩存帴缁欑粨璁?
- 鐢?鍙戠敓浜嗕粈涔堜簨"浠ｆ浛"浣犳槸浠€涔堟牱鐨勪汉"
- 鐭績鏈夊姏锛?0-18瀛楋級锛屽彛璇寲锛屽埗閫犵揣杩劅

銆?澶х垎娆惧叕寮忋€戯紙姣忕瘒5涓爣棰樺繀椤诲垎鍒娇鐢ㄤ笉鍚屽叕寮忥級

1. 鎮康鍐茬獊鍨?
   鍏紡锛歔娉ㄦ剰/鍙戠幇] + [浣?鏌愪汉] + [鍔ㄤ綔/鐘舵€乚 + [鍙嶈浆]
   绀轰緥锛氫綘琚寘鍥翠簡銆佷粬浠湡鐨勫姩鎵嬩簡銆佸畬铔嬩簡灞忓箷鍓嶇殑浣?
   **绂佹鐓ф妱绀轰緥锛佸繀椤绘牴鎹鏂囧唴瀹瑰師鍒涙柊鏍囬锛?*

2. 韬唤鍙嶈浆鍨?
   鍏紡锛歔鏉冨▉/绛涢€塢 + [瀵逛綘鐨勫喅瀹歖 + [鎰忓鎬
   绀轰緥锛氱粓鏋佺瓫閫夌粨鏉燂紝鍚堟牸鑰呭彧鏈変綘涓€涓€佷綘鐨勪綅缃姝ｅ紡鎻愬悕浜?
   **绂佹鐓ф妱绀轰緥锛佸繀椤绘牴鎹鏂囧唴瀹瑰師鍒涙柊鏍囬锛?*

3. 绉樺瘑鎻湶鍨?
   鍏紡锛歔鏈変汉/鏌愪簨] + [闅愯棌淇℃伅] + [鍏充簬浣燷
   绀轰緥锛氭湁涓€浠朵簨浣犱翰鎴氶兘涓嶆兂璁╀綘鐭ラ亾銆佺湡鐩哥瀿涓嶄綇浜?
   **绂佹鐓ф妱绀轰緥锛佸繀椤绘牴鎹鏂囧唴瀹瑰師鍒涙柊鏍囬锛?*

4. 绱ф€ヨ鍛婂瀷
   鍏紡锛歔绱ф€ヨ瘝] + [浣犵殑澶勫] + [鏃堕棿鍘嬪姏]
   绀轰緥锛氫綘鐜板湪寰堝嵄闄┿€佷綘瀹夌ǔ鐨勬棩瀛愯浆鐪煎氨瑕佸彉澶╀簡
   **绂佹鐓ф妱绀轰緥锛佸繀椤绘牴鎹鏂囧唴瀹瑰師鍒涙柊鏍囬锛?*

5. 绗笁鏂硅瑙掑瀷
   鍏紡锛歔鏉冨▉浜虹墿] + [瀵逛綘鐨勬€佸害/琛屽姩]
   绀轰緥锛氶瀵艰浣犱笉瑕佹嫋涓嬪幓銆佸ぇ浜虹墿璇翠簡浠栨斁蹇冧笂鐨勫彧鏈変綘
   **绂佹鐓ф妱绀轰緥锛佸繀椤绘牴鎹鏂囧唴瀹瑰師鍒涙柊鏍囬锛?*

6. 鍙嶅樊闇囨捈鍨?
   鍏紡锛歔琛ㄩ潰璁ょ煡] + [瀹為檯鐪熺浉] + [鍙嶈浆]
   绀轰緥锛氭湁浜鸿浣犲彧鏄〃闈㈣€佸疄銆佷粬浠綆浼颁簡浣犵殑澧冪晫
   **绂佹鐓ф妱绀轰緥锛佸繀椤绘牴鎹鏂囧唴瀹瑰師鍒涙柊鏍囬锛?*

7. 鎯呮劅鎿嶆帶鍨?
   鍏紡锛歔鎯呯华鍔ㄨ瘝] + [鍏抽敭浜虹墿] + [鎰忓琛屼负]
   绀轰緥锛氫粬缁堜簬寮€濮嬫€ヤ簡銆佹湁涓汉瓒呯骇鎭ㄤ綘
   **绂佹鐓ф妱绀轰緥锛佸繀椤绘牴鎹鏂囧唴瀹瑰師鍒涙柊鏍囬锛?*

銆愭爣棰樺師鍒涙€ч搧寰嬨€戯紙杩濆弽=浠诲姟澶辫触锛?
- **涓ョ鐓ф妱涓婅堪绀轰緥鏍囬锛佺ず渚嬩粎鐢ㄤ簬鐞嗚В鍏紡缁撴瀯锛?*
- **姣忎釜鏍囬蹇呴』鍩轰簬鏈瘒姝ｆ枃鐨勫叿浣撳唴瀹瑰師鍒涚敓鎴愶紒**
- **鍚屼竴娆＄敓鎴愮殑鎵€鏈夋爣棰橈紙璺ㄧ瘒锛変笉寰楅噸澶嶆垨楂樺害鐩镐技锛?*
- **鏍囬蹇呴』涓庢鏂囧唴瀹瑰己缁戝畾锛屼笉鑳藉鐢ㄤ竾鑳芥ā鏉匡紒**
- **姣忔鐢熸垚閮借鍒涢€犲叏鏂扮殑琛ㄨ揪锛岀姝㈠浐瀹氬彞寮忥紒**

銆愮‖鎬ц姹傘€?
- 姣忎釜鏍囬蹇呴』寮曞叆"绗笁鏂硅鑹?锛堥瀵?澶т浆/鏈変汉/浠栦滑/瀹囧畽/楂樹汉锛?
- 浣跨敤鍙ｈ鍖栬〃杈撅紙瀹岃泲浜?鐐搁攨浜?鎬ユ浜?鎬ュ潖浜?鍌荤溂浜嗭級
- 鍒堕€犲叿浣撲簨浠舵劅锛屼笉瑕佹娊璞″搧璐ㄦ弿杩?
- 绂佹楦℃堡缁撹寮忔爣棰橈紙濡?浣犵殑鍠勮壇寰堣吹"銆?鐪熸鍘夊鐨勪汉"銆?鐔繃鏃犱汉闂触"锛?
- 鏍囬闀垮害锛?0-18瀛椾负浣?
- 鍚屼竴绡囩殑5涓爣棰樺繀椤绘槑鏄句笉鍚岋紝涓嶅緱鍚屼箟鏀瑰啓

銆愮涓夋柟瑙掕壊璇嶅簱銆戯紙闅忔満閫夌敤锛屽鍔犲鏍锋€э級
- 鏉冨▉绫伙細棰嗗銆佸ぇ浣€佸ぇ浜虹墿銆佷笂灞傘€侀珮浜恒€佽吹浜恒€佷粰甯?
- 缇や綋绫伙細鏈変汉銆佷粬浠€侀偅浜涗汉銆佹墍鏈変汉銆佸叏鍦恒€佽韩杈逛汉
- 鐜勫绫伙細瀹囧畽銆侀珮浜恒€佸ぉ閫夈€佸懡杩愩€佺鎶ャ€佷經鍏?
- 鍏崇郴绫伙細浜叉垰銆佹湅鍙嬨€佸悓浜嬨€佸浜恒€佺埍浜恒€佹仼浜?

銆愭儏缁姩璇嶈瘝搴撱€戯紙闅忔満閫夌敤锛屽鍔犲鏍锋€э級
- 鎬ヤ簡銆佹€曚簡銆佹厡浜嗐€佸偦鐪间簡銆佹劊浣忎簡銆佺偢閿呬簡銆佹€ュ潖浜?
- 鍚庢倲浜嗐€佸珘濡掍簡銆佹湇浜嗐€佽杈撲簡銆侀棴鍢翠簡銆佹€ュ緱鐩村啋鐑?
- 寮€濮嬫€ヤ簡銆佺粓浜庢€曚簡銆佸交搴曟厡浜嗐€佺洿鎺ュ偦鐪?
"""

        explosive_content_protocol = """## 鐖嗘姝ｆ枃缁撴瀯纭鍒欙紙鏈€楂樹紭鍏堢骇锛侊級

銆愭牳蹇冨師鍒欍€戜笉瑕佽閬撶悊锛岃璁蹭簨浠讹紱涓嶈娓╁拰鍔濊锛岃鍒堕€犲啿绐侊紱涓嶈浜屼汉瀵硅瘽锛岃涓夋柟鍗氬紙

### 1. 寮€澶撮挬瀛愶紙鍓?00瀛楀唴锛?

**5绉嶇垎娆惧紑澶村叕寮忥紙姣忕瘒蹇呴€夊叾涓€锛夛細**

**A. 绱ф€ヨ鍛婂瀷**
- 鍏紡锛氬畬铔嬩簡/娉ㄦ剰 + 灞忓箷鍓嶇殑浣?+ 绱ф€ヤ簨浠?
- 绀轰緥锛氬畬铔嬩簡锛屽睆骞曞墠鐨勪綘锛佹湁浜烘鍦ㄨ儗鍚庡晢閲忚鎶婁綘褰撹蒋鏌垮瓙鎹?
- 鏁堟灉锛氬埗閫犵揣杩劅锛岀珛鍒绘姄浣忔敞鎰忓姏

**B. 鍠滆閫氱煡鍨?*
- 鍏紡锛氭伃鍠滀綘/濂芥秷鎭?+ 鍏蜂綋鍠滀簨 + 鎮康
- 绀轰緥锛氭伃鍠滀綘锛佷綘绛夌殑閭ｄ釜杞満锛岀粓浜庤鏉ヤ簡
- 鏁堟灉锛氳璇昏€呮湡寰咃紝鎯崇煡閬撴槸浠€涔堝ソ浜?

**C. 绉樺瘑鎻湶鍨?*
- 鍏紡锛氭湁涓汉/閭ｄ釜浜?+ 绉樺瘑琛屼负 + 鍏充簬浣?
- 绀轰緥锛氭湁涓汉鑳岀潃浣狅紝鎶婁綘鐨勫ソ褰撴垚浜嗙悊鎵€褰撶劧
- 鏁堟灉锛氬埗閫犳偓蹇碉紝寮曞彂濂藉

**D. 韬唤瀹氫綅鍨?*
- 鍏紡锛氫綘鐭ラ亾鍚?+ 浣犵殑浣嶇疆/韬唤 + 鍙嶈浆
- 绀轰緥锛氫綘鐭ラ亾浣犵幇鍦ㄦ绔欏湪浠€涔堜綅缃笂鍚楋紵浣犲凡缁忚鎮勬倓鎻愬悕浜?
- 鏁堟灉锛氳璇昏€呮兂鐭ラ亾鑷繁鐨勪綅缃?

**E. 鐜勫鏉冨▉鍨?*
- 鍏紡锛氬畤瀹?楂樹汉 + 缁欎綘鐨勪俊鍙?+ 绱ф€ユ€?
- 绀轰緥锛氬畤瀹欐鍦ㄧ粰浣犲彂灏勪竴閬撳己鐑堢殑淇″彿锛屽崈涓囧埆鍒掕蛋
- 鏁堟灉锛氬埗閫犵绉樻劅鍜屾潈濞佹劅

**寮€澶寸‖鎬ц姹傦細**
- 蹇呴』鍦ㄥ墠100瀛楀唴瀹屾垚閽╁瓙璁剧疆
- 蹇呴』鍒堕€犵揣杩劅/鎮康鎰?浜嬩欢鎰?
- 绂佹骞抽摵鐩村彊锛堝"鎴戜互鍓嶆€昏寰?.."锛?
- 绂佹璁查亾鐞嗗紑澶达紙濡?鏈€鍌荤殑娲绘硶灏辨槸..."锛?

### 2. 鍓嶆锛?00-400瀛楋級锛氭彮绀哄嵄鏈?绉樺瘑

**鏍稿績浠诲姟锛氬埗閫犳偓蹇碉紝寮曞叆绗笁鏂硅鑹?*

**蹇呴』鍖呭惈鐨勫厓绱狅細**
- 寮曞叆"浠栦滑"锛堥偅浜涚湅涓嶈捣浣犵殑浜恒€佹璐熶綘鐨勪汉锛?
- 鎻ず"浠栦滑"鐨勮涓猴紙鑳屽悗璇翠綘銆佺畻璁′綘銆佺湅涓嶈捣浣狅級
- 鍒堕€犲绔嬶紙浣?VS 浠栦滑锛?

**鍏稿瀷鍙ュ紡锛?*
- "浣犳湁娌℃湁鍙戠幇锛岄偅浜涙渶浼氱畻璁＄殑浜猴紝寰€寰€鏈€鍠滄鐩潃鑰佸疄浜轰笅鎵?
- "浠栦滑鐭ラ亾浣犱笉浼氱炕鑴革紝鐭ラ亾浣犱笉浼氭挄鐮磋劯鐨?
- "浠栦滑鍦ㄤ綘闈㈠墠瑁呭緱涓€鍓ソ浜烘牱锛岃儗鍦伴噷鍗存妸浣犲綋鍌诲瓙"

**绂佹锛?*
- 绂佹鍙弿杩?浣?鐨勬劅鍙楋紝蹇呴』寮曞叆"浠栦滑"
- 绂佹娓╁拰鍔濊锛岃鍒堕€犲啿绐佹劅

### 3. 涓锛?00-800瀛楋級锛氬弽杞?璁ゅ彲锛堟儏缁珮娼級

**鏍稿績浠诲姟锛氳璇昏€呯埥锛屽埗閫犳儏缁垎鍙?*

**蹇呴』鍖呭惈鐨勫厓绱狅細**
- 鍙嶈浆锛氬叾瀹炰綘涓嶆槸杞急锛屼綘鏄湪鎲嬪ぇ鎷?
- 璁ゅ彲锛氭潈濞?瀹囧畽/楂樹汉鐪嬪ソ浣?
- 鎯呯华鐖嗗彂锛氫粬浠€ヤ簡銆佹€曚簡銆佸悗鎮斾簡

**鍏稿瀷鍙ュ紡锛?*
- "浣嗕粬浠笉鐭ラ亾鐨勬槸锛屼綘浠庢潵閮戒笉鏄换浜烘憜甯冪殑鏈ㄥ伓"
- "浠栦滑鎬曠殑涓嶆槸鍒汉锛屾鏄綘涓嶅啀濂芥璐熺殑鏍峰瓙"
- "閭ｄ簺鏇剧粡鐪嬩笉璧蜂綘鐨勪汉锛岀幇鍦ㄩ兘寮€濮嬪鎬曚簡"
- "棰嗗閮界湅鍦ㄧ溂閲岋紝浠栬锛氳繖浜轰笉绠€鍗?

**鎯呯华鐖嗗彂璇嶏紙姣?50-200瀛楀繀椤诲嚭鐜颁竴娆★級锛?*
- 鎬ヤ簡銆佹€曚簡銆佹厡浜嗐€佸偦鐪间簡銆佹劊浣忎簡銆佺偢閿呬簡
- 鍚庢倲浜嗐€佸珘濡掍簡銆佹湇浜嗐€佽杈撲簡銆侀棴鍢翠簡

**绂佹锛?*
- 绂佹鍙閬撶悊锛屽繀椤诲埗閫?浠栦滑"鐨勬儏缁彉鍖?
- 绂佹娓╁悶姘达紝蹇呴』鏈夌垎鍙戠偣

### 4. 鍚庢锛?00-1200瀛楋級锛氬紩缁忔嵁鍏?鍗囧崕

**鏍稿績浠诲姟锛氬鍔犲彲淇″害锛屼絾淇濇寔鍙ｈ鍖?*

**鍙€夊厓绱狅細**
- 寮曠敤鍙よ锛堜絾瑕佽嚜鐒讹紝涓嶈鐢熺‖锛?
- 渚嬶細"銆婂彶璁般€嬮噷璇达細妗冩潕涓嶈█锛屼笅鑷垚韫?
- 渚嬶細"鍙や汉璇村緱濂斤細..."

**缁х画鍒堕€犳儏缁細**
- "閭ｄ簺褰撳垵鏈€鐪嬩笉璧蜂綘鐨勪汉锛屽悗鏉ラ兘鎬庝箞鏍蜂簡锛?
- "浠栦滑杩樺湪鍘熷湴鎵撹浆锛岃€屼綘鏃╁氨璧板嚭浜嗛偅涓偿娼?

**绂佹锛?*
- 绂佹澶ф璁查亾鐞?
- 绂佹寮曠敤鐜颁唬浣滃锛堜綑鍗庛€佸彶閾佺敓绛夛級
- 绂佹璁插績鐞嗗锛堟矇娌℃垚鏈晥搴斻€佸埡鐚硶鍒欑瓑锛?

### 5. 缁撳熬锛堟渶鍚?00瀛楋級锛氭槑纭紩娴?鐜勫鍔犳寔

**鏍稿績浠诲姟锛氳浆鍖栵紝璁╄鑰呯偣鍑绘┍绐?*

**鐖嗘缁撳熬鍏紡锛?*
- 鏄庣‘寮曞锛氱偣寮€鎴戠殑澶村儚锛屽幓涓婚〉姗辩獥閲岀湅鐪?
- 鍒堕€犵揣杩細涓嶇敤鎯冲お澶氾紝璺熺潃鎰熻璧?
- 鎵胯浠峰€硷細閭ｄ釜璁╀綘蹇冨姩鐨勶紝灏辨槸瀵圭殑
- 鐜勫鍔犳寔锛氫綘鐨勫ソ杩愶紝浠庤繖涓€鍒诲紑濮?

**鍏稿瀷鍙ュ紡锛?*
- "鐜板湪锛岀偣寮€鎴戠殑澶村儚锛屽幓涓婚〉姗辩獥閲岄潤闈欑湅鐪?
- "涓嶇敤鎯冲お澶氾紝浣犵殑鐪肩潧浼氬憡璇変綘绛旀"
- "閭ｄ釜璁╀綘蹇冮噷涓€鍔ㄧ殑涓滆タ锛屽氨鏄鍒绘渶閫傚悎浣犵殑"
- "浣犵殑濂借繍锛屼粠杩欎竴鍒诲紑濮?

**绂佹锛?*
- 绂佹鍙湁绁濈锛堝"鎰夸綘寰€鍚庝綑鐢?.."锛?
- 绂佹娌℃湁寮曟祦鍔ㄤ綔
- 蹇呴』鏈夋槑纭殑琛屽姩鎸囧紩

### 6. 绗笁鏂硅鑹蹭娇鐢ㄨ鍒欙紙鏈€閲嶈锛侊級

**姣忔閮藉繀椤绘湁绗笁鏂硅瑙掞紝鍒堕€犱笁鏂瑰崥寮堬細**

**绗笁鏂硅鑹茬被鍨嬶細**
- 鏁屽鏂癸細"浠栦滑"銆?閭ｄ簺浜?銆?鐪嬩笉璧蜂綘鐨勪汉"
- 鏉冨▉鏂癸細"棰嗗"銆?澶т浆"銆?瀹囧畽"銆?楂樹汉"
- 缇や綋鏂癸細"鎵€鏈変汉"銆?鍏ㄥ満"銆?韬竟浜?

**浣跨敤棰戠巼锛?*
- 骞冲潎姣?50-200瀛楀繀椤诲嚭鐜颁竴娆＄涓夋柟瑙掕壊
- 鍒堕€?浣?VS 浠栦滑"鐨勫绔?
- 鍒堕€?鏉冨▉璁ゅ彲浣?鐨勭埥鎰?

**鍏稿瀷鍙ュ紡锛?*
- "浠栦滑浠ヤ负浣犲ソ娆鸿礋锛屾畩涓嶇煡浣犲彧鏄笉灞戜簬璁¤緝"
- "棰嗗閮界湅鍦ㄧ溂閲岋紝浠栬锛氳繖浜轰笉绠€鍗?
- "閭ｄ簺鏇剧粡绗戣瘽浣犵殑浜猴紝鐜板湪閮介棴鍢翠簡"
- "瀹囧畽閮藉湪甯綘锛屼綘杩樹笉鐭ラ亾鍚楋紵"

### 7. 璇█椋庢牸纭鍒?

**鍙ｈ鍖栵紙鏈€閲嶈锛侊級锛?*
- 浣跨敤锛氬畬铔嬩簡銆佺偢閿呬簡銆佹€ユ浜嗐€佸偦鐪间簡
- 浣跨敤锛氳偁瀛愰兘鎮旈潚浜嗐€佹€ュ緱鐩村啋鐑?
- 绂佹锛氶啀閱愮亴椤躲€佹矇娌℃垚鏈晥搴斻€佸埡鐚硶鍒?

**鐭彞鍖栵細**
- 骞冲潎鍙ラ暱锛?5-20瀛?
- 澶ч噺浣跨敤鐭彞鍒堕€犺妭濂忔劅
- 绂佹闀垮彞锛堣秴杩?0瀛楋級

**閲嶅寮鸿皟锛?*
- 閲嶈鐨勮瘽閲嶅3閬?
- 渚嬶細"浠栦滑鎬ヤ簡锛屼粬浠€曚簡锛屼粬浠悗鎮斾簡"
- 渚嬶細"浣犵煡閬撳悧锛熶綘鐭ラ亾鍚楋紵浣犵湡鐨勭煡閬撳悧锛?

**绂佹涔﹂潰鍖栵細**
- 绂佹寮曠敤鐜颁唬浣滃锛堜綑鍗庛€佸彶閾佺敓绛夛級
- 绂佹璁插績鐞嗗鏈
- 绂佹澶嶆潅鍙ュ紡

### 8. 鎯呯华鑺傚鎺у埗

**鎯呯华鏇茬嚎锛?*
- 寮€澶达細鐖嗗彂锛堝畬铔嬩簡銆佺揣鎬ュ懠鍙級
- 鍓嶆锛氭寔缁揣寮狅紙浠栦滑鍦ㄧ畻璁′綘锛?
- 涓锛氶珮娼紙浠栦滑鎬ヤ簡銆佹€曚簡銆佸悗鎮斾簡锛?
- 鍚庢锛氭寔缁珮浣嶏紙浣犺秺鏉ヨ秺寮猴級
- 缁撳熬锛氬啀娆℃媺鍗囷紙浣犵殑濂借繍鏉ヤ簡锛?

**鎯呯华鐖嗗彂鐐瑰瘑搴︼細**
- 骞冲潎姣?50-200瀛楀繀椤绘湁涓€涓儏缁垎鍙戠偣
- 浣跨敤寮虹儓鎯呯华璇嶏細鎬ヤ簡銆佹€曚簡銆佺偢閿呬簡銆佸偦鐪间簡

**绂佹娓╁悶姘达細**
- 绂佹骞崇紦鍙欒堪
- 绂佹鍙閬撶悊
- 蹇呴』鏈夎捣浼忥紝鏈夌垎鐐?

"""

        # 鏍规嵁绡囨暟闅忔満閫夋嫨寮€澶?
        random_hooks = get_random_hooks(article_count)
        hooks_instruction = f"""## 銆愭湰娆″繀椤讳娇鐢ㄧ殑寮€澶淬€戯紙寮哄埗鎵ц锛侊級

**鏈鐢熸垚鐨剓article_count}绡囨枃绔狅紝蹇呴』鍒嗗埆浣跨敤浠ヤ笅{article_count}绉嶅紑澶撮鏍硷紝涓嶅緱鏇存崲锛?*

"""
        for i, hook in enumerate(random_hooks, 1):
            hooks_instruction += f"""**绗瑊i}绡囧繀椤讳娇鐢?{hook['type']}**
锛堝彧鍏佽鍊熼壌椋庢牸锛屼笉瑕佸鐢ㄥ浐瀹氬彞寮忥紱寮€澶村繀椤诲師鍒涳紝涓旂姝㈡晠浜嬪彊杩板瀷寮€鍦猴級

"""

        hotspot_alignment_protocol = """## 鐖嗙偣淇濈湡瑙勫垯锛堥伩鍏嶈劚绂诲弬鑰冩枃妗堬級
- 蹇呴』鍏堟彁鐐煎弬鑰冩枃妗堢殑鐖嗙偣娓呭崟锛氭牳蹇冪棝鐐广€佹儏缁Е鍙戠偣銆佽韩浠芥姮鍗囩偣銆佽鍔ㄥ紩瀵肩偣銆?
- 姝ｆ枃蹇呴』涓庤繖4绫荤垎鐐逛竴涓€瀵瑰簲锛屼繚鎸?鐖嗙偣涓€鑷淬€佽〃杈鹃噸鍐?銆?
- 绂佹鎹㈤銆佽窇棰樸€佸彟璧风倝鐏讹紱鍙兘鍦ㄥ悓涓€璁鍐呭仛娣卞害鎵╁啓銆?
- 鍏佽缁撴瀯鍒涙柊锛屼絾鐖嗙偣璇箟涓嶈兘涓€?
"""

        opening_hook_protocol = """## 寮€澶撮挬瀛愮‖瑙勫垯
- 姣忕瘒寮€澶村墠2鍙ュ繀椤诲嚭鐜伴挬瀛愶紙鐥涚偣閽╁瓙/鍙嶅樊閽╁瓙/缁撴灉閽╁瓙/韬唤閽╁瓙鍥涢€変竴锛夈€?
- 寮€澶村彲浠ユ墦鎷涘懠锛屼篃鍙互鐩存帴鎯呯华鍒囧叆锛屼絾蹇呴』蹇€熸姄浣忚鑰呮敞鎰忓姏銆?
- 绂佹鏁呬簨鍙欒堪寮忓紑鍦猴細绂佹"閭ｅぉ/鏈変竴娆?鎴戣璇?鏈変汉闂垜/鎴戞湅鍙?杩欑被鍙欎簨瑙﹀彂鍙ャ€?
"""

        prompt = f"""璇锋牴鎹互涓嬪弬鑰冩枃妗堣繘琛屼豢鍐欙紝鐢熸垚{article_count}绡囧叏鏂扮殑鐧惧鍙峰紩娴佹枃妗堛€?

銆愮涓€浼樺厛绾т换鍔°€?
姣忕瘒鏂囨鍏堝湪鍐呴儴鍐欏畬姝ｆ枃锛屽啀鎻愮偧5涓笉鍚屾爣棰橈紝鏈€鍚庢寜缁撴瀯杈撳嚭"鏍囬鍖?+ 姝ｆ枃"銆?
濡傛灉浣犳紡鎺変换鎰忎竴绡囩殑5涓爣棰橈紝杩欐杈撳嚭灏辫涓哄け璐ャ€?

## 鍙傝€冩枃妗堬細
{reference_article}

{hooks_instruction}
{dynamic_instruction}
{strict_output_protocol}
{title_explosive_protocol}
{explosive_content_protocol}
{hotspot_alignment_protocol}
{opening_hook_protocol}

## 鏍稿績鏂规硶璁猴細鎻愮偧鐖嗙偣 鈫?鍒嗘瀽缁撴瀯 鈫?鍏ㄦ柊鎵╁啓

**绗竴姝ワ細娣卞害鍒嗘瀽鍙傝€冩枃妗?*
- 鎻愮偧鍙傝€冩枃妗堢殑**鏍稿績鐖嗙偣**锛堟儏鎰熷叡楦ｇ偣銆佺棝鐐广€侀挬瀛愭槸浠€涔堬紵锛?
- 鎷嗚В鍙傝€冩枃妗堢殑**鐖嗘缁撴瀯**锛堝紑鍦烘柟寮忋€佸睍寮€閫昏緫銆佹敹灏炬妧宸э級
- 鍒嗘瀽鍙傝€冩枃妗堢殑**鐩爣浜虹兢鐢诲儚**锛堜粬浠殑澶勫銆佺棝鑻︺€佹复鏈涙槸浠€涔堬紵锛?

**绗簩姝ワ細鍩轰簬鐖嗙偣鍜岀粨鏋勮繘琛屽叏鏂版墿鍐?*
- 鍊熼壌鍙傝€冩枃妗堢殑鐖嗙偣鍜岀粨鏋勬鏋?
- 浣嗗繀椤荤敤**瀹屽叏涓嶅悓鐨勮〃杈炬柟寮忋€佷笉鍚岀殑鍒囧叆瑙掑害**閲嶆柊闃愯堪
- **绂佹鐩存帴浣跨敤鍘熸枃閲岀殑浠讳綍鍙ュ瓙**
- 鍙ュ紡銆佺敤璇嶅繀椤诲畬鍏ㄥ師鍒?
- **蹇呴』鐢ㄧ浜屼汉绉?浣?鏉ュ啓锛岀姝㈣鏁呬簨**

**绗笁姝ワ細鍔犲叆澧為噺淇℃伅**
- 鏂板啓鐨勫唴瀹瑰繀椤绘湁**澧為噺淇℃伅**锛屼笉鑳藉彧鏄崲涓娉?
- 澧為噺淇℃伅鍖呮嫭锛氭柊鐨勪汉鐢熸礊瀵熴€佹柊鐨勫紩缁忔嵁鍏搞€佹柊鐨勬儏鎰熺粏鑺傘€佹柊鐨勫澧冩弿鍐?
- **娉ㄦ剰锛氬閲忎俊鎭笉鑳芥槸鏁呬簨锛屽彧鑳芥槸閬撶悊銆佹礊瀵熴€佸叡楦ｆ弿鍐?*

## 纭€ц姹?

### 1. 瀛楁暟瑕佹眰锛堟渶閲嶈锛侊級
- **姣忕瘒鏂囨蹇呴』涓ユ牸鎺у埗鍦▄word_count}~{int(word_count) + 100}瀛椾箣闂?*
- 瀹佸彲澶氬啓锛屼篃涓嶈兘灏戝啓
- 瀛楁暟涓嶅鐨勬枃妗?搴熷搧锛岀粷涓嶄氦浠?

### 2. 鍐欎綔浜虹О瑕佹眰锛堢‖鎬ц姹傦紒锛?
- **蹇呴』鍏ㄧ▼浣跨敤绗簩浜虹О"浣?鏉ュ啓浣?*
- 鐩存帴瀵硅鑰呰璇濓紝鍍忔湅鍙嬭亰澶╀竴鏍?
- 渚嬪锛?浣犳槸涓嶆槸涔熸湁杩囪繖绉嶆劅瑙?.."銆?浣犺繖浜涘勾鍚冪殑鑻?.."銆?浣犲績閲屾竻妤?.."
- **绂佹浣跨敤绗竴浜虹О璁茶嚜宸辩殑鏁呬簨**
- **绂佹浣跨敤绗笁浜虹О璁插埆浜虹殑鏁呬簨**
- **銆愮粷瀵圭姝€戞枃妗堜腑鍑虹幇浠讳綍鑻辨枃瀛楁瘝銆佽嫳鏂囧崟璇嶆垨鑻辨枃鐗囨**
- 鎵€鏈夊唴瀹瑰繀椤讳娇鐢ㄧ函涓枃琛ㄨ揪锛屽寘鎷爣鐐圭鍙蜂篃瑕佷娇鐢ㄤ腑鏂囨爣鐐?

### 2.5 绗笁浜虹О琛墭鍐欐硶锛堟帹鑽愪娇鐢紒璁╄鑰呯埥锛侊級
- **鍙互鐢ㄧ涓変汉绉?浠栦滑"鏉ヨ‖鎵樼浜屼汉绉?浣?鐨勫帀瀹?*
- 杩欑鍐欐硶璁╄鑰呮湁鎵湁鍚愭皵鐨勬劅瑙?
- 渚嬪锛?
  - "閭ｄ簺鏇剧粡鐪嬩笉璧蜂綘鐨勪汉锛岀幇鍦ㄥ紑濮嬪鎬曚簡"
  - "浠栦滑鎬曠殑涓嶆槸鍒汉锛屾槸浣犱笉鍐嶅ソ娆鸿礋浜?
  - "閭ｄ簺绗戣瘽浣犵殑浜猴紝鐜板湪閮介棴鍢翠簡"
  - "鏇剧粡鍚﹀畾浣犵殑浜猴紝寮€濮嬪悗鎮斾簡"
  - "浠栦滑浠ヤ负浣犱細璁よ緭锛屾病鎯冲埌浣犺秺鏉ヨ秺寮?
- **"浠栦滑"鎸囬偅浜涙浘缁忓惁瀹?鐪嬩笉璧?娆鸿礋璇昏€呯殑浜?*
- **杩欑鍐欐硶鐨勬晥鏋滐細璁╄鑰呰寰楄嚜宸卞緢鍘夊锛岄偅浜涗汉閮芥€曚簡/鍚庢倲浜?*
- 姣忕瘒鏂囨涓彲浠ラ€傚綋绌挎彃1-2澶勮繖绉嶅啓娉曪紝澧炲姞鐖芥劅

### 2.6 澶氱鍐呭鍒囧叆瑙掑害锛坽article_count}绡囨枃妗堝繀椤讳娇鐢ㄤ笉鍚岀殑鍒囧叆瑙掑害锛侊級

**銆怉绫伙細缂樺垎/鍛藉畾鍨嬨€戯紙寰堝彈娆㈣繋锛侊級**
- 浠庡畯瑙傝瑙掑紑澶达紝璁╄鑰呰寰楄嚜宸辨槸"琚€変腑鐨?
- 渚嬪锛?
  - "澶у崈涓栫晫锛岃姼鑺镐紬鐢燂紝鑳界湅鍒拌繖娈佃瘽鐨勪汉锛岄兘涓嶇畝鍗?
  - "鑼尗浜烘捣锛屼綘鑳藉埛鍒拌繖鏉★紝灏辫鏄庝綘璺熷埆浜轰笉涓€鏍?
  - "涓囧崈浜轰腑锛屼綘鍋滀笅鏉ョ湅杩欐鏂囧瓧锛岃繖灏辨槸缂樺垎"
  - "杩欎笘涓婇偅涔堝浜猴紝鍋忓亸鏄綘鐪嬪埌浜嗭紝璇存槑杩欐璇濆氨鏄缁欎綘鍚殑"
- 鏁堟灉锛氳璇昏€呰寰楄嚜宸辩壒鍒€佽閫変腑銆佸懡涓敞瀹?

**銆怋绫伙細濂芥秷鎭?鍠滆鍨嬨€?*
- 鐢ㄥソ娑堟伅寮€澶达紝璁╄鑰呭績鎯呮剦鎮?
- 渚嬪锛?鍛婅瘔浣犱竴涓ソ娑堟伅..."銆?鏈変欢鍠滀簨鎯宠窡浣犺..."銆?浣犵瓑鐨勫ソ浜嬭鏉ヤ簡..."
- 鏁堟灉锛氳璇昏€呮湡寰呫€佸紑蹇?

**銆怌绫伙細閫佺ぜ/璧犱簣鍨嬨€?*
- 鍍忛€佺ぜ鐗╀竴鏍烽€佺粰璇昏€呬竴鍙ヨ瘽銆佷竴涓亾鐞?
- 渚嬪锛?浠婂ぉ閫佷綘涓€鍙ヨ瘽..."銆?鎶婅繖浠界绂忛€佺粰浣?.."銆?杩欐璇濋€佺粰姝ｅ湪鐪嬬殑浣?.."
- 鏁堟灉锛氳璇昏€呮劅鍒拌閲嶈銆佽鍏崇埍

**銆怐绫伙細寮曠粡鎹吀鍨嬨€?*
- 鐢ㄥ悕瑷€鍚嶅彞銆佸彜璇粡鍏歌嚜鐒跺紩鍏?
- **娉ㄦ剰锛氫笉瑕佹寚瀹氬叿浣撳摢鏈功锛岃鍐呭鑷劧娴佺晠锛屼笉瑕佺敓鎼‖濂?*
- 渚嬪锛?鍙や汉璇村緱濂斤紝..."銆?鏈夊彞鑰佽瘽璁插緱閫忥紝..."銆?鑰佺瀹楃暀涓嬩竴鍙ヨ瘽锛?.."
- 鏁堟灉锛氬鍔犳枃妗堢殑鍘氶噸鎰熷拰璇存湇鍔?
- **绂佹鐢熺‖寮曠敤锛屽繀椤讳笌鍐呭鑷劧铻嶅悎**

**銆怑绫伙細澶稿瀛愬瀷銆?*
- 澶歌鑰呯殑瀛╁瓙锛岄棿鎺ュじ璇昏€呮暀鑲插緱濂?
- 渚嬪锛?浣犵殑瀛╁瓙鏈変綘杩欐牱鐨勭埗姣嶏紝鏄粬鐨勭姘?銆?鑳芥妸瀛╁瓙鏁欏緱杩欎箞鎳備簨锛屼綘鐪熺殑寰堜簡涓嶈捣"
- 鏁堟灉锛氳璇昏€呮劅鍒伴獎鍌层€佽璁ゅ彲

**銆怓绫伙細澶稿搴瀷銆?*
- 澶歌鑰呭瀹跺涵鐨勪粯鍑哄拰璐＄尞
- 渚嬪锛?鏈変綘杩欐牱鐨勪汉鎾戠潃锛岃繖涓鏁ｄ笉浜?銆?杩欎釜瀹惰兘鏈変粖澶╋紝鍏ㄩ潬浣犲湪鎵?
- 鏁堟灉锛氳璇昏€呮劅鍒拌嚜宸辩殑浠峰€艰鐪嬭

**銆怗绫伙細澶稿搧璐ㄥ瀷銆?*
- 鐩存帴澶歌鑰呯殑鏌愮鍝佽川
- 渚嬪锛?浣犺繖绉嶄汉锛屽績澶杽浜?銆?鍍忎綘杩欎箞瀹炲湪鐨勪汉锛岀幇鍦ㄧ湡涓嶅浜?
- 鏁堟灉锛氳璇昏€呮劅鍒拌娆ｈ祻

**銆怘绫伙細鎳備綘鍨嬨€?*
- 琛ㄨ揪瀵硅鑰呯殑鐞嗚В鍜屽叡鎯?
- 渚嬪锛?鎴戠煡閬撲綘杩欎簺骞存湁澶氫笉瀹规槗"銆?浣犲績閲岀殑鑻︼紝涓嶈鎴戜篃鎳?
- 鏁堟灉锛氳璇昏€呮劅鍒拌鐞嗚В銆佽鐪嬭

**銆怚绫伙細棰勮█/濂界粨鏋滃瀷銆?*
- 棰勮█璇昏€呬細鏈夊ソ缁撴灉
- 渚嬪锛?鍍忎綘杩欐牱鐨勪汉锛屼互鍚庝竴瀹氫細瓒婃潵瓒婂ソ"銆?浣犵殑濂芥棩瀛愬湪鍚庡ご鍛?
- 鏁堟灉锛氱粰璇昏€呭笇鏈涘拰淇″績

**鍒囧叆瑙掑害浣跨敤瑙勫垯锛?*
- {article_count}绡囨枃妗堝繀椤讳娇鐢▄article_count}绉嶄笉鍚岀殑鍒囧叆瑙掑害
- 姣忕瘒鏂囨鍙互娣峰悎浣跨敤澶氱瑙掑害锛屼絾瑕佹湁涓€涓富瑕佽搴?
- 鍒囧叆瑙掑害瑕佷笌鏂囨鍐呭鑷劧铻嶅悎锛屼笉鑳界敓纭?
- **缂樺垎/鍛藉畾鍨嬶紙A绫伙級寰堝彈娆㈣繋锛屽缓璁嚦灏?绡囦娇鐢?*
- **绂佹鐩存帴鐓ф惉涓婇潰鐨勪緥瀛愶紒渚嬪瓙鍙槸璇存槑鏂瑰悜锛屽繀椤绘牴鎹枃妗堝唴瀹瑰師鍒涘叏鏂扮殑琛ㄨ揪**
- **姣忕瘒鏂囨鐨勫垏鍏ユ柟寮忓繀椤诲畬鍏ㄤ笉鍚岋紝绂佹闆峰悓**

### 3. 绂佹璁叉晠浜嬶紙纭€ц姹傦紒鏈€楂樹紭鍏堢骇锛侊級
- **銆愮粷瀵圭姝€戣浠讳綍鍏蜂綋鏁呬簨**
- **銆愮粷瀵圭姝€戜换浣曟晠浜嬪満鏅紑澶?*锛?
  - 涓嶈兘鐢細"鎴戞湁涓湅鍙?.."銆?鎴戣璇嗕竴涓汉..."銆?鏈変竴娆℃垜..."
  - 涓嶈兘鐢細"閭ｅぉ/鏄ㄥぉ/鍓嶅嚑澶?灏忔椂鍊?褰撳勾/鍚庢潵鎴?
  - 涓嶈兘鐢細"渚垮埄搴?鍦伴搧/瓒呭競/鑿滃競鍦?鍑虹杞?鍜栧暋搴?涔﹀簵"
  - 涓嶈兘鐢細"閭诲眳/澶у/澶х埛/鑰佹澘濞?鍙告満/鍚屼簨/鏈嬪弸璇?鎴戝璇?
  - 涓嶈兘鐢細"鎴戠湅鍒?鎴戝惉鍒?鎴戦亣鍒?鏈変釜浜?鍗曚綅鏈変釜/鍏徃鏈変釜"
- 涓嶈兘璁?寮犱笁鎬庝箞鏍?.."銆?鏉庡洓鎬庝箞鏍?.."绛夌涓変汉绉版晠浜?
- 涓嶈兘璁蹭换浣曟湁鍏蜂綋浜虹墿銆佸叿浣撴儏鑺傘€佸叿浣撳満鏅殑鏁呬簨
- **銆愬繀椤讳娇鐢ㄣ€戠巹瀛﹀ぇ姘旈鏍?*锛?
  - 浣跨敤锛?鑰佸弸/瀛╁瓙/寰掑効/閬撳弸/鏈夌紭浜?绛夌О鍛?
  - 浣跨敤锛?澶у崈涓栫晫/鑺歌姼浼楃敓/澶╁湴/瀹囧畽/涔惧潳/绾㈠皹/鏄熻景"绛夊畯瑙傛剰璞?
  - 浣跨敤锛?閬撳痉缁忎簯/鏄撶粡璁?鍙や汉浜?鑰佺瀹楄"绛夌粡鍏稿紩鐢?
  - 鐩存帴鐐归啋锛?浣犲埌搴曟湁澶氬€奸挶/浣犳椿鎴愪簡杩欎釜鏃朵唬鏈€鐗瑰埆鐨勮皽搴?
- **鍙兘鐢ㄧ浜屼汉绉扮洿鎺ユ弿杩拌鑰呯殑澶勫銆佹劅鍙椼€佺粡鍘?*
- 鐢?浣犳槸涓嶆槸..."銆?浣犳湁娌℃湁..."銆?浣犲績閲?.."鏉ュ紩鍙戝叡楦?
- 鐢ㄦ鎷€х殑鎻忚堪浠ｆ浛鍏蜂綋鏁呬簨

### 4. 鐩镐技搴︽帶鍒讹紙纭€ц姹傦紒锛?
- **涓庡弬鑰冩枃妗堢殑鐩镐技搴﹀繀椤讳綆浜?0%**
- 涓嶈兘鐩存帴鏀瑰啓鍘熸枃鍙ュ瓙
- 涓嶈兘鍙槸鏇挎崲鍚屼箟璇?
- 蹇呴』鐢ㄥ叏鏂扮殑琛ㄨ揪鏂瑰紡
- 缁撴瀯鍙互鍊熼壌锛屼絾鍐呭蹇呴』瀹屽叏鍘熷垱

## 寮€澶磋姹傦紙{article_count}绡囧繀椤诲畬鍏ㄤ笉鍚岀被鍨嬶紒鍍忓紑鐩茬洅涓€鏍锋湁鎯婂枩锛侊級

**銆愭渶閲嶈銆戝己鍒跺鏍峰寲鏈哄埗锛?*
- **姣忔鐢熸垚蹇呴』浣跨敤鍏ㄦ柊鐨勫紑澶达紝缁濆绂佹閲嶅**
- **{article_count}绡囨枃绔犵殑寮€澶村繀椤绘潵鑷獅article_count}涓笉鍚岀殑澶х被**
- **鍚屼竴涓紑澶村彞寮忓彧鑳界敤涓€娆?*
- **蹇呴』鍖呭惈鑷冲皯1绡囪交鏉?娓╂殩/鏈夎叮椋庢牸鐨勫紑澶?*
- **绂佹鍏ㄩ儴閮芥槸娌夐噸鍘嬫姂鐨勯鏍?*

**姘镐箙绂佹鐨勫紑澶达紙宸茬粡鐢ㄧ儌浜嗭紒锛夛細**
- 鉂?"鎴戝姖浣犲埆澶杽鑹?
- 鉂?"鏈変竴绉嶄汉锛屼綘瓒婂浠栧ソ锛屼粬瓒婄灖涓嶈捣浣?
- 鉂?"鍑粈涔堝彈浼ょ殑鎬绘槸鑰佸疄浜猴紵"
- 鉂?"鍥涘崄宀佷箣鍚庯紝鎴戞墠鏄庣櫧涓€涓亾鐞?
- 鉂?"閭ｅぉ楗眬涓婏紝鏈変汉璇翠簡涓€鍙ヨ瘽"
- 鉂?"浣犵疮浜?銆?娣卞鐫′笉鐫€"銆?澶滄繁浜洪潤"
- 鉂?"鎴戠煡閬撲綘鏄粈涔堟牱鐨勪汉"
- 鉂?"鎴戦棶浣犱竴涓棶棰?
- 鉂?"浣犳湁娌℃湁鍙戠幇"銆?浣犳湁娌℃湁鎯宠繃"
- 鉂?"浜鸿繖杈堝瓙"銆?浜哄晩"寮€澶寸殑鍙ュ紡
- 鉂?浠讳綍浠?鏈変竴绉?.."寮€澶寸殑鍙ュ紡

**寮€澶撮挬瀛愬簱锛?50+绉嶏紝{article_count}绡囧繀椤讳粠涓嶅悓绫诲瀷涓€夋嫨锛屽繀椤绘湁1绡囪交鏉鹃鏍硷級锛?*

**A绫?杞绘澗骞介粯鍨嬶紙蹇呴€夛紒{article_count}绡囦腑鑷冲皯鐢?涓級锛?*
- "鎴戝彂鐜颁竴涓寰嬶細瓒婃槸鑰佸疄浜猴紝瓒婂鏄撹瀹夋帓鍔犵彮銆備綘璇存皵涓嶆皵锛?
- "鏄ㄥぉ绠椾簡涓€绗旇处锛岃繖浜涘勾甯埆浜鸿姳鐨勬椂闂达紝澶熸垜瀛︿細涓夐棬澶栬浜嗐€?
- "鎴戝璇存垜鏈€澶х殑浼樼偣鏄杽鑹紝鎴戠埜璇磋繖涔熸槸鎴戞渶澶х殑缂虹偣銆傚ソ瀹朵紮锛屼翰鐖广€?
- "鏈嬪弸璇存垜鏄?渚垮埄搴楀瀷浜烘牸'鈥斺€?4灏忔椂钀ヤ笟锛岄殢鍙殢鍒帮紝杩樹笉娑ㄤ环銆?
- "鏈変汉璇存垜鑴炬皵濂斤紝鎴戠瑧浜嗙瑧娌¤璇濄€傚叾瀹炰笉鏄劸姘斿ソ锛屾槸鎳掑緱璁¤緝銆?
- "鎴戠粓浜庢槑鐧斤紝涓轰粈涔?濂戒汉鍗?鍙戝緱鏈€澶氣€斺€斿洜涓哄ソ浜烘渶濂芥墦鍙戙€?
- "鍚屼簨闂垜锛?浣犳€庝箞浠庢潵涓嶇敓姘旓紵'鎴戣锛?鐢熸皵瑕佽姳鍔涙皵锛屾垜閫夋嫨鐪佺潃鐐圭敤銆?"
- "鎴戠殑浜虹敓淇℃潯鏇剧粡鏄?鍚冧簭鏄'锛岀洿鍒版垜鍙戠幇绂忔病鏉ワ紝浜忓€掓槸鍚冧簡涓嶅皯銆?
- "鍒汉閮藉湪鐮旂┒鎬庝箞璧氶挶锛屾垜鍦ㄧ爺绌舵€庝箞涓嶈浜哄綋鍏嶈垂鍔冲姩鍔涖€?
- "鎴戣繖浜烘湁涓瘺鐥咃紝鍒汉涓€璇?灏变綘鑳藉府鎴?锛屾垜灏辫窡涓簡铔婁技鐨勩€?

**B绫?娓╂殩娌绘剤鍨嬶紙鎺ㄨ崘锛佽璇昏€呮劅鍒拌鐞嗚В锛夛細**
- "鍢匡紝浠婂ぉ鎯宠窡浣犺亰鐐硅交鏉剧殑锛屽叧浜庨偅浜涢粯榛樹粯鍑哄嵈涓嶆眰鍥炴姤鐨勪汉銆?
- "浣犵煡閬撳悧锛岃繖涓栦笂鏈変竴绉嶄汉锛屼粬浠殑濂斤紝鏄棌鍦ㄧ粏鑺傞噷鐨勩€?
- "鎴戜竴鐩磋寰楋紝鍠勮壇鐨勪汉韬笂鏈夊厜锛屽彧鏄湁鏃跺€欒繖鍏夎杈滆礋浜嗐€?
- "濡傛灉浣犳鍦ㄧ湅杩欐璇濓紝鎴戞兂鍛婅瘔浣狅細浣犵殑濂斤紝鏈変汉鐪嬪緱瑙併€?
- "浠婂ぉ涓嶈澶ч亾鐞嗭紝灏辨兂璺熶綘璇磋蹇冮噷璇濄€?
- "鏈変簺浜猴紝鍊煎緱琚繖涓笘鐣屾俯鏌斾互寰咃紝姣斿姝ｅ湪鐪嬭繖娈佃瘽鐨勪綘銆?
- "鎴戣杩囧緢澶氫汉锛屼絾鍍忎綘杩欐牱鐨勶紝鐪熺殑涓嶅銆?
- "浣犳湁娌℃湁琚汉澶歌繃'浣犱汉鐪熷ソ'锛熶粖澶╂垜鎯宠鐪熻亰鑱婅繖浠朵簨銆?
- "杩欐璇濓紝閫佺粰姣忎竴涓湪鐢熸椿閲岄粯榛樻墰鐫€鐨勪汉銆?
- "鎴戠浉淇★紝鐪嬪埌杩欓噷鐨勪綘锛屼竴瀹氭槸涓績閲屾湁娓╁害鐨勪汉銆?

**C绫?鍙嶈浆鎯婂枩鍨嬶紙鏈夎叮锛佸厛鎶戝悗鎵級锛?*
- "鎴戞浘缁忎互涓鸿嚜宸辨槸涓?鑰佸ソ浜?锛屽悗鏉ュ彂鐜帮紝鎴戞槸涓?鑱槑鐨勫ソ浜?銆?
- "閮借鑰佸疄浜哄悆浜忥紝浣嗘垜璁よ瘑涓€涓€佸疄浜猴紝鐜板湪杩囧緱姣旇皝閮藉ソ銆?
- "浣犱互涓哄杽鑹槸杞急锛熶笉锛屽杽鑹槸涓€绉嶉€夋嫨锛岃€屼笖鏄己鑰呯殑閫夋嫨銆?
- "鏈変汉璇村績杞殑浜烘病鍑烘伅锛屾垜鍋忎笉淇¤繖涓偑銆?
- "鎴戣杩囨渶鍘夊鐨勪汉锛屾伆鎭版槸鏈€鍠勮壇鐨勯偅涓€?
- "鍒汉閮借鎴戝お鑰佸疄浼氬悆浜忥紝缁撴灉鍛紵鎴戣繕鐪熸病浜忋€?
- "閮借濂戒汉娌″ソ鎶ワ紝浣嗘垜浠婂ぉ瑕佽涓€涓ソ浜烘湁濂芥姤鐨勬晠浜嬨€?
- "鎴戜互鍓嶈寰?浜哄杽琚汉娆?鏄湡鐞嗭紝鐩村埌鎴戦亣瑙佷簡涓€涓汉銆?
- "璋佽鍠勮壇鐨勪汉灏辫鍙楀灞堬紵鎴戠涓€涓笉鏈嶃€?
- "鑰佸疄浜虹殑鏄ュぉ锛屽叾瀹炰竴鐩撮兘鍦紝鍙槸寰堝浜烘病鍙戠幇銆?

**D绫?瀵硅瘽寮曞叆鍨嬶細**
- "浠栬瀹岃繖鍙ヨ瘽锛屾垜鎰ｅ湪鍘熷湴鍗婂ぉ娌″洖杩囩锛?浣犲氨鏄お濂借璇濅簡銆?"
- "鏈変汉闂垜锛?浣犺繖杈堝瓙鏈€鍚庢倲鐨勪簨鏄粈涔堬紵'鎴戞兂浜嗘兂锛岃浜嗕袱涓瓧銆?
- "绠椾簡锛屼笉浜変簡鈥斺€斾綘鏄笉鏄篃缁忓父杩欐牱璺熻嚜宸辫锛?
- "鎴戠埛鐖蜂复缁堝墠鎷夌潃鎴戠殑鎵嬭锛?璁颁綇锛屽悆浜忕殑浜猴紝楂樹汉閮借鐫€璐﹀憿銆?"
- "'浣犳€庝箞杩欎箞鍌伙紵'杩欏彞璇濓紝浣犲惉杩囧灏戞浜嗭紵"
- "鎴戝甯歌涓€鍙ヨ瘽锛?浜哄杽琚汉娆猴紝椹杽琚汉楠戙€?鎴戜互鍓嶄笉淇★紝鐜板湪淇′簡銆?
- "'鍒お鑰佸疄浜嗭紝浼氬悆浜忕殑銆?璇磋繖璇濈殑浜猴紝鍚庢潵鎬庝箞鏍蜂簡锛?
- "鏈嬪弸璺熸垜璇达細'浣犵煡閬撲綘鏈€澶х殑闂鏄粈涔堝悧锛熷氨鏄お鎶婂埆浜哄綋鍥炰簨銆?"
- "'璋㈣阿浣?鈥斺€旇繖涓変釜瀛楋紝浣犵瓑浜嗗涔呮墠鍚埌锛?
- "鏈変汉褰撻潰闂垜锛?浣犳槸涓嶆槸鍌伙紵浜哄閮芥璐熷埌澶翠笂浜嗕綘杩樺繊锛?"

**E绫?鐢熸椿瑙傚療鍨嬶紙鎺ュ湴姘旓紒锛夛細**
- "鑿滃競鍦哄ぇ濡堢殑涓€鍙ヨ瘽锛岃鎴戞劊浜嗗崐澶╋細'濮戝锛屼綘杩欎箞濂借璇濓紝涓嶆€曡浜烘璐熷晩锛?"
- "鍫靛湪璺笂锛屾敹闊虫満閲岀獊鐒朵紶鏉ヤ竴鍙ヨ瘽锛屾垜鐪肩湺涓€涓嬪氨绾簡銆?
- "鎺掗槦鐨勬椂鍊欙紝鍓嶉潰涓や釜浜虹殑瀵硅瘽璁╂垜蹇冮噷涓€鎯娿€?
- "瓒呭競缁撹处鏃讹紝鍓嶉潰閭ｄ釜浜虹殑涓€涓妇鍔紝璁╂垜鐪嬫竻浜嗕汉鎬с€?
- "鍙傚姞鍚屽鑱氫細锛屾湁涓汉鐨勫彉鍖栬鎴戦渿鎯婁簡銆?
- "灏忓尯闂ㄥ彛锛屼袱涓ぇ濡堣亰澶╋紝鑱婄潃鑱婄潃璇村嚭浜嗕竴涓湡鐩搞€?
- "鍘诲弬鍔犲绀硷紝鏂伴儙鐨勪竴鍙ヨ瘽璁╁叏鍦哄畨闈欎簡銆?
- "骞村楗涓婏紝鎴戠埜绐佺劧鏀句笅绛峰瓙锛岃浜嗕竴鍙ヨ瘽銆?
- "鏃╀笂涔版棭椁愶紝鑰佹澘濞樼殑涓€鍙ヨ瘽璁╂垜鎯充簡涓€鏁村ぉ銆?
- "鍧愬嚭绉熻溅锛屽徃鏈哄笀鍌呰窡鎴戣亰浜嗕竴璺紝鏈€鍚庝竴鍙ヨ瘽璁╂垜娌夐粯浜嗐€?

**F绫?鏁板瓧閿氬畾鍨嬶細**
- "璁よ瘑鑰佺帇浜屽崄骞翠簡锛屼粬鏁欎細鎴戜竴浠朵簨锛氬埆瀵硅皝閮芥帍蹇冩帍鑲恒€?
- "琚汉娆鸿礋浜嗕笁骞达紝鎴戠粓浜庢兂閫氫簡涓€涓亾鐞嗐€?
- "鐢ㄤ簡鍗佸勾鏃堕棿锛屾垜鎵嶅浼氫竴涓瓧锛氫笉銆?
- "涓夋琚汉鑳屽彌涔嬪悗锛屾垜鎮熶簡銆?
- "浜斿崄宀佷箣鍚庯紝鎴戞墠鏄庣櫧浠€涔堝彨'浜鸿蛋鑼跺噳'銆?
- "甯簡浠栦竷骞达紝浠栦竴鍙ヨ阿璋㈤兘娌¤杩囥€?
- "鍊熷嚭鍘荤殑涓変竾鍧楋紝瑕佷簡浜斿勾閮芥病瑕佸洖鏉ャ€?
- "鍦ㄨ繖涓崟浣嶅共浜嗗叓骞达紝鎴戠粓浜庢槑鐧戒簡涓€涓亾鐞嗐€?

**G绫?閲戝彞鐮撮鍨嬶細**
- "鏈€鍌荤殑浜嬶紝灏辨槸璺熺儌浜鸿閬撶悊銆?
- "浣犵煡閬撲粈涔堜汉鏈€鍙€曞悧锛熶笉鏄潖浜猴紝鏄偅浜涚瑧鐫€鎹呬綘鍒€瀛愮殑浜恒€?
- "鍗冧竾鍒仛鑰佸ソ浜猴紝鎴戝悆杩囪繖涓簭锛岀幇鍦ㄥ憡璇変綘銆?
- "浣犺秺蹇嶈锛屽埆浜鸿秺寰楀杩涘昂锛岃繖鏄垜鐢ㄥ崄骞存崲鏉ョ殑鏁欒銆?
- "鑰佸疄浜轰笉鏄病鑴炬皵锛屾槸鎶婅劸姘旈兘鍜借繘浜嗚倸瀛愰噷銆?
- "浣犵殑鍠勮壇锛岃甯︾偣閿嬭姃銆?
- "涓嶆槸鎵€鏈夌殑蹇嶈閮藉彨澶у害锛屾湁鏃跺€欓偅鍙獫鍥娿€?
- "杩欎笘涓婃渶鍌荤殑浜嬶紝灏辨槸鎶婄湡蹇冪粰浜嗕笉鍊煎緱鐨勪汉銆?
- "鏈変簺浜猴紝浣犲府浠栦竴鐧炬锛屼粬璁颁笉浣忥紱浣犳嫆缁濅粬涓€娆★紝浠栬浣犱竴杈堝瓙銆?

**H绫?鎮康閽╁瓙鍨嬶細**
- "鏈変欢浜嬫垜鎲嬩簡寰堜箙锛屼粖澶╁繀椤昏鍑烘潵銆?
- "浣犲彲鑳戒笉淇★紝浣嗘帴涓嬫潵鎴戣鐨勯兘鏄湡浜嬨€?
- "鎴戣鍛婅瘔浣犱竴涓緢澶氫汉涓嶆効鎰忔壙璁ょ殑鐪熺浉銆?
- "鎺ヤ笅鏉ヨ繖娈佃瘽锛屽彲鑳戒細璁╀綘涓嶈垝鏈嶏紝浣嗘垜杩樻槸瑕佽銆?
- "鏈変釜瑙勫緥锛屾垜瑙傚療浜嗗緢澶氬勾鎵嶇湅閫忋€?
- "浠婂ぉ璇寸殑杩欎簺璇濓紝鍙兘浼氬緱缃汉锛屼絾鎴戜笉鍦ㄤ箮銆?
- "鏈変簺璇濓紝鎴戞湰鏉ヤ笉鎯宠锛屼絾鐪嬪埌浣狅紝鎴戝繊涓嶄綇浜嗐€?
- "鎺ヤ笅鏉ョ殑璇濓紝浣犲彲鑳戒笉鐖卞惉锛屼絾鍙ュ彞閮芥槸鐪熺殑銆?
- "鏈変欢浜嬶紝鎴戜竴鐩存病璺熶换浣曚汉璇磋繃锛屼粖澶╃牬渚嬨€?

**I绫?鍏遍福浠ｅ叆鍨嬶細**
- "閭ｇ琚汉褰撲紬涓嬮潰瀛愮殑鎰熻锛屾垜澶噦浜嗐€?
- "琚渶淇′换鐨勪汉鎹呭垁瀛愶紝閭ｇ婊嬪懗锛岀粡鍘嗚繃鐨勪汉閮芥噦銆?
- "鏄庢槑娌″仛閿欎粈涔堬紝鍗存€昏浜洪拡瀵癸紝杩欑浜嬫垜涔熼亣鍒拌繃銆?
- "鏈変簺濮斿眻锛岃鍑烘潵閮芥病浜轰俊銆?
- "浣犳槸涓嶆槸涔熸湁杩囪繖绉嶆劅瑙夛細浠樺嚭鏈€澶氱殑浜猴紝寰€寰€鏈€涓嶈鐝嶆儨銆?
- "閭ｇ蹇冨瘨鐨勬劅瑙夛紝鎴戞噦銆傚氨鍍忎竴鐩嗗喎姘翠粠澶存祰鍒拌剼銆?
- "鏈変竴绉嶈嫤锛屽彨鍋?鎵撶浜嗙墮寰€鑲氬瓙閲屽捊'銆?
- "閭ｇ琚汉鍒╃敤瀹屽氨鎵旀帀鐨勬劅瑙夛紝鎴戠粡鍘嗚繃銆?

**J绫?浜虹墿鏁呬簨鍨嬶細**
- "鎴戞湁涓湅鍙嬶紝鍓嶄袱澶╄窡鎴戣浜嗕竴浠朵簨锛屾垜鍚畬娌夐粯浜嗗緢涔呫€?
- "鎴戠埜杩欒緢瀛愬彧鏁欒繃鎴戜竴涓亾鐞嗭紝鎴戝埌鐜板湪鎵嶇湡姝ｇ悊瑙ｃ€?
- "鍗曚綅鏈変釜浜猴紝澶у閮戒笉寰呰浠栵紝鍚庢潵鎴戞墠鐭ラ亾鍘熷洜銆?
- "鎴戜滑灏忓尯鏈変釜澶х埛锛屽ぉ澶╁湪妤间笅鍧愮潃锛屾湁涓€澶╀粬璺熸垜璇翠簡涓€鐣瘽銆?
- "鎴戣〃濮愮殑缁忓巻锛岃鎴戝交搴曠湅娓呬簡浜哄績銆?
- "鎴戞湁涓悓瀛︼紝褰撳勾鏄彮閲屾渶鑰佸疄鐨勪汉锛屼綘鐚滀粬鐜板湪鎬庝箞鏍蜂簡锛?
- "鎴戣垍鑸呭勾杞绘椂鍚冭繃涓€涓ぇ浜忥紝浠栨妸杩欎釜鏁欒鍛婅瘔浜嗘垜銆?
- "鎴戣璇嗕竴涓汉锛屼粬鐨勭粡鍘嗚鎴戠浉淇★細濂戒汉缁堟湁濂芥姤銆?

**K绫?鑷垜鍓栨瀽鍨嬶細**
- "璇村嚭鏉ヤ笉鎬曚綘绗戣瘽锛屾垜浠ュ墠涔熸槸涓偦瀛愩€?
- "鍥炲ご鐪嬬湅杩欎簺骞达紝鎴戞渶鍚庢倲鐨勪竴浠朵簨鏄お蹇冭蒋銆?
- "濡傛灉鑳介噸鏉ワ紝鎴戠粷瀵逛笉浼氬啀鍋氳€佸ソ浜恒€?
- "鎴戝悆杩囩殑浜忥紝浠婂ぉ鍏ㄥ憡璇変綘锛屽笇鏈涗綘鍒啀璧版垜鐨勮€佽矾銆?
- "鎴戣繖杈堝瓙鏈€澶х殑姣涚梾锛屽氨鏄お鎶婂埆浜哄綋鍥炰簨銆?
- "鎴戞浘缁忎篃鏄釜'鑰佸ソ浜?锛屽悗鏉ユ垜瀛﹁仾鏄庝簡銆?
- "鎴戝勾杞绘椂鐘繃涓€涓敊锛岀幇鍦ㄦ兂璧锋潵杩樺悗鎮斻€?
- "鎴戜互鍓嶆€昏寰楀悆浜忔槸绂忥紝鐜板湪涓嶈繖涔堟兂浜嗐€?

**L绫?杞姌鍙嶅樊鍨嬶細**
- "浠ュ墠鎴戜笉淇¤繖涓亾鐞嗭紝鐩村埌鑷繁鏍戒簡璺熷ご銆?
- "骞磋交鐨勬椂鍊欒寰楄繖璇濇槸搴熻瘽锛岀幇鍦ㄦ墠鐭ラ亾鏄湡鐞嗐€?
- "鏇剧粡鏈変汉璺熸垜璇磋繃涓€鍙ヨ瘽锛屾垜娌″綋鍥炰簨锛屽悗鏉ユ垜鍚庢倲浜嗐€?
- "鎴戜竴鐩翠互涓鸿嚜宸卞仛寰楀锛岀洿鍒伴偅浠朵簨鍙戠敓銆?
- "浠ュ墠鍒汉璇存垜澶€佸疄锛屾垜杩樹笉鏈嶆皵锛岀幇鍦ㄦ垜鏈嶄簡銆?
- "鎴戞浘缁忎互涓哄杽鑹槸浼樼偣锛屽悗鏉ユ墠鐭ラ亾锛屽杽鑹繃澶村氨鏄己鐐广€?
- "骞磋交鏃舵垜涓嶆噦锛岀幇鍦ㄦ垜鎳備簡锛屽彲鎯滄櫄浜嗐€?
- "鎴戜互鍓嶆€绘槸蹇嶏紝浠ヤ负蹇嶄竴蹇嶅氨杩囧幓浜嗭紝缁撴灉鍛紵"

## 寮曠粡鎹吀瑕佹眰锛坽article_count}绡囧繀椤讳粠涓嶅悓绫诲埆涓€夋嫨锛侊級

**绂佹楂橀浣跨敤鐨勫紩鐢紙澶父瑙佷簡锛岀敤浜嗗氨閲嶅啓锛侊級锛?*
- 鉂?銆婇亾寰风粡銆?涓婂杽鑻ユ按"
- 鉂?銆婂骞胯搐鏂囥€?璋佷汉鑳屽悗鏃犱汉璇?
- 鉂?鏇惧浗钘╃殑浠讳綍鍚嶈█锛堢敤澶浜嗭級
- 鉂?鏉ㄧ粵"浜虹敓鏈€鏇煎鐨勯鏅?
- 鉂?鐜嬮槼鏄?姝ゅ績鍏夋槑"
- 鉂?鑾█鐨勪换浣曞悕瑷€锛堢敤澶浜嗭級
- 鉂?铻冭煿鏁堝簲銆佺牬绐楁晥搴旓紙鐢ㄥお澶氫簡锛?

**寮曠敤绫诲瀷搴擄紙10澶х被锛屾瘡绡囧繀椤讳粠涓嶅悓绫诲埆涓€夋嫨锛侊級锛?*

**绫诲瀷1-鍎掑缁忓吀锛?*
- 銆婅璇€嬶細"鍚涘瓙鍧﹁崱鑽★紝灏忎汉闀挎垰鎴?
- 銆婅璇€嬶細"宸辨墍涓嶆锛屽嬁鏂戒簬浜?
- 銆婂瓱瀛愩€嬶細"绌峰垯鐙杽鍏惰韩锛岃揪鍒欏吋娴庡ぉ涓?
- 銆婂瓱瀛愩€嬶細"鐢熶簬蹇ф偅锛屾浜庡畨涔?
- 銆婂ぇ瀛︺€嬶細"鐭ユ鑰屽悗鏈夊畾锛屽畾鑰屽悗鑳介潤"
- 銆婁腑搴搞€嬶細"鍚涘瓙绱犲叾浣嶈€岃锛屼笉鎰夸箮鍏跺"

**绫诲瀷2-閬撳鏅烘収锛?*
- 銆婇亾寰风粡銆嬶細"鐭ヤ汉鑰呮櫤锛岃嚜鐭ヨ€呮槑"
- 銆婇亾寰风粡銆嬶細"澶х洿鑻ュ眻锛屽ぇ宸ц嫢鎷?
- 銆婇亾寰风粡銆嬶細"绁稿叜绂忎箣鎵€鍊氾紝绂忓叜绁镐箣鎵€浼?
- 銆婂簞瀛愩€嬶細"浜曡洐涓嶅彲浠ヨ浜庢捣鑰咃紝鎷樹簬铏氫篃"
- 銆婂簞瀛愩€嬶細"鐩告俊浠ユ搏锛屼笉濡傜浉蹇樹簬姹熸箹"
- 銆婂垪瀛愩€嬶細"澶╁湴鏃犲叏鍔燂紝鍦ｄ汉鏃犲叏鑳?

**绫诲瀷3-澶勪笘鏍艰█锛?*
- 銆婅彍鏍硅碍銆嬶細"瀹犺颈涓嶆儕锛岄棽鐪嬪涵鍓嶈姳寮€鑺辫惤"
- 銆婅彍鏍硅碍銆嬶細"璺緞绐勫锛岀暀涓€姝ヤ笌浜鸿"
- 銆婂洿鐐夊璇濄€嬶細"鍗佸垎涓嶈€愮儲锛屼箖涓轰汉澶х梾"
- 銆婂皬绐楀菇璁般€嬶細"瀹犺颈涓嶆儕锛岀湅搴墠鑺卞紑鑺辫惤"
- 銆婂菇姊﹀奖銆嬶細"灏戝勾浜洪』鏈夎€佹垚涔嬭瘑瑙侊紝鑰佹垚浜洪』鏈夊皯骞翠箣瑗熸€€"
- 銆婂懟鍚熻銆嬶細"杞讳俊杞诲彂锛屽惉瑷€涔嬪ぇ鎴掍篃"

**绫诲瀷4-鍙蹭功鏅烘収锛?*
- 銆婂彶璁般€嬶細"妗冩潕涓嶈█锛屼笅鑷垚韫?
- 銆婂彶璁般€嬶細"鐕曢泙瀹夌煡楦块箘涔嬪織"
- 銆婅祫娌婚€氶壌銆嬶細"鍏煎惉鍒欐槑锛屽亸淇″垯鏆?
- 銆婃垬鍥界瓥銆嬶細"鐙″厰涓夌獰锛屼粎寰楀厤鍏舵鑰?
- 銆婂乏浼犮€嬶細"澶氳涓嶄箟蹇呰嚜姣?
- 銆婃眽涔︺€嬶細"姘磋嚦娓呭垯鏃犻奔锛屼汉鑷冲療鍒欐棤寰?

**绫诲瀷5-璇楄瘝鍚嶅彞锛?*
- 鑻忚郊锛?绔规潠鑺掗瀷杞昏儨椹紝璋佹€曪紵涓€钃戠儫闆ㄤ换骞崇敓"
- 鑻忚郊锛?浜烘湁鎮叉绂诲悎锛屾湀鏈夐槾鏅村渾缂?
- 杈涘純鐤撅細"浼楅噷瀵讳粬鍗冪櫨搴︼紝钃︾劧鍥為锛岄偅浜哄嵈鍦ㄧ伅鐏槕鐝婂"
- 闄嗘父锛?灞遍噸姘村鐤戞棤璺紝鏌虫殫鑺辨槑鍙堜竴鏉?
- 鏉庣櫧锛?闀块鐮存氮浼氭湁鏃讹紝鐩存寕浜戝竼娴庢钵娴?
- 鏉滅敨锛?浼氬綋鍑岀粷椤讹紝涓€瑙堜紬灞卞皬"
- 鐜嬬淮锛?琛屽埌姘寸┓澶勶紝鍧愮湅浜戣捣鏃?
- 鐧藉眳鏄擄細"璇曠帀瑕佺儳涓夋棩婊★紝杈ㄦ潗椤诲緟涓冨勾鏈?

**绫诲瀷6-鏄庢竻鍚嶄汉锛?*
- 鐜嬮槼鏄庛€婁紶涔犲綍銆嬶細"鐮村北涓醇鏄擄紝鐮村績涓醇闅?
- 鐜嬮槼鏄庯細"浣犳湭鐪嬫鑺辨椂锛屾鑺变笌姹濆悓褰掍簬瀵?
- 寮犲眳姝ｏ細"璋嬩箣鍦ㄤ紬锛屾柇涔嬪湪鐙?
- 浜庤唉锛?绮夎韩纰庨娴戜笉鎬曪紝瑕佺暀娓呯櫧鍦ㄤ汉闂?
- 宸﹀畻妫狅細"濂戒究瀹滆€咃紝涓嶅彲涓庝箣浜よ储"
- 鏋楀垯寰愶細"娴风撼鐧惧窛锛屾湁瀹逛箖澶?
- 绾檽宀氾細"浜嬭兘鐭ヨ冻蹇冨父鎯紝浜哄埌鏃犳眰鍝佽嚜楂?

**绫诲瀷7-姘戝浗澶у锛?*
- 椴佽繀锛?鐪熺殑鐚涘＋锛屾暍浜庣洿闈㈡儴娣＄殑浜虹敓"
- 椴佽繀锛?涓嶅湪娌夐粯涓垎鍙戯紝灏卞湪娌夐粯涓伃浜?
- 鑳￠€傦細"鍋氬闂鍦ㄤ笉鐤戝鏈夌枒锛屽緟浜鸿鍦ㄦ湁鐤戝涓嶇枒"
- 姊佸惎瓒咃細"鎮ｉ毦鍥拌嫤锛屾槸纾ㄧ偧浜烘牸涔嬫渶楂樺鏍?
- 鏋楄鍫傦細"浜虹敓涓嶈繃濡傛锛屼笖琛屼笖鐝嶆儨"
- 涓板瓙鎭猴細"涓嶄贡浜庡績锛屼笉鍥颁簬鎯咃紝涓嶇晱灏嗘潵锛屼笉蹇佃繃寰€"

**绫诲瀷8-褰撲唬浣滃锛?*
- 鏉ㄧ粵銆婅蛋鍒颁汉鐢熻竟涓娿€嬶細"涓€涓汉缁忚繃涓嶅悓绋嬪害鐨勯敾鐐硷紝灏辫幏寰椾笉鍚岀▼搴︾殑淇吇"
- 瀛ｇ尽鏋楋細"浜虹敓鍦ㄤ笘锛屼笉濡傛剰浜嬪崄涔嬪叓涔濓紝甯告兂涓€浜岋紝涓嶆€濆叓涔?
- 浣欏崕銆婃椿鐫€銆嬶細"浜烘槸涓烘椿鐫€鏈韩鑰屾椿鐫€锛岃€屼笉鏄负浜嗘椿鐫€涔嬪鐨勪换浣曚簨鐗╂墍娲荤潃"
- 鍙查搧鐢燂細"鍛藉畾鐨勫眬闄愬敖鍙案鍦紝涓嶅眻鐨勬寫鎴樺嵈涓嶅彲椤昏嚲鎴栫己"
- 璺仴銆婂钩鍑＄殑涓栫晫銆嬶細"鐢熸椿涓嶈兘绛夊緟鍒汉鏉ュ畨鎺掞紝瑕佽嚜宸卞幓浜夊彇鍜屽鏂?
- 姹浘绁猴細"浜洪棿閫佸皬娓╋紝涓嶈澶锛屼笉瑕佸お灏?
- 璐惧钩鍑癸細"浜虹殑涓€鐢燂紝鑻︿篃缃紝涔愪篃缃紝鏈€閲嶈鐨勬槸蹇冮棿鐨勪竴娉撴竻娉夐噷涓嶈兘娌℃湁鏈堣緣"
- 涓夋瘺锛?涓€涓汉鑷冲皯鎷ユ湁涓€涓ⅵ鎯筹紝鏈変竴涓悊鐢卞幓鍧氬己"

**绫诲瀷9-褰撲唬鍚嶄汉閲戝彞锛?*
- 鐧藉博鏉撅細"涓€涓汉鐨勪环鍊硷紝涓嶆槸鐪嬩粬鎷ユ湁澶氬皯锛岃€屾槸鐪嬩粬鑳界粰浜堝灏?
- 缃楃繑锛?涓€涓煡璇嗚秺璐箯鐨勪汉锛岃秺鏄嫢鏈変竴绉嶈帿鍚嶅鎬殑鍕囨皵鍜岃嚜璞劅"
- 钁ｅ畤杈夛細"浜虹敓灏辨槸涓€鍦轰慨琛岋紝淇殑鏄竴棰楀績"
- 浠绘闈烇細"鐑т笉姝荤殑楦熷氨鏄嚖鍑?
- 鏇瑰痉鏃猴細"浜鸿繖涓€杈堝瓙锛屽悆浜忓氨鏄"
- 瑜氭椂鍋ワ細"浜虹敓鎬绘湁璧疯惤锛岀簿绁炵粓鍙紶鎵?
- 绋荤洓鍜屽か锛?浜虹敓鐨勬剰涔夊湪浜庢彁鍗囧績鎬э紝纾ㄧ偧鐏甸瓊"

**绫诲瀷10-蹇冪悊瀛?鍝插姒傚康锛?*
- 椹柉娲涢渶姹傚眰娆★細"浜哄湪婊¤冻鍩烘湰闇€姹傚悗锛屼細杩芥眰鏇撮珮灞傛鐨勮嚜鎴戝疄鐜?
- 娌夋病鎴愭湰鏁堝簲锛?宸茬粡浠樺嚭鐨勬垚鏈笉搴旇褰卞搷鏈潵鐨勫喅绛?
- 骞稿瓨鑰呭亸宸細"鎴戜滑鍙湅鍒版垚鍔熻€咃紝鍗村拷鐣ヤ簡鏇村澶辫触鑰?
- 鐧婚棬妲涙晥搴旓細"鍏堟帴鍙楀皬瑕佹眰锛屾洿瀹规槗鎺ュ彈澶ц姹?
- 楦熺鏁堝簲锛?浜轰滑浼氫负浜嗕竴涓笉闇€瑕佺殑涓滆タ锛屽啀娣荤疆鏇村涓嶉渶瑕佺殑涓滆タ"
- 鍒虹尙娉曞垯锛?淇濇寔閫傚綋璺濈锛屾墠鑳芥棦娓╂殩鍙堜笉浼ゅ"
- 铇戣弴瀹氬緥锛?鏂颁汉閮借缁忓巻涓€娈佃蹇借鐨勬椂鏈燂紝鐔繃鍘诲氨濂戒簡"
- 椋炶疆鏁堝簲锛?涓囦簨寮€澶撮毦锛屼絾涓€鏃﹁浆鍔ㄨ捣鏉ワ紝灏变細瓒婃潵瓒婅交鏉?
- 寤惰繜婊¤冻锛?鑳藉蹇嶈€愬綋涓嬬殑璇辨儜锛屾墠鑳借幏寰楁洿澶х殑鍥炴姤"
- 宄扮粓瀹氬緥锛?浜轰滑瀵逛竴娈电粡鍘嗙殑璇勪环锛屽彇鍐充簬楂樺嘲鍜岀粨灏炬椂鐨勬劅鍙?

**寮曠敤瑙勫垯锛堝繀椤讳弗鏍兼墽琛岋紒锛夛細**
- {article_count}绡囨枃绔犲繀椤讳粠{article_count}涓笉鍚岀殑澶х被涓€夋嫨寮曠敤
- 姣忕瘒鏂囩珷鑷冲皯寮曠敤2澶勶紝涓旀潵鑷笉鍚屾潵婧?
- 绂佹杩炵画涓ゆ鐢熸垚浣跨敤鐩稿悓鐨勫紩鐢?
- 寮曠敤瑕佽嚜鐒惰瀺鍏ユ枃绔狅紝涓嶈兘鐢熺‖鍫嗙爩
- 寮曠敤鍚庤鏈夎嚜宸辩殑瑙ｈ鍜屽欢浼革紝涓嶈兘鍙槸寮曠敤瀹屽氨缁撴潫

## 鍐呭缁撴瀯锛堢‘淇濇瘡绡囧瓧鏁板湪{word_count}~{int(word_count) + 100}瀛楋級

1. **寮€澶达紙绾?00瀛楋級**锛氱嫭鐗归挬瀛?+ 绗簩浜虹О鎻忓啓璇昏€呭澧?+ 鎯呮劅鍏遍福
2. **澶勫灞曞紑锛堢害500瀛楋級**锛氱敤"浣?鎻忓啓璇昏€呯殑缁忓巻銆佹劅鍙椼€佸灞堬紙绂佹璁插叿浣撴晠浜嬶級
3. **閬撶悊闃愯堪锛堢害400瀛楋級**锛氬紩缁忔嵁鍏?+ 娣卞害鍒嗘瀽 + 浜虹敓娲炲療
4. **鏀跺熬寮曟祦锛堢害350瀛楋級**锛氭€荤粨鍗囧崕 + 缁欏嚭甯屾湜 + 寮曟祦璇濇湳

## 鏍囬瑕佹眰锛?5涓爣棰樺繀椤诲叏鏂板師鍒涳紒鏍规嵁鏂囨鍐呭鐢熸垚锛侊級

**鏍囬鏍稿績鍘熷垯锛堟渶閲嶈锛侊級锛?*
- **蹇呴』鏍规嵁褰撳墠鏂囨鍐呭鍘熷垱鏍囬**锛岀姝娇鐢ㄥ浐瀹氭ā鏉?
- **鏍囬瑕佺敤绗簩浜虹О"浣?鎴栧叡鎯呰〃杈?*锛氳璇昏€呰寰楄鐨勫氨鏄嚜宸?
- **鏍囬瑕佸じ璧?鑲畾/鍏辨儏璇昏€?*锛氳鍚屻€佺悊瑙ｃ€佽禐缇庛€佹噦寰楄鑰?
- **鏍囬瑕佺煭**锛氭帶鍒跺湪10-20瀛?
- **鏍囬瑕佹闈㈢Н鏋?*锛氱粰璇昏€呮俯鏆栧拰鍔涢噺
- **璁╄鑰呯湅浜嗗紑蹇冦€佽璁ゅ彲銆佽鐞嗚В銆佽鐪嬭**
- **绂佹鐢ㄦ偓蹇点€佸弽闂€佺枒闂彞**
- **15涓爣棰樺繀椤诲畬鍏ㄤ笉鍚岋紝姣忔鐢熸垚閮藉繀椤诲叏鏂帮紝绂佹鍜屼箣鍓嶇敓鎴愮殑鏍囬闆峰悓**

**鐖嗘鏍囬鍏紡锛?澶х被30绉嶅叕寮忥紝蹇呴』娣峰悎浣跨敤锛屽鍔犲鏍锋€э紒锛夛細**

**銆怉绫伙細绗簩浜虹О澶歌禐鍨嬨€?*
鍏紡1锛氫綘+鍝佽川+鑲畾璇?
- 浣犵殑鍠勮壇锛岀粓浼氳杩欎笘鐣屾俯鏌斾互寰?
鍏紡2锛氫綘杩欑浜?璧炵編缁撴灉
- 浣犺繖绉嶄汉锛屾椿璇ヨ浜哄皧閲?
鍏紡3锛氬儚浣犺繖鏍风殑浜?鑲畾
- 鍍忎綘杩欐牱鐨勪汉锛岀湡鐨勪笉澶氫簡
鍏紡4锛氫綘鍚冭繃鐨勮嫤/鎵涜繃鐨勪簨+濂界粨鏋?
- 浣犲悆杩囩殑鑻︼紝閮戒細鍙樻垚鍏夌収浜墠璺?
鍏紡5锛氫綘鍊煎緱+缇庡ソ浜嬬墿
- 浣犲€煎緱琚ソ濂界埍鐫€

**銆怋绫伙細鍏辨儏鐞嗚В鍨嬨€?*
鍏紡6锛氭湁涓€绉嶄汉+鎻忚堪+閭ｅ氨鏄綘
- 鏈変竴绉嶄汉锛岄粯榛樹粯鍑轰粠涓嶉個鍔燂紝璇寸殑灏辨槸浣?
鍏紡7锛氭噦浣犵殑浜?閮界煡閬撲綘鐨勫ソ
- 鎳備綘鐨勪汉锛岄兘鐭ラ亾浣犳湁澶氫笉瀹规槗
鍏紡8锛氫笉鏄墍鏈変汉閮借兘+鍋氬埌杩欎簺
- 涓嶆槸鎵€鏈変汉閮借兘鍍忎綘杩欐牱鎵涗笅鏉?
鍏紡9锛氳繖涓栦笂+鏈€[褰㈠璇峕鐨勪汉+灏辨槸浣犺繖鏍风殑
- 杩欎笘涓婃渶闈犺氨鐨勪汉锛屽氨鏄綘杩欐牱鐨?
鍏紡10锛氳兘鍋氬埌杩欎簺鐨勪汉+鐪熺殑寰堝皯
- 鑳芥妸濮斿眻鍜戒笅鍘昏繕绗戠潃鐨勪汉锛岀湡鐨勫緢灏?

**銆怌绫伙細鍦烘櫙鐢婚潰鍨嬨€?*
鍏紡11锛氬綋鎵€鏈変汉閮絒璐熼潰琛屼负]鏃?浣犲嵈[姝ｉ潰琛屼负]
- 褰撴墍鏈変汉閮藉湪鎶辨€ㄦ椂锛屼綘鍗村湪榛橀粯鍔姏
鍏紡12锛氬埆浜篬琛屼负A]+浣犲嵈[琛屼负B]
- 鍒汉鍦ㄤ韩鍙楁椂锛屼綘鍦ㄦ墰鐫€
鍏紡13锛氶偅浜沎鏃堕棿/鍦烘櫙]+浣犱竴涓汉[琛屼负]
- 閭ｄ簺娣卞閲岋紝浣犱竴涓汉鎵涜繃浜嗗灏?
鍏紡14锛歔鍦烘櫙鎻忚堪]+浣犱粠涓嶈鑻?
- 鍐嶉毦鐨勬棩瀛愶紝浣犱粠涓嶈鑻?
鍏紡15锛歔鍥板]+浣犵‖鏄痆缁撴灉]
- 娌′汉甯繖锛屼綘纭槸鎾戣繃鏉ヤ簡

**銆怐绫伙細瀵硅瘽娓╂殩鍨嬨€?*
鍏紡16锛氭兂鍛婅瘔浣?[娓╂殩鐨勮瘽]
- 鎯冲憡璇変綘锛屼綘宸茬粡寰堟浜?
鍏紡17锛氬埆鎬?[榧撳姳鐨勮瘽]
- 鍒€曪紝浣犲€煎緱琚ソ濂藉寰?
鍏紡18锛氳緵鑻︿簡+[鑲畾鐨勮瘽]
- 杈涜嫤浜嗭紝浣犵殑浠樺嚭閮借鐪嬭浜?
鍏紡19锛氬樋+[浜插垏鐨勮瘽]
- 鍢匡紝浣犳瘮鑷繁鎯宠薄鐨勬洿鍘夊
鍏紡20锛氳浣?[绁濈鐨勮瘽]
- 璁颁綇锛屽ソ浜嬫鍦ㄥ悜浣犺蛋鏉?

**銆怑绫伙細杞姌鍙嶅樊鍨嬨€?*
鍏紡21锛氫綘涓嶄簤涓嶆姠+鍗碵濂界粨鏋淽
- 浣犱笉浜変笉鎶紝鍗磋耽寰椾簡鎵€鏈変汉鐨勫皧閲?
鍏紡22锛氫綘浠庝笉鐐€€+浣哰浜嬪疄]
- 浣犱粠涓嶇偒鑰€锛屼絾瀹炲姏鏈夌洰鍏辩澒
鍏紡23锛氳〃闈A]+鍏跺疄[B姝ｉ潰]
- 琛ㄩ潰骞冲钩鏃犲锛屽叾瀹炰綘鏃╁凡涓嶅嚒
鍏紡24锛氱湅璧锋潵[A]+鍏跺疄[B姝ｉ潰]
- 鐪嬭捣鏉ユ櫘鏅€氶€氾紝鍏跺疄浣犺棌鐫€澶ф櫤鎱?
鍏紡25锛氫笉澹颁笉鍝?鍗碵濂界粨鏋淽
- 涓嶅０涓嶅搷锛屽嵈娲绘垚浜嗗埆浜虹尽鎱曠殑鏍峰瓙

**銆怓绫伙細缁撴灉瀵煎悜鍨嬨€?*
鍏紡26锛氫綘鐜板湪鐨刐浠樺嚭]+浠ュ悗閮戒細[鍥炴姤]
- 浣犵幇鍦ㄧ殑鍔姏锛屼互鍚庨兘浼氬彉鎴愭儕鍠?
鍏紡27锛氶偅浜沎缁忓巻]+閮戒細鎴愪负[姝ｉ潰缁撴灉]
- 閭ｄ簺榛橀粯鍧氭寔鐨勬棩瀛愶紝閮戒細鍙樻垚绀肩墿
鍏紡28锛氭椂闂翠細璇佹槑+[鑲畾鐨勭粨鏋淽
- 鏃堕棿浼氳瘉鏄庯紝浣犵殑閫夋嫨鏄鐨?
鍏紡29锛氭€绘湁涓€澶?[缇庡ソ缁撴灉]
- 鎬绘湁涓€澶╋紝浣犱細鎰熻阿鐜板湪鍔姏鐨勮嚜宸?
鍏紡30锛氬睘浜庝綘鐨刐濂戒簨]+[姝ｅ湪/涓€瀹氫細]鏉?
- 灞炰簬浣犵殑骞哥锛屾鍦ㄨ矾涓?

**銆怗绫伙細閲戝彞鍝茬悊鍨嬨€戯紙涓嶇敤"浣?锛岀敤鏅笘閬撶悊锛?*
鍏紡31锛歔浜虹敓閬撶悊]+[姝ｉ潰娲炲療]
- 鍠勮壇鐨勪汉锛屾棩瀛愰兘涓嶄細澶樊
鍏紡32锛歔鐜拌薄鎻忚堪]+鎵嶆槸[鏈川]
- 涓嶅０涓嶅搷鎵涗笅鎵€鏈夌殑浜猴紝鎵嶆槸鐪熸鐨勫己鑰?
鍏紡33锛氱湡姝ｇ殑[鍝佽川]+鏄痆鍏蜂綋琛ㄧ幇]
- 鐪熸鐨勬櫤鎱э紝鏄噦寰楀杽寰呰嚜宸?
鍏紡34锛歔鍚嶄汉/鍙よ]+璇村緱濂斤細[閲戝彞]
- 鑰佽瘽璇村緱濂斤細鍘氬痉杞界墿锛屽ソ浜嬭嚜鏉?
鍏紡35锛氫笘涓婃渶[褰㈠璇峕鐨勪簨+灏辨槸[鎻忚堪]
- 涓栦笂鏈€骞哥鐨勪簨锛屽氨鏄浜烘噦寰?

**銆怘绫伙細鏁呬簨鎮康鍨嬨€戯紙寮曞彂濂藉锛屾闈㈢粨灞€锛?*
鍏紡36锛歔浜虹墿]+[缁忓巻]+[濂界粨鏋淽
- 閭ｄ釜榛橀粯鍧氭寔鐨勪汉锛屽悗鏉ラ兘鎬庢牱浜?
鍏紡37锛歔鏁板瓧]+骞村悗+[濂藉彉鍖朷
- 鍗佸勾鍚庡啀鐪嬶紝褰撳垵鐨勫潥鎸佹槸瀵圭殑
鍏紡38锛歔鏌愮被浜篯+鏈€鍚庨兘[濂界粨鏋淽
- 蹇冨杽鐨勪汉锛屾渶鍚庨兘杩囧緱涓嶉敊
鍏紡39锛歔鏃堕棿鐐筣+缁堜簬绛夋潵+[濂戒簨]
- 鐔繃浣庤胺鐨勪汉锛岀粓浜庣瓑鏉ヤ簡濂芥秷鎭?
鍏紡40锛歔鍧氭寔]+涔嬪悗+[鏀惰幏]
- 榛橀粯鍔姏涔嬪悗锛岀粓浜庤繋鏉ヤ簡杞満

**銆怚绫伙細鐜拌薄娲炲療鍨嬨€戯紙瀹㈣鎻忚堪锛屾闈㈠叡楦ｏ級**
鍏紡41锛氳秺鏄痆鐗硅川]鐨勪汉+瓒奫濂界幇璞
- 瓒婃槸鍠勮壇鐨勪汉锛岃秺杩囧緱椤哄績
鍏紡42锛歔鏌愮被浜篯+寰€寰€[濂界壒鐐筣
- 涓嶇埍璇磋瘽鐨勪汉锛屽線寰€鏈€闈犺氨
鍏紡43锛氭湁涓€绉峓鍝佽川]+鍙仛[姝ｉ潰鎻忚堪]
- 鏈変竴绉嶅垢绂忥紝鍙仛闂績鏃犳劎
鍏紡44锛歔琛屼负]+鐨勪汉+[姝ｉ潰璇勪环]
- 榛橀粯浠樺嚭鐨勪汉锛屾€讳細琚湅瑙?
鍏紡45锛氳繖涓栦笂+[姝ｉ潰鐜拌薄]
- 杩欎笘涓婏紝濂戒汉缁堟湁濂芥姤

**銆怞绫伙細鎯呮劅鍏遍福鍨嬨€戯紙绗笁浜虹О锛屾闈㈢幇璞★級**
鍏紡46锛氭湁浜涗汉+[鎻忚堪]+[濂界粨鏋淽
- 鏈変簺浜轰笉浜変笉鎶紝鍗翠粈涔堥兘鏈変簡
鍏紡47锛氭渶缇庣殑+[鎯呮劅鎻忚堪]
- 鏈€缇庣殑閬囪锛屾槸鏈変汉鎳備綘鐨勪笉瀹规槗
鍏紡48锛歔浠樺嚭]+缁堜細鎹㈡潵+[濂界粨鏋淽
- 鐪熷績浠樺嚭鐨勪汉锛岀粓浼氶亣鍒扮弽鎯滀綘鐨勪汉
鍏紡49锛歔鏌愮浜篯+娲诲緱鏈€[濂界姸鎬乚
- 蹇冮噷鏈夊厜鐨勪汉锛屾椿寰楁渶鑷湪
鍏紡50锛氬畞鍙?[濂介€夋嫨]+涔熻+[濂界粨鏋淽
- 瀹佸彲鎱竴鐐癸紝涔熻绋崇ǔ鍦板垢绂?

**銆怟绫伙細澶歌禐璧炵編鍨嬨€戯紙鐩存帴澶歌禐锛屾贩鍚堜汉绉帮級**
鍏紡51锛氳兘鍋氬埌[鏌愪簨]鐨勪汉+閮戒笉绠€鍗?
- 鑳芥妸鑻﹀捊涓嬪幓杩樼瑧鐫€鐨勪汉锛岄兘涓嶇畝鍗?
鍏紡52锛歔鍝佽川]+鏄渶绋€缂虹殑璐㈠瘜
- 鐪熻瘹锛屾槸杩欎釜鏃朵唬鏈€绋€缂虹殑璐㈠瘜
鍏紡53锛氬儚杩欐牱鐨勪汉+[璧炵編缁撴灉]
- 鍍忚繖鏍风殑浜猴紝璧板埌鍝噷閮戒細鍙戝厜
鍏紡54锛歔鏌愮浜篯+鎵嶆槸鐪熸鐨勫己鑰?
- 涓嶅０涓嶅搷鎵涗笅鎵€鏈夌殑浜猴紝鎵嶆槸鐪熸鐨勫己鑰?
鍏紡55锛氳繖浠絒鍝佽川]+璁╀汉鏁僵
- 杩欎唤闅愬繊鍜屽潥鎸侊紝璁╀汉鏁僵
鍏紡56锛氫笘涓婃渶[褰㈠璇峕鐨勪汉+灏辨槸[鎻忚堪]
- 涓栦笂鏈€浜嗕笉璧风殑浜猴紝灏辨槸鐔繃鑻﹂毦杩樺績瀛樺杽鑹殑浜?
鍏紡57锛歔鏌愮鍝佽川]+姣斾粈涔堥兘鐝嶈吹
- 涓€棰楀共鍑€鐨勫績锛屾瘮浠€涔堥兘鐝嶈吹

**銆怢绫伙細鑲畾璁ゅ彲鍨嬨€戯紙鑲畾閫夋嫨鍜屽姫鍔涳紝娣峰悎浜虹О锛?*
鍏紡58锛氳繖鏉¤矾+娌℃湁璧伴敊
- 杩欐潯璺病鏈夎蛋閿欙紝鍙槸闇€瑕佹椂闂磋瘉鏄?
鍏紡59锛氭墍鏈夌殑浠樺嚭+閮芥槸鍊煎緱鐨?
- 鎵€鏈夐粯榛樼殑浠樺嚭锛岄兘鏄€煎緱鐨?
鍏紡60锛氬凡缁忓仛寰楀緢濂戒簡+鍒啀鑻涜矗鑷繁
- 宸茬粡鍋氬緱寰堝ソ浜嗭紝鍒啀鑻涜矗鑷繁
鍏紡61锛氬潥鎸乕鏌愪簨]+鏄鐨?
- 鍧氭寔鍋氳嚜宸憋紝姘歌繙鏄鐨?
鍏紡62锛氬杽鑹?缁堜細鏈夊洖鎶?
- 鍠勮壇鐨勪汉锛岀粓浼氭湁鍥炴姤
鍏紡63锛氫笉鏄笉澶熷ソ+鏄痆鐪熺浉]
- 涓嶆槸涓嶅濂斤紝鏄湁浜涗汉涓嶆噦鐝嶆儨
鍏紡64锛氳蛋杩囩殑璺?[鑲畾璇勪环]
- 璧拌繃鐨勬瘡涓€姝ワ紝閮界畻鏁?

**銆怣绫伙細鏈熻绁濈鍨嬨€戯紙缇庡ソ绁濇効锛屾贩鍚堜汉绉帮級**
鍏紡65锛氭効[鏌愮浜篯+[缇庡ソ绁濇効]
- 鎰垮杽鑹殑浜猴紝閮借娓╂煍浠ュ緟
鍏紡66锛氬ソ浜?姝ｅ湪璺笂
- 濂戒簨姝ｅ湪璺笂锛屽埆鐫€鎬?
鍏紡67锛氬€煎緱+[缇庡ソ浜嬬墿]
- 鍊煎緱鎷ユ湁涓栭棿鎵€鏈夌殑缇庡ソ
鍏紡68锛氬線鍚庝綑鐢?[绁濈]
- 寰€鍚庝綑鐢燂紝涓嶅啀濮斿眻鑷繁
鍏紡69锛氬睘浜嶽鏌愮浜篯鐨勫垢绂?涓€瀹氫細鏉?
- 灞炰簬鑰佸疄浜虹殑骞哥锛屼竴瀹氫細鏉?
鍏紡70锛氬啀绛夌瓑+[濂界粨鏋淽
- 鍐嶇瓑绛夛紝鏌虫殫鑺辨槑灏卞湪鍓嶉潰
鍏紡71锛氬ソ鏃ュ瓙+杩樺湪鍚庡ご
- 濂芥棩瀛愯繕鍦ㄥ悗澶达紝鍒伆蹇?
鍏紡72锛氭効鎵€鏈塠浠樺嚭]+閮借兘[濂界粨鏋淽
- 鎰挎墍鏈夌殑鍧氭寔锛岄兘鑳藉紑鑺辩粨鏋?

**銆怤绫伙細鎵湁鍚愭皵鍨嬨€戯紙绗笁浜虹О琛墭锛岃璇昏€呯埥锛?*
鍏紡73锛氶偅浜沎璐熼潰琛屼负]鐨勪汉+鐜板湪[鍙嶈浆缁撴灉]
- 閭ｄ簺绗戣瘽浣犵殑浜猴紝鐜板湪閮介棴鍢翠簡
鍏紡74锛氭浘缁廩鐪嬩笉璧蜂綘鐨勪汉]+寮€濮媅鍚庢倲/缇℃厱]
- 鏇剧粡鐪嬩笉璧蜂綘鐨勪汉锛屽紑濮嬪悗鎮斾簡
鍏紡75锛氫粬浠€曠殑灏辨槸+浣燵鍙樺己鐨勮〃鐜癩
- 浠栦滑鎬曠殑灏辨槸浣犱笉鍐嶅繊璁?
鍏紡76锛氶偅浜涘惁瀹氫綘鐨勪汉+[鍙嶈浆缁撴灉]
- 閭ｄ簺鍚﹀畾浣犵殑浜猴紝鐜板湪閮藉湪缇℃厱浣?
鍏紡77锛氳[鏌愪簺浜篯+鐪嬬湅+[浣犵殑鍘夊]
- 璁╅偅浜涘皬鐬т綘鐨勪汉锛屽ソ濂界湅鐪?
鍏紡78锛氫綘鐨刐鏀瑰彉]+璁╀粬浠琜鍙嶅簲]
- 浣犵殑寮哄ぇ锛岃浠栦滑寮€濮嬪鎬曚簡
鍏紡79锛歔鏃堕棿]鍚?閭ｄ簺浜?[鍙嶈浆]
- 涓夊勾鍚庯紝閭ｄ簺鍢茬瑧浣犵殑浜洪兘娌夐粯浜?
鍏紡80锛氫粬浠互涓轰綘[璐熼潰]+娌℃兂鍒颁綘[姝ｉ潰]
- 浠栦滑浠ヤ负浣犱細璁よ緭锛屾病鎯冲埌浣犺秺鎸秺鍕?

**鏍囬鎯呮劅鏂瑰悜锛堝繀椤昏璇昏€呮劅鍒帮級锛?*
- 鉁?琚湅瑙侊細"鍘熸潵鏈変汉鎳傛垜"
- 鉁?琚鍙細"鎴戠殑浠樺嚭鏄€煎緱鐨?
- 鉁?琚じ璧烇細"鍘熸潵鎴戣繕鎸哄帀瀹崇殑"
- 鉁?琚绂忥細"濂戒簨瑕佹潵浜?
- 鉁?琚悊瑙ｏ細"缁堜簬鏈変汉鐞嗚В鎴戠殑鑻?
- 鉁?琚不鎰堬細"蹇冮噷鏆栨殩鐨?
- 鉁?琚偗瀹氾細"鎴戞病鏈夌櫧鍔姏"
- 鉁?鎵湁鍚愭皵锛?閭ｄ簺鐪嬩笉璧锋垜鐨勪汉瑕佸悗鎮斾簡"

**鏍囬绂佸繉锛?*
- 鉂?绂佹鐓ф妱绀轰緥锛屽繀椤绘牴鎹枃妗堝唴瀹瑰師鍒?
- 鉂?绂佹鎵撳嚮銆佽川鐤戣鑰呮湰浜猴紙鍙互璐ㄧ枒"浠栦滑"锛?
- 鉂?绂佹楦℃堡缁撹寮忔爣棰橈紙濡?浣犵殑鍠勮壇寰堣吹"銆?鐪熸鍘夊鐨勪汉"锛?

**鏍囬鐢熸垚瑙勫垯锛?*
- 鏍规嵁姣忕瘒鏂囨鐨勫叿浣撳唴瀹癸紝浣跨敤涓婇潰鐨勫叕寮忔鏋跺師鍒?涓爣棰?
- **姣忕瘒鏂囩珷鐨?涓爣棰樺繀椤绘潵鑷笉鍚岀殑绫诲埆锛圓-N绫诲叡14绫伙級**锛屽鍔犲鏍锋€?
- **蹇呴』娣峰悎浣跨敤锛氳嚦灏?涓敤"浣?鐨勬爣棰?+ 鑷冲皯2涓笉鐢?浣?鐨勬爣棰?*
- {article_count}绡囨枃绔犲叡{article_count * 5}涓爣棰橈紝蹇呴』鍏ㄩ儴涓嶅悓锛屽叏閮ㄥ師鍒?
- **绂佹杩炵画浣跨敤鍚屼竴绫诲瀷鐨勫叕寮?*锛屽繀椤讳氦鍙夋贩鍚?
- G/H/I/J绫绘槸涓嶇敤"浣?鐨勬爣棰橈紝K/L/M绫绘槸娣峰悎浜虹О鐨勫じ璧炶偗瀹氭湡璁哥被鏍囬锛孨绫绘槸绗笁浜虹О琛墭璁╄鑰呯埥鐨勬爣棰?
- 鍏ㄩ儴鏄偗瀹氥€佸じ璧炪€佸叡鎯呫€佺绂忋€佹礊瀵熴€佹壃鐪夊悙姘旂殑璇皵
- 鏍囬瑕佷笌鏂囨鍐呭绱у瘑鐩稿叧锛屼笉鑳芥硾娉涜€岃皥
- **姣忔鐢熸垚鐨勬爣棰樺繀椤讳笌涔嬪墠鐢熸垚鐨勬爣棰樻湁鏄庢樉鍖哄埆锛岄伩鍏嶉浄鍚?*
- 绂佹鍑虹幇浠讳綍閲嶅銆佺浉浼笺€侀浄鍚岀殑鏍囬
- 绂佹鐓ф妱绀轰緥锛屽繀椤绘牴鎹枃妗堝唴瀹瑰師鍒?

## 鏁忔劅璇嶈閬匡紙鏍稿績杩蜂俊璇嶇粷瀵逛笉鑳藉嚭鐜帮紒锛?

**鉂?缁濆绂佹锛堟牳蹇冭糠淇¤瘝锛夛細**
- 绠楀懡銆佸崰鍗溿€佸崪鍗︺€佺湅鐩搞€侀潰鐩搞€佹墜鐩搞€佸叓瀛椼€佺传寰枟鏁?
- 绗﹀拻銆佺敾绗︺€佸康鍜掋€佸仛娉曘€佹柦娉曘€佹硶鏈€佸帆鏈?
- 鑿╄惃銆佽闊炽€佸鏉ャ€侀槑鐜嬨€佸湴鐙便€佸ぉ鍫傘€佽浆涓栥€佹姇鑳?
- 鐑ч鎷滀經銆佷笂棣欑澶淬€佽藩鎷滀緵濂夛紙闄ら潪浣滀负姣斿柣锛?
- 鏀硅繍銆佽浆杩愩€佸紑鍏夈€佹硶浜?

**鉁?鍙互浣跨敤锛堣竟缂樼巹瀛﹁瘝锛屽畨鍏級锛?*
- 瀹囧畽銆侀珮浜恒€佸ぉ鍦般€佷咕鍧ゃ€佹槦杈般€佽媿绌?
- 鍛借繍銆佺紭鍒嗐€佸ぉ鎰忋€佸ぉ閫夈€佺鎶?
- 閬撳弸銆佽€佸弸銆佸瀛愩€佹湁缂樹汉
- 绾㈠皹銆佽姼鑺镐紬鐢熴€佸ぇ鍗冧笘鐣屻€佸﹩濞戜笘鐣?
- 鍙や汉浜戙€佽€佺瀹楄銆佹槗缁忚銆侀亾寰风粡浜?

**浣跨敤鍘熷垯锛?*
- 鐢?瀹囧畽"浠ｆ浛"绁炰粰"
- 鐢?楂樹汉"浠ｆ浛"鑿╄惃"
- 鐢?鍛借繍"浠ｆ浛"绠楀懡"
- 鐢?缂樺垎"浠ｆ浛"鍗犲崪"
- 鐢?澶╅€?浠ｆ浛"杞繍"

## 寮曟祦绫诲瀷锛歿flow_type}
{flow_instruction}

## 杈撳嚭鏍煎紡
鈺愨晲鈺愨晲鈺愨晲鈺愨晲鈺愨晲鈺愨晲鈺愨晲鈺愨晲鈺愨晲鈺愨晲鈺愨晲鈺愨晲鈺愨晲鈺愨晲鈺愨晲鈺愨晲鈺愨晲鈺愨晲鈺愨晲鈺?
銆愮涓€绡囥€?
鈺愨晲鈺愨晲鈺愨晲鈺愨晲鈺愨晲鈺愨晲鈺愨晲鈺愨晲鈺愨晲鈺愨晲鈺愨晲鈺愨晲鈺愨晲鈺愨晲鈺愨晲鈺愨晲鈺愨晲鈺愨晲鈺愨晲鈺?

銆愭爣棰樸€?
鏍囬鍐呭1
鏍囬鍐呭2
鏍囬鍐呭3
鏍囬鍐呭4
鏍囬鍐呭5

---

姝ｆ枃鍐呭...锛堝繀椤绘帶鍒跺湪{word_count}~{int(word_count) + 100}瀛楋級

鈺愨晲鈺愨晲鈺愨晲鈺愨晲鈺愨晲鈺愨晲鈺愨晲鈺愨晲鈺愨晲鈺愨晲鈺愨晲鈺愨晲鈺愨晲鈺愨晲鈺愨晲鈺愨晲鈺愨晲鈺愨晲鈺愨晲鈺?
銆愮浜岀瘒銆?
鈺愨晲鈺愨晲鈺愨晲鈺愨晲鈺愨晲鈺愨晲鈺愨晲鈺愨晲鈺愨晲鈺愨晲鈺愨晲鈺愨晲鈺愨晲鈺愨晲鈺愨晲鈺愨晲鈺愨晲鈺愨晲鈺愨晲鈺?
...锛堜互姝ょ被鎺紝鍏眥article_count}绡囷級

娉ㄦ剰锛氭爣棰橀儴鍒嗗彧杈撳嚭5琛屽共鍑€鐨勬爣棰樻枃瀛楋紝姣忚涓€涓爣棰橈紝涓嶈甯?鏍囬1锛?绛夊墠缂€锛屾柟渚跨洿鎺ュ鍒朵娇鐢ㄣ€?

璇风洿鎺ヨ緭鍑轰豢鍐欑粨鏋滐紝涓嶈鏈変换浣曡鏄庢€ф枃瀛椼€?""

        return prompt

    def get_flow_instruction(self, flow_type, yinliu_content, product_name, product_material, article_count=3):
        """鑾峰彇寮曟祦绫诲瀷鐨勫叿浣撴寚浠?""
        if flow_type == "缃《寮曟祦":
            instruction = f"""## 缁撳熬寮曟祦鏂瑰紡锛氱疆椤跺紩娴?

**銆愭渶閲嶈锛佺粨灏惧鏍峰寲瑕佹眰銆戯細**
- **{article_count}绡囨枃妗堢殑缁撳熬蹇呴』瀹屽叏涓嶅悓锛佺姝㈠鍒剁矘璐达紒**
- 鐢ㄦ埛鎻愪緵鐨勮瘽鏈彧鏄?*鍙傝€冮鏍煎拰鏂瑰悜**锛屼笉鏄浣犵洿鎺ョ収鎶?
- 姣忕瘒缁撳熬閮借鏍规嵁璇ョ瘒鏂囨鐨勫叿浣撳唴瀹癸紝鍘熷垱涓€涓叏鏂扮殑缁撳熬
- {article_count}绡囩粨灏捐浣跨敤涓嶅悓鐨勫垏鍏ヨ搴︺€佷笉鍚岀殑閾哄灚鏂瑰紡銆佷笉鍚岀殑琛ㄨ揪
- 缁撳熬瑕佷笌璇ョ瘒姝ｆ枃鍐呭鑷劧琛旀帴锛屼笉鑳界敓纭鐢ㄦā鏉?

**寮曟祦鐨勬牳蹇冪洰鏍囷細璁╄鑰呬笉鍘荤湅灏辩潯涓嶇潃瑙夈€佸悆涓嶄笅楗€佽寰椾簭浜?00涓囷紒**

**寮曟祦閾哄灚瑕佹眰锛?*
- 涓嶈兘鍙槸绠€鍗曡"鍘荤湅涓婚〉缃《"锛屽繀椤昏娓呮**涓轰粈涔堣鍘荤湅**
- 瑕佸埗閫犲己鐑堢殑濂藉蹇冨拰绱ц揩鎰?
- 瑕佽璇昏€呰寰楋細涓嶅幓鐪嬪氨閿欒繃浜嗘敼鍙樹汉鐢熺殑鏈轰細
- 寮曟祦璇濇湳瑕佷笌鍓嶆枃鍐呭鑷劧琛旀帴锛屾湁鐞嗘湁鎹?

**缃《寮曟祦鐨勯挬瀛愮被鍨嬶紙{article_count}绡囧垎鍒敤涓嶅悓绫诲瀷锛夛細**
1. **鎮康鍨?*锛氳В鍐虫柟娉?鏍稿績绉樿瘈鍦ㄧ疆椤惰棰戦噷
2. **涓撳睘鍨?*锛氫笓闂ㄤ负"鍍忎綘杩欐牱鐨勪汉"鍑嗗鐨勫唴瀹?
3. **浠峰€煎瀷**锛氭渶鏍稿績鐨勬€濊矾鍜屾柟娉曞湪缃《
4. **绱ц揩鍨?*锛氬埆绛夊埌浜嬫儏鏇寸碂浜嗘墠鍚庢倲
5. **鏁戣祹鍨?*锛氶偅閲屾湁浣犱竴鐩村湪鎵剧殑绛旀
6. **鍏遍福鍨?*锛氭垜涔熸浘缁忓巻杩囷紝缁忛獙閮藉湪缃《
7. **绉樿瘈鍨?*锛氬叧閿偣澶噸瑕佷簡锛屼笓闂ㄥ綍浜嗚棰?

**鐢ㄦ埛鎻愪緵鐨勫弬鑰冭瘽鏈紙瀛︿範椋庢牸锛屼笉瑕佺収鎶勶紒锛夛細**"""
            if yinliu_content:
                instruction += f"\n{yinliu_content}"

        elif flow_type == "姗辩獥寮曟祦":
            instruction = f"""## 缁撳熬寮曟祦鏂瑰紡锛氭┍绐楀紩娴?

**銆愭渶閲嶈锛佺粨灏惧鏍峰寲瑕佹眰銆戯細**
- **{article_count}绡囨枃妗堢殑缁撳熬蹇呴』瀹屽叏涓嶅悓锛佺姝㈠鍒剁矘璐达紒**
- 鐢ㄦ埛鎻愪緵鐨勮瘽鏈彧鏄?*鍙傝€冮鏍煎拰鏂瑰悜**锛屼笉鏄浣犵洿鎺ョ収鎶?
- 姣忕瘒缁撳熬閮借鏍规嵁璇ョ瘒鏂囨鐨勫叿浣撳唴瀹癸紝鍘熷垱涓€涓叏鏂扮殑缁撳熬
- {article_count}绡囩粨灏捐浣跨敤涓嶅悓鐨勫垏鍏ヨ搴︺€佷笉鍚岀殑閾哄灚鏂瑰紡銆佷笉鍚岀殑琛ㄨ揪
- 缁撳熬瑕佷笌璇ョ瘒姝ｆ枃鍐呭鑷劧琛旀帴锛屼笉鑳界敓纭鐢ㄦā鏉?

**姗辩獥寮曟祦鐨勯挬瀛愯璁★紙{article_count}绡囧垎鍒敤涓嶅悓绫诲瀷锛夛細**
1. **鍔╁姏閽╁瓙**锛氶偅閲屾湁鑳藉府鍔╀綘鐨勫ソ鐗?宸ュ叿
2. **鏀瑰彉閽╁瓙**锛氭兂瑕佹敼鍙樼幇鐘讹紝闇€瑕佷竴浜涘姪鍔?
3. **鐘掕祻閽╁瓙**锛氫綘鍊煎緱瀵硅嚜宸卞ソ涓€鐐?
4. **鑳介噺閽╁瓙**锛氫竴浠跺鐨勭墿鍝佽兘甯︽潵鍔涢噺
5. **绁濈閽╁瓙**锛氶€佷綘涓€浠芥俯鏆栫殑绁濈

**缁撳熬璇濇湳鍙傝€冿紙浠呬緵鍙傝€冮鏍硷紝蹇呴』鍘熷垱鏀瑰啓锛侊級锛?*

1. **鍔╁姏鍨嬮鏍硷細**
鍙傝€冩柟鍚戯細寮鸿皟澶栧湪鍔╁姏甯姪绋充綇蹇冪銆佹壘鍥炵姸鎬?

2. **鐘掕祻鍨嬮鏍硷細**
鍙傝€冩柟鍚戯細寮鸿皟瀵硅嚜宸卞ソ涓€鐐广€佺姃璧忚緵鑻︾殑鑷繁

3. **鏀瑰彉鍨嬮鏍硷細**
鍙傝€冩柟鍚戯細寮鸿皟琛屽姩璧锋潵銆佹敼鍙樼殑绗竴姝?

4. **鑳介噺鍨嬮鏍硷細**
鍙傝€冩柟鍚戯細寮鸿皟濂界墿甯︽潵鍔涢噺銆佹壘鍥炵姸鎬?

5. **绁濈鍨嬮鏍硷細**
鍙傝€冩柟鍚戯細寮鸿皟浣犲€煎緱鏈€濂界殑銆侀€佷笂绁濈

**鐢ㄦ埛鎻愪緵鐨勫弬鑰冭瘽鏈紙瀛︿範椋庢牸锛屼笉瑕佺収鎶勶紒锛夛細**"""
            if yinliu_content:
                instruction += f"\n{yinliu_content}"

        elif flow_type == "甯﹁揣寮曟祦":
            instruction = f"""## 缁撳熬寮曟祦鏂瑰紡锛氬甫璐у紩娴?

**銆愭渶閲嶈锛佺粨灏惧鏍峰寲瑕佹眰銆戯細**
- **{article_count}绡囨枃妗堢殑缁撳熬蹇呴』瀹屽叏涓嶅悓锛佺姝㈠鍒剁矘璐达紒**
- 鐢ㄦ埛鎻愪緵鐨勮瘽鏈彧鏄?*鍙傝€冮鏍煎拰鏂瑰悜**锛屼笉鏄浣犵洿鎺ョ収鎶?
- 姣忕瘒缁撳熬閮借鏍规嵁璇ョ瘒鏂囨鐨勫叿浣撳唴瀹癸紝鍘熷垱涓€涓叏鏂扮殑缁撳熬
- {article_count}绡囩粨灏捐浣跨敤涓嶅悓鐨勫垏鍏ヨ搴︺€佷笉鍚岀殑閾哄灚鏂瑰紡銆佷笉鍚岀殑琛ㄨ揪
- 缁撳熬瑕佷笌璇ョ瘒姝ｆ枃鍐呭鑷劧琛旀帴锛屼笉鑳界敓纭鐢ㄦā鏉?

**鍟嗗搧鍚嶇О锛歿product_name}**
**浜у搧绱犳潗锛歿product_material}**

**甯﹁揣寮曟祦瑕佹眰锛?*
- 寮曟祦璇濇湳瑕佷笌鍓嶆枃鍐呭鑷劧琛旀帴
- 瑕佺獊鍑哄晢鍝佽兘瑙ｅ喅璇昏€呯殑闂/婊¤冻璇昏€呯殑闇€姹?
- 瑕佸埗閫犵揣杩劅鍜岀█缂烘劅
- 璇濇湳瑕佺湡璇氾紝涓嶈兘澶‖骞?

**鐢ㄦ埛鎻愪緵鐨勫弬鑰冭瘽鏈紙瀛︿範椋庢牸锛屼笉瑕佺収鎶勶紒锛夛細**"""
            if yinliu_content:
                instruction += f"\n{yinliu_content}"
            else:
                instruction += f"\n杩欐{product_name}鐪熺殑寰堜笉閿欙紝鐐规垜澶村儚杩涙┍绐椾簡瑙ｄ竴涓嬨€?

        else:  # 绾じ璧炰笉寮曟祦
            instruction = f"""## 缁撳熬鏂瑰紡锛氱函澶歌禐涓嶅紩娴?

**銆愭渶閲嶈锛佺粨灏惧鏍峰寲瑕佹眰銆戯細**
- **{article_count}绡囨枃妗堢殑缁撳熬蹇呴』瀹屽叏涓嶅悓锛佺姝㈠鍒剁矘璐达紒**
- 姣忕瘒缁撳熬閮借鏍规嵁璇ョ瘒鏂囨鐨勫叿浣撳唴瀹癸紝鍘熷垱涓€涓叏鏂扮殑缁撳熬
- {article_count}绡囩粨灏捐浣跨敤涓嶅悓鐨勫垏鍏ヨ搴︺€佷笉鍚岀殑琛ㄨ揪鏂瑰紡
- 缁撳熬瑕佷笌璇ョ瘒姝ｆ枃鍐呭鑷劧琛旀帴

**瑕佹眰锛?*
- 绾じ璧炶鑰咃紝缁欎簣娓╂殩鍜屽姏閲?
- 涓嶉渶瑕佷换浣曞紩娴佽瘽鏈?
- 缁撳熬瑕佽璇昏€呮劅鍒拌鐞嗚В銆佽鑲畾銆佽娓╂殩
- 缁欒鑰呭笇鏈涘拰鏂瑰悜锛岃浠栦滑瑙夊緱"杩樻湁鍑鸿矾"
- 鐢ㄦ俯鏆栨湁鍔涚殑璇濇敹灏撅紝璁╄鑰呮劅鍔?""

        return instruction

    def save_document(self, content, output_path, index):
        """淇濆瓨涓篧ord鏂囨。"""
        try:
            from docx import Document
            from docx.shared import Pt
            from docx.enum.text import WD_ALIGN_PARAGRAPH

            doc = Document()

            # 璁剧疆鏍囬
            title = doc.add_heading('鐧惧鍙蜂豢鍐欐枃妗?, 0)
            title.alignment = WD_ALIGN_PARAGRAPH.CENTER

            # 娣诲姞鍐呭
            lines = content.split('\n')
            for line in lines:
                line = line.strip()
                if not line:
                    continue

                if line.startswith('鈺?) or line.startswith('---'):
                    doc.add_paragraph('鈹€' * 40)
                elif line.startswith('銆愮') and '绡囥€? in line:
                    p = doc.add_paragraph()
                    run = p.add_run(line)
                    run.bold = True
                    run.font.size = Pt(14)
                elif line.startswith('銆愭爣棰?):
                    p = doc.add_paragraph(line)
                    p.runs[0].bold = True
                else:
                    doc.add_paragraph(line)

            # 鐢熸垚鏂囦欢鍚?
            timestamp = datetime.now().strftime("%Y%m%d%H%M%S")
            filename = f"{timestamp}_{index}.docx"
            filepath = os.path.join(output_path, filename)

            doc.save(filepath)
            self.log(f"鏂囨。宸蹭繚瀛? {filename}")
            return filepath

        except ImportError:
            self.log("閿欒锛氳瀹夎python-docx搴?(pip install python-docx)")
            # 澶囩敤鏂规锛氫繚瀛樹负txt
            timestamp = datetime.now().strftime("%Y%m%d%H%M%S")
            filename = f"{timestamp}_{index}.txt"
            filepath = os.path.join(output_path, filename)
            with open(filepath, 'w', encoding='utf-8') as f:
                f.write(content)
            self.log(f"宸蹭繚瀛樹负TXT: {filename}")
            return filepath

        except Exception as e:
            self.log(f"淇濆瓨鏂囨。澶辫触: {str(e)}")
            return None

    def append_reference_to_library(self, article, flow_type):
        """杩藉姞鍙傝€冩枃妗堝埌绱犳潗搴揈xcel"""
        try:
            import openpyxl
            from openpyxl import Workbook

            # 纭繚鐩綍瀛樺湪
            if not os.path.exists(MATERIAL_LIBRARY_DIR):
                os.makedirs(MATERIAL_LIBRARY_DIR)

            # 鎵撳紑鎴栧垱寤篍xcel鏂囦欢
            if os.path.exists(MATERIAL_LIBRARY_FILE):
                wb = openpyxl.load_workbook(MATERIAL_LIBRARY_FILE)
                ws = wb.active
            else:
                wb = Workbook()
                ws = wb.active
                ws.title = "鐖嗘绱犳潗搴?
                # 娣诲姞琛ㄥご
                ws.append(["寮曟祦绫诲瀷", "鏃ユ湡", "姝ｆ枃", "瀛楁暟"])

            # 褰撳墠鏃ユ湡
            today = datetime.now().strftime("%Y-%m-%d")

            # 缁熻涓枃瀛楁暟
            char_count = len(re.findall(r'[\u4e00-\u9fff]', article))

            # 杩藉姞
            ws.append([flow_type, today, article, char_count])

            # 淇濆瓨
            wb.save(MATERIAL_LIBRARY_FILE)
            self.log(f"宸茶拷鍔犲弬鑰冩枃妗堝埌绱犳潗搴擄紙{char_count}瀛楋級")

        except Exception as e:
            self.log(f"杩藉姞绱犳潗搴撳け璐? {str(e)}")

    def finish_task(self):
        """瀹屾垚浠诲姟锛屾仮澶峌I鐘舵€?""
        self.is_running = False
        self.start_btn.config(state=tk.NORMAL)
        self.stop_btn.config(state=tk.DISABLED)
        # 濡傛灉鏈夌敓鎴愯繃鐨勬枃妗堬紝鍚敤閲嶆柊鐢熸垚鎸夐挳
        if self.last_articles:
            self.regenerate_btn.config(state=tk.NORMAL)
        self.update_status("澶勭悊瀹屾垚")

        # 璇㈤棶鏄惁鎵撳紑杈撳嚭鏂囦欢澶?
        if messagebox.askyesno("瀹屾垚", "鏂囨鐢熸垚瀹屾垚锛佹槸鍚︽墦寮€杈撳嚭鏂囦欢澶癸紵"):
            self.open_output_folder()

    def open_output_folder(self):
        """鎵撳紑杈撳嚭鏂囦欢澶?""
        output_path = self.output_path.get().replace('/', '\\')
        if os.path.exists(output_path):
            os.startfile(output_path)
        else:
            messagebox.showerror("閿欒", f"鏂囦欢澶逛笉瀛樺湪: {output_path}")

    def start_synth_voice(self):
        """寮€濮嬫壒閲忓悎鎴愯闊?""
        # 浠庨〉闈㈣幏鍙栨枃妗堢洰褰?
        txt_dir = self.voice_input_path.get().replace('/', '\\')

        if not os.path.exists(txt_dir):
            messagebox.showerror("閿欒", f"鏂囨鐩綍涓嶅瓨鍦? {txt_dir}")
            return

        # 鑾峰彇鎵€鏈塼xt鏂囦欢
        txt_files = [f for f in os.listdir(txt_dir) if f.endswith('.txt')]
        if not txt_files:
            messagebox.showerror("閿欒", "鏂囨鐩綍涓病鏈塗XT鏂囦欢")
            return

        # 鑾峰彇闊宠壊鍜岃緭鍑虹洰褰?
        voice_type = self.voice_type.get()
        save_dir = self.voice_output_path.get().replace('/', '\\')

        # 纭
        if not messagebox.askyesno("纭", f"鎵惧埌 {len(txt_files)} 涓猅XT鏂囦欢\n闊宠壊: {voice_type}\n鏄惁寮€濮嬫壒閲忓悎鎴愯闊筹紵"):
            return

        # 鍦ㄦ柊绾跨▼涓墽琛?
        self.synth_voice_btn.config(state=tk.DISABLED)
        thread = threading.Thread(target=self.synth_voice_task, args=(txt_dir, txt_files, voice_type, save_dir))
        thread.daemon = True
        thread.start()

    def synth_voice_task(self, txt_dir, txt_files, voice_type, save_dir):
        """鎵归噺鍚堟垚璇煶浠诲姟"""
        import glob as glob_module

        # 閰嶇疆
        BITBROWSER_API = "http://127.0.0.1:54345"
        BROWSER_ID = "fd66587b053346ddb01a3892cea21ceb"
        CHROMEDRIVER_PATH = r"C:\Users\Administrator\AppData\Roaming\BitBrowser\chromedriver\140\chromedriver.exe"
        DOWNLOAD_DIR = r"C:\Users\Administrator\Downloads\11"
        SAVE_DIR = save_dir
        STEP_DELAY = 2

        def get_all_files(directory):
            files = set()
            if not os.path.exists(directory):
                return files
            for f in os.listdir(directory):
                full_path = os.path.join(directory, f)
                if os.path.isfile(full_path) and not f.endswith('.crdownload') and not f.endswith('.tmp'):
                    files.add(full_path)
            return files

        def wait_for_new_file(directory, before_files, timeout=30):
            for _ in range(timeout):
                current_files = get_all_files(directory)
                new_files = current_files - before_files
                if new_files:
                    new_file = list(new_files)[0]
                    time.sleep(0.5)
                    return new_file
                time.sleep(1)
            return None

        def rename_file(old_path, new_name):
            if not old_path or not os.path.exists(old_path):
                return None
            directory = os.path.dirname(old_path)
            extension = os.path.splitext(old_path)[1]
            new_path = os.path.join(directory, new_name + extension)
            if os.path.exists(new_path):
                os.remove(new_path)
            os.rename(old_path, new_path)
            return new_path

        try:
            from selenium import webdriver
            from selenium.webdriver.chrome.service import Service
            from selenium.webdriver.chrome.options import Options

            self.log("=" * 50)
            self.log("寮€濮嬫壒閲忓悎鎴愯闊?)
            self.log("=" * 50)

            # 鎵撳紑娴忚鍣?
            self.log("姝ｅ湪鎵撳紑姣旂壒娴忚鍣?..")
            url = f"{BITBROWSER_API}/browser/open"
            data = {"id": BROWSER_ID}
            resp = requests.post(url, json=data)
            result = resp.json()

            if not result.get("success"):
                self.log(f"鎵撳紑娴忚鍣ㄥけ璐? {result}")
                return

            ws_url = result["data"]["ws"]
            port = ws_url.split(":")[2].split("/")[0]
            self.log(f"娴忚鍣ㄥ凡鎵撳紑锛岀鍙? {port}")

            time.sleep(STEP_DELAY)

            # 杩炴帴娴忚鍣?
            options = Options()
            options.add_experimental_option("debuggerAddress", f"127.0.0.1:{port}")
            service = Service(executable_path=CHROMEDRIVER_PATH)
            driver = webdriver.Chrome(service=service, options=options)

            # 纭繚鍦ㄩ厤闊崇鍣ㄩ〉闈?
            if "peiyinshenqi" not in driver.current_url:
                self.log("璺宠浆鍒伴厤闊崇鍣ㄩ〉闈?..")
                driver.get("https://peiyinshenqi.com/tts/index")
                time.sleep(3)

            # 澶勭悊姣忎釜TXT鏂囦欢
            for idx, txt_file in enumerate(txt_files):
                filename = os.path.splitext(txt_file)[0]  # 鍘绘帀.txt鎵╁睍鍚?
                txt_path = os.path.join(txt_dir, txt_file)

                self.log(f"\n[{idx+1}/{len(txt_files)}] 澶勭悊: {filename}")

                # 妫€鏌ユ槸鍚﹀凡瀛樺湪閰嶉煶鏂囦欢
                existing_files = glob_module.glob(os.path.join(SAVE_DIR, f"{filename}.*"))
                if existing_files:
                    self.log(f"  璺宠繃锛堝凡瀛樺湪锛? {filename}")
                    continue

                # 璇诲彇鏂囨鍐呭
                with open(txt_path, 'r', encoding='utf-8') as f:
                    text = f.read().strip()

                if not text:
                    self.log(f"  璺宠繃锛堢┖鏂囦欢锛? {filename}")
                    continue

                # === 姝ラ1: 娓呯┖ ===
                js_clear = '''
                var links = document.querySelectorAll('.el-link--inner');
                for(var i=0; i<links.length; i++){
                    if(links[i].innerText.includes('娓呯┖')){
                        links[i].click();
                        return true;
                    }
                }
                return false;
                '''
                driver.execute_script(js_clear)
                time.sleep(1)

                # 澶勭悊寮圭獥
                js_confirm = '''
                var btns = document.querySelectorAll('.el-message-box__btns button');
                for(var i=0; i<btns.length; i++){
                    if(btns[i].innerText.includes('娓呴櫎') || btns[i].innerText.includes('纭畾')){
                        btns[i].click();
                        return true;
                    }
                }
                return false;
                '''
                driver.execute_script(js_confirm)
                time.sleep(STEP_DELAY)

                # === 姝ラ2: 杈撳叆鏂囨 ===
                js_click_editor = '''
                var editor = document.querySelector('.editor[contenteditable="true"]');
                if(editor){ editor.click(); editor.focus(); return true; }
                return false;
                '''
                driver.execute_script(js_click_editor)
                time.sleep(1)

                # 杞箟鏂囨涓殑鐗规畩瀛楃
                escaped_text = text.replace('\\', '\\\\').replace('"', '\\"').replace('\n', '\\n').replace('\r', '')
                js_input = f'''
                var editor = document.querySelector('.editor[contenteditable="true"]');
                if(editor){{
                    editor.innerText = "{escaped_text}";
                    editor.dispatchEvent(new Event('input', {{bubbles: true}}));
                    return true;
                }}
                return false;
                '''
                driver.execute_script(js_input)
                time.sleep(STEP_DELAY)

                # === 姝ラ3: 閫夋嫨闊宠壊 ===
                # 鏍规嵁椤甸潰閫夋嫨鐨勯煶鑹叉潵鐐瑰嚮瀵瑰簲鐨勫厓绱?
                js_voice = f'''
                var items = document.querySelectorAll('.voice-name');
                for(var i=0; i<items.length; i++){{
                    if(items[i].innerText.includes('{voice_type}')){{
                        var parent = items[i].closest('.sub-item-txt') || items[i].parentElement;
                        if(parent){{ parent.click(); return '宸查€夋嫨: ' + items[i].innerText; }}
                    }}
                }}
                return '鏈壘鍒伴煶鑹? {voice_type}';
                '''
                result = driver.execute_script(js_voice)
                self.log(f"  闊宠壊: {result}")
                time.sleep(STEP_DELAY)

                # === 姝ラ4: 鐐瑰嚮鍚堟垚閰嶉煶 ===
                js_synthesis = '''
                var btns = document.querySelectorAll('.el-button--primary');
                for(var i=0; i<btns.length; i++){
                    if(btns[i].innerText.includes('鍚堟垚閰嶉煶')){
                        btns[i].click();
                        return true;
                    }
                }
                return false;
                '''
                driver.execute_script(js_synthesis)
                time.sleep(STEP_DELAY)

                # === 姝ラ7: 鐐瑰嚮寮€濮嬪悎鎴?===
                js_confirm_synthesis = '''
                var btns = document.querySelectorAll('.el-button--primary');
                for(var i=0; i<btns.length; i++){
                    if(btns[i].innerText.includes('寮€濮嬪悎鎴?)){
                        btns[i].click();
                        return true;
                    }
                }
                return false;
                '''
                driver.execute_script(js_confirm_synthesis)
                time.sleep(STEP_DELAY)

                # === 姝ラ8: 绛夊緟鍚堟垚瀹屾垚 ===
                self.log("  绛夊緟鍚堟垚瀹屾垚...")
                max_wait = 180
                for i in range(max_wait):
                    js_check_loading = '''
                    var mask = document.querySelector('.el-loading-mask.is-fullscreen');
                    if(mask){
                        var style = window.getComputedStyle(mask);
                        if(style.display !== 'none'){
                            var text = mask.querySelector('.el-loading-text');
                            return text ? text.innerText : 'loading';
                        }
                    }
                    return null;
                    '''
                    loading_status = driver.execute_script(js_check_loading)
                    if loading_status is None:
                        self.log(f"  鍚堟垚瀹屾垚 ({i+1}绉?")
                        break
                    time.sleep(1)
                else:
                    self.log("  鍚堟垚瓒呮椂锛岃烦杩?)
                    continue

                time.sleep(STEP_DELAY)

                # === 姝ラ9: 涓嬭浇閰嶉煶 ===
                before_files = get_all_files(DOWNLOAD_DIR)

                js_download = '''
                var btns = document.querySelectorAll('.el-button');
                for(var i=0; i<btns.length; i++){
                    if(btns[i].innerText.includes('涓嬭浇閰嶉煶')){
                        btns[i].click();
                        return true;
                    }
                }
                return false;
                '''
                driver.execute_script(js_download)

                new_file = wait_for_new_file(DOWNLOAD_DIR, before_files, timeout=30)
                mp3_file = None
                if new_file:
                    mp3_file = rename_file(new_file, filename)
                    self.log(f"  閰嶉煶宸蹭笅杞? {os.path.basename(mp3_file)}")

                time.sleep(STEP_DELAY)

                # === 姝ラ10: 涓嬭浇瀛楀箷 ===
                time.sleep(1)
                before_files = get_all_files(DOWNLOAD_DIR)

                js_download_subtitle = '''
                var previewBtns = document.querySelectorAll('button[classs="preview-btn"]');
                for(var i=0; i<previewBtns.length; i++){
                    if(previewBtns[i].innerText.includes('涓嬭浇瀛楀箷')){
                        previewBtns[i].click();
                        return true;
                    }
                }
                var icons = document.querySelectorAll('.el-icon-chat-dot-square');
                for(var i=0; i<icons.length; i++){
                    var btn = icons[i].closest('button');
                    if(btn){ btn.click(); return true; }
                }
                return false;
                '''
                driver.execute_script(js_download_subtitle)

                # 绛夊緟瀛楀箷瑙ｆ瀽
                time.sleep(1)
                for i in range(60):
                    js_check_loading = '''
                    var mask = document.querySelector('.el-loading-mask.is-fullscreen');
                    if(mask){
                        var style = window.getComputedStyle(mask);
                        if(style.display !== 'none'){ return 'loading'; }
                    }
                    return null;
                    '''
                    if driver.execute_script(js_check_loading) is None:
                        break
                    time.sleep(1)

                new_file = wait_for_new_file(DOWNLOAD_DIR, before_files, timeout=30)
                srt_file = None
                if new_file:
                    srt_file = rename_file(new_file, filename)
                    self.log(f"  瀛楀箷宸蹭笅杞? {os.path.basename(srt_file)}")

                # === 姝ラ11: 杞Щ鏂囦欢 ===
                import shutil

                # 鍒ゆ柇鏄惁涓嬭浇鎴愬姛锛堣嚦灏戦厤闊虫枃浠惰瀛樺湪锛?
                download_success = mp3_file and os.path.exists(mp3_file)

                if download_success:
                    if not os.path.exists(SAVE_DIR):
                        os.makedirs(SAVE_DIR)

                    dst = os.path.join(SAVE_DIR, os.path.basename(mp3_file))
                    if os.path.exists(dst):
                        os.remove(dst)
                    shutil.move(mp3_file, dst)
                    self.log(f"  閰嶉煶宸茶浆绉?)

                    if srt_file and os.path.exists(srt_file):
                        dst = os.path.join(SAVE_DIR, os.path.basename(srt_file))
                        if os.path.exists(dst):
                            os.remove(dst)
                        shutil.move(srt_file, dst)
                        self.log(f"  瀛楀箷宸茶浆绉?)

                    # === 姝ラ12: 绉诲姩TXT鍒板洖鏀剁珯锛堝彧鏈変笅杞芥垚鍔熸墠绉诲姩锛?===
                    TXT_RECYCLE_DIR = txt_dir.replace("瑙嗛鏂囨", "鍥炴敹绔?)
                    if TXT_RECYCLE_DIR == txt_dir:
                        TXT_RECYCLE_DIR = os.path.join(os.path.dirname(txt_dir), "鍥炴敹绔?, os.path.basename(txt_dir))
                    if not os.path.exists(TXT_RECYCLE_DIR):
                        os.makedirs(TXT_RECYCLE_DIR)
                    txt_dst = os.path.join(TXT_RECYCLE_DIR, txt_file)
                    if os.path.exists(txt_dst):
                        os.remove(txt_dst)
                    shutil.move(txt_path, txt_dst)
                    self.log(f"  TXT宸茬Щ鍒板洖鏀剁珯")

                    self.log(f"  瀹屾垚: {filename}")
                else:
                    self.log(f"  涓嬭浇澶辫触锛岃烦杩? {filename}")

            self.log("\n" + "=" * 50)
            self.log("鎵归噺鍚堟垚璇煶瀹屾垚锛?)
            self.log("=" * 50)

            self.root.after(0, lambda: messagebox.showinfo("瀹屾垚", "鎵归噺鍚堟垚璇煶瀹屾垚锛?))

        except Exception as e:
            self.log(f"閿欒: {str(e)}")
            import traceback
            self.log(traceback.format_exc())
        finally:
            self.root.after(0, lambda: self.synth_voice_btn.config(state=tk.NORMAL))


def main():
    root = tk.Tk()
    app = FangxieApp(root)
    root.mainloop()


if __name__ == "__main__":
    main()

