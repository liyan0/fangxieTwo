# -*- coding: utf-8 -*-
"""
百家号文案+图片生成器（不生成视频）
功能：1篇参考文案 → 3篇仿写文案 → 每篇4张图片 → Word文档
"""

import os
import requests
import time
import urllib3
from datetime import datetime
from docx import Document
from docx.shared import Inches, Pt
from docx.enum.text import WD_ALIGN_PARAGRAPH

# 禁用SSL警告
urllib3.disable_warnings(urllib3.exceptions.InsecureRequestWarning)

# 配置
JIMENG_API_URL = "https://jimengly.zeabur.app/v1/images/generations"
JIMENG_API_KEY = "31f6ed72dd0f109538bea4323ba48a5f,e4377e9da2db16832ab65270fddccbd9,240ddc9646fbd088634ded292572ab67"
JIMENG_MODEL = "jimeng"

OUTPUT_DIR = "D:/A百家号带货视频/成品文案"
TEMP_DIR = "D:/A百家号带货视频/临时文件/images"


def ensure_dirs():
    os.makedirs(OUTPUT_DIR, exist_ok=True)
    os.makedirs(TEMP_DIR, exist_ok=True)


# ==================== 图片生成 ====================

def generate_image_prompts():
    """图片提示词 - 宇宙仙侠风格"""
    return [
        "宇宙星空背景，一个孤独的身影站在悬崖边，仙侠风格，紫色星云，金色光芒，史诗感，电影级画质",
        "浩瀚宇宙中漂浮的仙山，云雾缭绕，星河璀璨，一道人影负手而立，仙侠古风，神秘氛围",
        "星空下的古老宫殿，银河倒映，仙鹤飞舞，紫金色调，东方玄幻风格，大气磅礴",
        "宇宙深处的修仙者，周身环绕星辰，长袍飘逸，背对观众望向星河，孤独而坚定，仙侠史诗风格"
    ]


def generate_image(prompt, save_path):
    """调用即梦API生成图片"""
    headers = {
        "Authorization": f"Bearer {JIMENG_API_KEY}",
        "Content-Type": "application/json"
    }
    payload = {
        "model": JIMENG_MODEL,
        "prompt": prompt,
        "n": 1,
        "ratio": "16:9"
    }

    try:
        response = requests.post(JIMENG_API_URL, headers=headers, json=payload, timeout=180, verify=False)
        if response.status_code == 200:
            result = response.json()
            if result.get("data") and len(result["data"]) > 0:
                image_url = result["data"][0].get("url")
                if image_url:
                    img_response = requests.get(image_url, timeout=60, verify=False)
                    with open(save_path, "wb") as f:
                        f.write(img_response.content)
                    return True
        return False
    except Exception as e:
        print(f"    错误: {e}")
        return False


def generate_images_for_article(article_index):
    """为一篇文章生成4张图片"""
    prompts = generate_image_prompts()
    image_paths = []

    for i, prompt in enumerate(prompts):
        save_path = os.path.join(TEMP_DIR, f"article{article_index}_img{i+1}.png")
        print(f"    生成图片 {i+1}/4: {prompt[:20]}...")

        if generate_image(prompt, save_path):
            image_paths.append(save_path)
            print(f"    [OK] 图片 {i+1} 完成")
        else:
            print(f"    [X] 图片 {i+1} 失败")

        time.sleep(2)  # 避免API限流

    return image_paths


# ==================== 文案模板 ====================

# 文案1：一人战胜一群人
WENAN_TEMPLATE_1 = [
    # 第1篇
    {
        "title": "那个被所有人否定的人，后来怎么样了",
        "content": """停下来，我有几句话，不说憋得慌。

你是不是也有过这样的时刻：明明自己没做错什么，却被所有人否定？明明你比谁都努力，却没有一个人站出来替你说句公道话？

那种感觉，我太懂了。

你站在人群中间，所有人都在朝一个方向走，只有你，倔强地站在原地。他们用异样的眼神看你，背后的议论你听得清清楚楚——"这人脑子有问题"、"就他能耐"、"等着看笑话吧"。你不是不想解释，而是你知道，解释了也没用。他们根本不想听你说什么，他们只想看到你失败，只想证明他们是对的，你是错的。

这些年，你一个人扛过来的东西，没人知道。

凌晨三点睡不着的时候，你躺在床上盯着天花板，脑子里全是那些质疑的声音。"你凭什么？""你以为你是谁？""别做梦了，认清现实吧。"这些话像刀子一样，一遍遍割着你的心。你问自己：我是不是真的错了？我是不是应该放弃？我是不是太高估自己了？那些夜晚，你一个人躲在被子里，眼泪流了一枕头，第二天还要装作若无其事的样子面对所有人。没人知道你有多难，也没人想知道。

但你知道吗？你能走到今天，本身就是一种胜利。

大多数人在遇到阻力的时候，第一反应是退缩。他们会说"算了吧"、"何必呢"、"差不多得了"。但你不一样。你心里有一团火，怎么浇都浇不灭。那些否定你的话，不但没有让你停下来，反而成了你前进的燃料。你把委屈咽进肚子里，把眼泪擦干净，然后咬着牙继续往前走。这份倔强，这份不服输，是多少人想学都学不来的。

王阳明说过："此心光明，亦复何言。"只要你内心坦荡，问心无愧，又何必在意别人怎么说？那些质疑你的人，他们站在岸上，从来没有下过水，却对你在水里的姿势指指点点。他们不知道水有多深，不知道暗流有多急，不知道你为了不被淹没付出了多少。他们只会说"你游得不好看"，却从来不问一句"你累不累"。

但你不需要他们问。你早就学会了自己扛。

我见过太多像你这样的人。在单位里被排挤，因为不愿意同流合污；在家里被否定，因为选择不符合别人的期待；在亲戚朋友面前被看不起，因为走的路和别人不一样。但就是这些人，最后往往活得最通透。因为他们不是在跟别人比，他们是在跟自己的心走。

《菜根谭》里说："伏久者飞必高，开先者谢独早。"那些厚积薄发的人，往往要经历更长的沉默期。你现在的默默无闻，不是失败，而是在为将来积蓄力量。那些嘲笑你的人，他们只看到你现在的处境，看不到你未来的辽阔。

你没有铠甲，就把伤疤当成勋章；没有后援，就把自己活成一支队伍。那些独自熬过的夜，那些咬着牙扛过的难关，那些没人看见的崩溃和重建，都在悄悄塑造着一个更强大的你。你以为自己只是在苦撑，其实你已经在蜕变。

如果你正处在这样的阶段——被误解、被否定、被孤立，感觉全世界都在跟你作对，我想告诉你：你走的路没有错，只是走在了大多数人前面。

关于如何在质疑声中保持内心的笃定，如何把这些年吃的苦转化为后半生的底气，我在主页置顶的那条视频里，说了一些可能会改变你后半生的话。很多人看完都说，早点看到就好了。你也去看看，也许那里面有一句话，正是你现在最需要听到的。

点我头像，置顶第一条，等你。"""
    },
    # 第2篇
    {
        "title": "为什么越善良的人，越容易被辜负",
        "content": """我说句扎心的话，你别不爱听。

你有没有发现，这辈子你为别人做了那么多，到头来最累的还是自己？你掏心掏肺对别人好，换来的却是理所当然？你委屈自己成全所有人，却没有一个人真正心疼过你？

如果你也是这样的人，这条内容你一定要看完。

从小到大，你就是那个懂事的人。小时候，别的孩子在撒娇，你已经学会了察言观色；长大后，别人在享受生活，你在扛起一个家。你省吃俭用，把最好的留给孩子、留给家人；你任劳任怨，从来不说一句累。你以为只要自己付出够多，别人就会看见你的好，就会记着你的情。

可结果呢？

你起早贪黑忙了一辈子，换来的是"你做的还不够"。你掏心掏肺对别人好，换来的是"这是你应该的"。你委屈自己成全所有人，换来的是没有一个人问你"你累不累"。那些你帮过的人，转头就忘了你的好；那些你养大的孩子，觉得你的付出理所当然；那些你迁就了一辈子的人，从来没有真正把你放在心上。

我知道，你心里有多苦。

夜深人静的时候，你也会想：我这辈子图什么？我付出了这么多，值得吗？为什么我对别人那么好，别人却从来不把我当回事？这些话你不敢说出口，因为你怕别人说你矫情，说你计较，说你不够大度。你只能一个人扛着，一个人消化，一个人在没人的时候偷偷抹眼泪。

但我想告诉你一句话：你的付出，不是没有价值，只是给错了人。

曾国藩说过："轻财足以聚人，律己足以服人，量宽足以得人，身先足以率人。"你这辈子做到了所有，唯独忘了善待自己。你把所有的温柔都给了别人，却从来没有留一点给自己。你值得被珍惜，值得被善待，值得有人心疼你、在乎你、把你放在心尖上。

不是你不够好，是有些人根本不配。

那些不懂得感恩的人，你对他再好也是白搭。那些只知道索取的人，你付出再多他也觉得不够。与其把真心浪费在不值得的人身上，不如留一些给自己。你已经辛苦了大半辈子，接下来的日子，该为自己活一回了。

杨绛先生说："我们曾如此渴望外界的认可，到最后才知道，世界是自己的，与他人毫无关系。"你不需要向任何人证明什么，你也不欠任何人一个交代。你这辈子吃过的苦、受过的累、扛过的委屈，老天都看在眼里，总有一天会还给你。

那些亏待你的人，不值得你再为他们伤心；那些不懂你的人，不值得你再为他们解释。从今往后，把心收回来，好好爱自己。

关于如何在下半场人生里活出自己，如何不再为不值得的人消耗自己，我在主页置顶的那条视频里说得更透。那里面有几句话，是我特别想说给你听的，很多人看完都哭了。如果你也想给自己的心找个出口，去看看吧。

点我头像，看置顶第一条，我在那里等你。"""
    },
    # 第3篇
    {
        "title": "你吃过的苦，老天都记着账",
        "content": """我想问你一个问题：这辈子，你有没有好好心疼过自己？

先别急着回答，听我把话说完。

你总是在照顾别人。照顾父母，照顾孩子，照顾爱人，照顾身边所有需要你的人。你把自己活成了一棵大树，所有人都在你的树荫下乘凉，却没有人问过你：站了这么久，你累不累？根扎得那么深，你疼不疼？

你当然累，你当然疼。可你从来不说。

年轻的时候，你以为咬咬牙就过去了。孩子要上学，你省吃俭用供他读书；家里要用钱，你起早贪黑拼命赚；遇到难处了，你一个人扛着，从来不跟别人诉苦。你觉得这是你的责任，是你应该做的。你把所有的苦都咽进肚子里，把所有的累都扛在肩上，硬生生把自己熬成了现在的样子。

可是，有谁真正看见过你的付出？

那些你养大的孩子，觉得你的付出是理所当然；那些你帮过的人，转头就把你忘了；那些你牺牲自己成全的人，从来没有说过一句谢谢。你不是没有委屈，你只是不敢说。你怕说出来显得自己矫情，怕说出来让别人觉得你在邀功。于是你把所有的委屈都憋在心里，一个人默默消化，一个人偷偷流泪。

但我想告诉你：你的苦，不是白吃的。你流的每一滴汗，熬的每一个夜，扛的每一份重担，老天都在给你记着账。

《道德经》里说："天之道，损有余而补不足。"你这辈子付出的每一分每一秒，都在悄悄地为你积攒着什么。也许现在你还看不到，但总有一天，这些都会以另一种方式回报给你。那些你扛过的苦难，会变成你的铠甲；那些你熬过的夜晚，会变成你的勋章；那些你流过的眼泪，会变成你后半生的底气。

你知道吗？能扛事的人，才是真正厉害的人。

这个世界上，有太多人遇到一点困难就退缩，遇到一点委屈就崩溃，遇到一点挫折就放弃。但你不一样。你被生活打倒过无数次，每一次你都爬了起来；你被人否定过无数次，每一次你都咬牙扛了过来；你在黑暗里走过无数个夜晚，每一次你都等到了天亮。这份韧性，这份坚强，是多少人一辈子都学不会的。

季羡林先生说过："人生在世，有些事情是不必在乎的，有些东西是必须清空的。"那些不理解你的人，不必在乎；那些辜负你的人，必须清空。你已经为别人活了大半辈子，接下来的日子，该为自己活了。

不要再委屈自己了。你值得被善待，值得被珍惜，值得拥有幸福的晚年。那些你吃过的苦，会变成你后半生的甜；那些你扛过的难，会变成你余生的福气。

关于如何在下半场人生里放下执念、善待自己，如何把前半生的苦熬成后半生的福，我在主页置顶的那条视频里聊得更深。那里有一些话，是我特别想说给像你这样的人听的。很多人看完说，感觉有人终于懂我了。

点我头像，置顶第一条，你一定要去看看。你这辈子最应该心疼的人，是你自己。"""
    }
]


# ==================== Word文档生成 ====================

def create_word_document(articles_with_images, output_path):
    """创建Word文档，包含文案和图片"""
    doc = Document()

    # 总标题
    title = doc.add_heading('百家号引流文案 + 配图', 0)
    title.alignment = WD_ALIGN_PARAGRAPH.CENTER

    doc.add_paragraph(f"生成时间：{datetime.now().strftime('%Y-%m-%d %H:%M:%S')}")
    doc.add_paragraph()

    for i, article in enumerate(articles_with_images):
        # 文章标题
        doc.add_heading(f"【第{i+1}篇】{article['title']}", level=1)
        doc.add_paragraph()

        # 文案内容
        doc.add_heading("文案内容：", level=2)
        for para in article['content'].split('\n\n'):
            if para.strip():
                p = doc.add_paragraph(para.strip())
                p.paragraph_format.line_spacing = 1.5

        doc.add_paragraph()

        # 配图
        doc.add_heading("配图（4张）：", level=2)
        for j, img_path in enumerate(article['images']):
            if os.path.exists(img_path):
                doc.add_paragraph(f"图片 {j+1}:")
                doc.add_picture(img_path, width=Inches(5.5))
                doc.add_paragraph()

        # 分隔线
        doc.add_paragraph("─" * 50)
        doc.add_paragraph()

    doc.save(output_path)
    return output_path


# ==================== 主函数 ====================

def main():
    """主函数：生成3篇文案 + 每篇4张图片 → Word文档"""
    print("=" * 60)
    print("百家号文案+图片生成器")
    print("=" * 60)

    ensure_dirs()
    timestamp = datetime.now().strftime("%Y%m%d_%H%M%S")

    articles_with_images = []

    # 使用模板生成3篇文案
    for i, article in enumerate(WENAN_TEMPLATE_1):
        print(f"\n【第{i+1}篇】{article['title']}")
        print("-" * 40)

        # 生成4张图片
        print("  生成图片...")
        image_paths = generate_images_for_article(i + 1)

        articles_with_images.append({
            "title": article['title'],
            "content": article['content'],
            "images": image_paths
        })

        print(f"  [OK] 完成，共{len(image_paths)}张图片")

    # 生成Word文档
    print("\n" + "=" * 60)
    print("生成Word文档...")
    output_path = os.path.join(OUTPUT_DIR, f"{timestamp}_文案和配图.docx")
    create_word_document(articles_with_images, output_path)

    print("\n" + "=" * 60)
    print("全部完成！")
    print(f"输出文件：{output_path}")
    print("=" * 60)

    return output_path


if __name__ == "__main__":
    main()
